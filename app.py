from __future__ import annotations

import io
import json
import base64
import asyncio
import hashlib
import hmac
import re
import subprocess
import threading
import time
import uuid
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import BadZipFile, ZipFile

import httpx
import uvicorn
try:
    import lark_oapi as lark
    import lark_oapi.ws.client as lark_ws_client
    from lark_oapi.api.im.v1 import P2ImMessageReceiveV1
except ImportError:
    lark = lark_ws_client = None
    P2ImMessageReceiveV1 = None
try:
    from cryptography.hazmat.primitives import padding
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
except ImportError:
    Cipher = algorithms = modes = padding = None
try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None
    WD_CELL_VERTICAL_ALIGNMENT = WD_TABLE_ALIGNMENT = WD_ALIGN_PARAGRAPH = None
    OxmlElement = qn = Inches = Pt = RGBColor = None
from fastapi import BackgroundTasks, Cookie, Depends, FastAPI, File, Header, HTTPException, Request, Response, UploadFile
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from openpyxl.styles import Alignment, Border, Font, Side
from pydantic import BaseModel, Field
from pypdf import PdfReader

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / 'static'
TEMPLATE_DIR = BASE_DIR / 'templates'
PACKAGE_TEMPLATE_DIR = BASE_DIR / '模板文件'
DATA_DIR = BASE_DIR / 'data'
TMP_DIR = DATA_DIR / 'tmp'
PROJECTS_DIR = DATA_DIR / 'projects'
MARKET_SKILL_DIR = DATA_DIR / 'market_skill'
FEISHU_DIR = DATA_DIR / 'feishu'
FEISHU_SESSIONS_DIR = FEISHU_DIR / 'sessions'
FEISHU_EVENTS_DIR = FEISHU_DIR / 'events'
FEISHU_AGENT_LOGS_DIR = FEISHU_DIR / 'agent_logs'
CONFIG_PATH = DATA_DIR / 'config.json'
USERS_PATH = DATA_DIR / 'users.json'
PACKAGE_EXTRACTION_TEMPLATE_PATH = PACKAGE_TEMPLATE_DIR / 'extraction_template.xlsx'
PACKAGE_REGISTER_TEMPLATE_PATH = PACKAGE_TEMPLATE_DIR / '招标登记.xlsx'
SESSION_TTL = timedelta(hours=4)
FEISHU_SESSION_TTL = timedelta(hours=24)
FEISHU_EVENT_TTL = timedelta(days=7)
AUTH_TTL = timedelta(hours=12)
PASSWORD_SALT = 'ruico-bid-parser'
DEFAULT_ADMIN_USERNAME = 'ruico'
DEFAULT_ADMIN_PASSWORD = 'Ruico668@'
AUTH_COOKIE_NAME = 'bid_parser_token'

for folder in (DATA_DIR, TMP_DIR, PROJECTS_DIR, MARKET_SKILL_DIR, FEISHU_DIR, FEISHU_SESSIONS_DIR, FEISHU_EVENTS_DIR, FEISHU_AGENT_LOGS_DIR):
    folder.mkdir(parents=True, exist_ok=True)

app = FastAPI(title='标书解析与登记 Demo')
app.mount('/static', StaticFiles(directory=STATIC_DIR), name='static')


@app.on_event('startup')
def startup_feishu_ws() -> None:
    start_feishu_ws_client()

EXTRACTION_FIELDS = [
    {'key': 'open_time', 'label': '开标时间', 'cell': 'B2'},
    {'key': 'submission_method', 'label': '投标文件递交方式', 'cell': 'B3'},
    {'key': 'deposit_amount', 'label': '投标保证金', 'cell': 'B4'},
    {'key': 'deposit_method', 'label': '保证金缴纳方式', 'cell': 'E4'},
    {'key': 'deposit_notes', 'label': '保证金备注', 'cell': 'G4'},
    {'key': 'performance_bond', 'label': '履约保证金', 'cell': 'B5'},
    {'key': 'agency_fee', 'label': '招标代理服务费', 'cell': 'F5'},
    {'key': 'service_period', 'label': '服务期限', 'cell': 'B6'},
    {'key': 'service_scope_note', 'label': '服务范围要求', 'cell': 'B7'},
    {'key': 'qualification_requirements', 'label': '投标人资格要求', 'cell': 'B8'},
    {'key': 'package_rule', 'label': '合同包/标段规则', 'cell': 'B9'},
    {'key': 'vessel_requirements', 'label': '船舶要求', 'cell': 'B10'},
    {'key': 'evaluation_method', 'label': '评标方法', 'cell': 'B11'},
    {'key': 'technical_business', 'label': '技术/商务要求', 'cell': 'B12'},
    {'key': 'quotation', 'label': '报价要求', 'cell': 'B13'},
]

REGISTER_COLUMNS = [
    {'key': 'receive_date', 'label': '接收日期'},
    {'key': 'project_name', 'label': '项目名称'},
    {'key': 'is_awarded', 'label': '是否中标'},
    {'key': 'bid_no', 'label': '招标编号'},
    {'key': 'tenderer', 'label': '招标人'},
    {'key': 'bid_deadline', 'label': '投标截止日期'},
    {'key': 'deposit_amount', 'label': '保证金金额'},
    {'key': 'contract_note', 'label': '保证金退款备注（合同期限）'},
    {'key': 'deposit_return_status', 'label': '保证金退还情况'},
    {'key': 'payment_method', 'label': '缴纳方式'},
    {'key': 'acquisition_method', 'label': '招标文件获取途径'},
    {'key': 'submission_method', 'label': '投标文件递交方式'},
    {'key': 'remark', 'label': '备注'},
]

DEFAULT_CONFIG = {
    'base_url': '',
    'api_key': '',
    'model': '',
    'temperature': 0.1,
    'feishu_enabled': False,
    'feishu_receive_mode': 'ws',
    'feishu_app_id': '',
    'feishu_app_secret': '',
    'feishu_verification_token': '',
    'feishu_encrypt_key': '',
    'feishu_allowed_open_ids': '',
    'feishu_allowed_chat_ids': '',
}

FEISHU_TOKEN_CACHE: dict[str, Any] = {'token': '', 'expires_at': 0.0}
FEISHU_WS_THREAD: threading.Thread | None = None
FEISHU_WS_STARTED_FOR = ''
FEISHU_TOOL_NAMES = {
    'help',
    'search_project',
    'parse_bid_file',
    'extract_cargo',
    'extract_newbuilding',
    'market_report_cargo',
    'market_report_newbuilding',
    'confirm_save_market_record',
    'cancel_pending',
}

FEISHU_SKILLS: dict[str, dict[str, Any]] = {
    'help': {
        'title': '系统总览',
        'description': '告诉用户系统能做什么，并给出下一步输入示例。',
        'keywords': ('帮助', '你好', '你能做什么', '菜单', '功能'),
        'example': '你好',
    },
    'bid_parse': {
        'title': '标书解析',
        'description': '上传标书文件后，提取核心字段、风险点和原文依据。',
        'keywords': ('解析标书', '标书解析', '招标文件', '投标文件'),
        'example': '解析标书',
        'needs_file': True,
        'confirmation': False,
    },
    'project_search': {
        'title': '项目查询',
        'description': '按项目名、招标编号、年份查询历史项目台账。',
        'keywords': ('查询', '查询项目', '历史', '台账', '项目'),
        'example': '查询 福海创',
        'confirmation': False,
    },
    'cargo_opportunity': {
        'title': '商机收集',
        'description': '识别货盘信息并沉淀到市场情报台账。',
        'keywords': ('商机', '货盘', '货源', '成交价', '运价'),
        'example': '4000吨甲苯，张家港～东莞，要求月内装出',
        'kind': 'cargo',
        'confirmation': True,
    },
    'newbuilding_info': {
        'title': '新造船收集',
        'description': '识别新造船信息并沉淀到市场情报台账。',
        'keywords': ('新造船', '造船', '船厂', '交付', '出厂'),
        'example': '某油化船，船厂A，船东B，DWT 50000，预计2027年出厂',
        'kind': 'newbuilding',
        'confirmation': True,
    },
    'cargo_report': {
        'title': '商机市场分析',
        'description': '按用户选择的时间范围生成商机市场分析报告。',
        'keywords': ('商机分析', '商机报告', '市场分析', '周报', '月报', '报告'),
        'example': '生成商机报告，本月',
        'kind': 'cargo',
        'needs_period': True,
        'confirmation': False,
    },
    'newbuilding_report': {
        'title': '新造船市场分析',
        'description': '按用户选择的时间范围生成新造船市场分析报告。',
        'keywords': ('新造船分析', '新造船报告', '市场分析', '周报', '月报', '报告'),
        'example': '生成新造船报告，本月',
        'kind': 'newbuilding',
        'needs_period': True,
        'confirmation': False,
    },
}

FEISHU_SKILL_ORDER = tuple(FEISHU_SKILLS.keys())

STATUS_LABELS = {'ok', '已明确满足', '需人工确认', '疑似风险/否决项'}
ELLIPSIS_MARKERS = ('...', '…', '⋯', '。。。')
BID_STATUS_OPTIONS = ('待跟进', '准备投标', '已投标', '放弃', '未投标')
AWARD_STATUS_OPTIONS = ('未知', '待定', '已中标', '未中标')
TIMELINE_TYPE_OPTIONS = ('parse', 'quote', 'award', 'note')

PARSER_SYSTEM_PROMPT = """
你是航运招投标标书解析助手，只能输出一个 JSON 对象，不能输出 Markdown 或额外解释。

请根据用户上传的招标文件正文，提取标书核心信息，并结合以下公司流程输出分析：
1. 标书解析：投标截止时间、报价要求、资质文件、技术方案、商务条款、评标标准。
2. 业务匹配：只输出“需人工确认项”，不要擅自判断公司运力或运量一定满足。
3. 安全匹配：只输出“需人工确认项”和明显否决风险提示。
4. 询比价场景提示：
   - 老客户 + 已承运货种 + 已靠泊港口：可直接询比价。
   - 新货种 / 新港口：先做适装适靠性评估。
   - 首次合作新客户：先做客户资质及评级审核。
   - 高风险运区：提示上报航运部经理并组织海务、法务、安全联合评估。
5. 航运部客户经理应于接收当日完成《招标文件接收登记表》登记。
6. 如果文档包含多个合同包或标段，请在 register_rows 中按“每个标段一行”返回；无法明确拆分时返回一行。
7. 不确定的字段留空，不要编造。
8. extraction_fields 中的 source_excerpt 必须是招标文件中的原文连续片段，不得使用省略号，不得改写。
9. source_excerpt 的原则是“宁长勿短、宁全勿漏”：
   - 不要只摘一句标题或半句结论。
   - 对于保证金、履约保证金、招标代理费、合同包规则、资格要求、船舶要求、评标办法、技术商务要求、报价要求，必须尽量摘取完整条款块。
   - 如果原文是多行、多项列表、分合同包金额、特别说明、备注、补充条件，必须把相关原文整体保留到 source_excerpt 中，保留换行和分项，不要压缩成一句短语。
   - 如果某字段的依据分散在相邻几段正文中，应合并摘取这些连续原文，优先保留完整条件、金额、时间、适用范围、例外说明。
   - 尤其像“合同包1-8分别对应不同保证金/履约保证金”“资格条件第1-7条”“评标规则9.1/9.2/9.3”“船舶要求及特别说明”这类内容，必须保留完整原文区块，不能只截最前面一行。
10. extraction_fields.value 可以是提炼后的摘要，但 source_excerpt 必须尽可能接近人工摘取标准，详细保留原文上下文，便于直接粘贴到“原文依据”模板页。
11. 如果一个文件中存在多个标的、多个合同包、多个标段、多个航线或多个独立报价单元，必须进行统计并逐项拆分：
   - register_rows 中每个标的/标段/合同包至少返回一行。
   - analysis.summary 中明确说明总共识别到多少个标的/标段/合同包。
   - package_rule 中尽量摘取原文中关于“可兼投兼中、分包、分标段、分航线、分别报价”的完整说明。

JSON 顶层字段必须完整包含：
{
  "document_summary": {
    "project_name": "",
    "bid_no": "",
    "tenderer": "",
    "bid_deadline": "",
    "open_time": "",
    "submission_method": "",
    "deposit_amount": "",
    "service_period": "",
    "qualification_requirements": "",
    "vessel_requirements": "",
    "technical_business": "",
    "quotation": "",
    "evaluation_method": ""
  },
  "extraction_fields": {
    "open_time": {"value": "", "source_excerpt": ""},
    "submission_method": {"value": "", "source_excerpt": ""},
    "deposit_amount": {"value": "", "source_excerpt": ""},
    "deposit_method": {"value": "", "source_excerpt": ""},
    "deposit_notes": {"value": "", "source_excerpt": ""},
    "performance_bond": {"value": "", "source_excerpt": ""},
    "agency_fee": {"value": "", "source_excerpt": ""},
    "service_period": {"value": "", "source_excerpt": ""},
    "service_scope_note": {"value": "", "source_excerpt": ""},
    "qualification_requirements": {"value": "", "source_excerpt": ""},
    "package_rule": {"value": "", "source_excerpt": ""},
    "vessel_requirements": {"value": "", "source_excerpt": ""},
    "evaluation_method": {"value": "", "source_excerpt": ""},
    "technical_business": {"value": "", "source_excerpt": ""},
    "quotation": {"value": "", "source_excerpt": ""}
  },
  "register_rows": [
    {
      "receive_date": "",
      "project_name": "",
      "is_awarded": "",
      "bid_no": "",
      "tenderer": "",
      "bid_deadline": "",
      "deposit_amount": "",
      "contract_note": "",
      "deposit_return_status": "",
      "payment_method": "",
      "acquisition_method": "",
      "submission_method": "",
      "remark": ""
    }
  ],
  "analysis": {
    "summary": "",
    "qualification_files": [],
    "business_points": [],
    "scoring_points": [],
    "risks": []
  },
  "match_review": [
    {
      "area": "",
      "item": "",
      "status": "已明确满足|需人工确认|疑似风险/否决项",
      "reason": "",
      "action": ""
    }
  ]
}
""".strip()

CHAT_SYSTEM_PROMPT = """
你是标书问答助手，只能基于本次标书解析结果和摘取的原文片段回答。
如果问题没有明确依据，直接回答：
未在本次标书中找到明确依据，请人工复核原文。
回答尽量简洁，优先引用结构化结果。
""".strip()

JSON_REPAIR_SYSTEM_PROMPT = """
你是 JSON 修复助手。
你的唯一任务是把用户给出的内容修复成一个严格有效的 JSON 对象。
要求：
1. 只能输出 JSON 对象本身。
2. 不要输出 Markdown，不要输出解释，不要输出多余文字。
3. 保留原始语义，不要擅自新增不存在的业务结论。
4. 如果原文里有换行、引号、列表，按 JSON 字符串规范正确转义。
5. 如果存在尾逗号、智能引号、格式错误、半截内容，尽最大努力修复为合法 JSON。
""".strip()

SESSIONS: dict[str, dict[str, Any]] = {}
AUTH_SESSIONS: dict[str, dict[str, Any]] = {}


def hash_password(password: str) -> str:
    # ponytail: 先用 salted sha256 落地单机 demo；要上公网时再换成 bcrypt/argon2。
    return hashlib.sha256(f'{PASSWORD_SALT}:{password}'.encode('utf-8')).hexdigest()


def default_admin_user() -> dict[str, Any]:
    return {
        'username': DEFAULT_ADMIN_USERNAME,
        'password_hash': hash_password(DEFAULT_ADMIN_PASSWORD),
        'display_name': '系统管理员',
        'role': 'admin',
        'created_at': datetime.now().isoformat(timespec='seconds'),
    }


def load_users() -> dict[str, Any]:
    if not USERS_PATH.exists():
        payload = {'users': [default_admin_user()]}
        USERS_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        return payload
    try:
        payload = json.loads(USERS_PATH.read_text(encoding='utf-8'))
    except json.JSONDecodeError:
        payload = {'users': [default_admin_user()]}
        USERS_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        return payload
    users = payload.get('users') if isinstance(payload, dict) else None
    if not isinstance(users, list):
        payload = {'users': [default_admin_user()]}
        USERS_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        return payload
    if not any((item or {}).get('username') == DEFAULT_ADMIN_USERNAME for item in users if isinstance(item, dict)):
        users.insert(0, default_admin_user())
        USERS_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
    return payload


def save_users(payload: dict[str, Any]) -> dict[str, Any]:
    USERS_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
    return payload


def find_user(username: str) -> dict[str, Any] | None:
    username = compact_text(username).lower()
    for user in load_users().get('users', []):
        if compact_text((user or {}).get('username', '')).lower() == username:
            return user
    return None


def sanitize_user(user: dict[str, Any]) -> dict[str, Any]:
    return {
        'username': compact_text(user.get('username', '')),
        'display_name': compact_text(user.get('display_name', '')),
        'role': compact_text(user.get('role', 'user')) or 'user',
        'created_at': compact_text(user.get('created_at', '')),
    }


def cleanup_auth_sessions() -> None:
    now = datetime.now()
    expired = [
        token
        for token, payload in AUTH_SESSIONS.items()
        if now - payload['created_at'] > AUTH_TTL
    ]
    for token in expired:
        AUTH_SESSIONS.pop(token, None)


def create_auth_session(user: dict[str, Any]) -> dict[str, Any]:
    cleanup_auth_sessions()
    token = uuid.uuid4().hex
    payload = {
        'token': token,
        'user': sanitize_user(user),
        'created_at': datetime.now(),
    }
    AUTH_SESSIONS[token] = payload
    return {'token': token, 'user': payload['user']}


def extract_auth_token(authorization: str | None = None, auth_cookie: str | None = None) -> str:
    if authorization and authorization.startswith('Bearer '):
        return authorization.split(' ', 1)[1].strip()
    return compact_text(auth_cookie)


def require_auth(
    authorization: str | None = Header(default=None),
    auth_cookie: str | None = Cookie(default=None, alias=AUTH_COOKIE_NAME),
) -> dict[str, Any]:
    cleanup_auth_sessions()
    token = extract_auth_token(authorization, auth_cookie)
    if not token:
        raise HTTPException(status_code=401, detail='请先登录。')
    session = AUTH_SESSIONS.get(token)
    if not session:
        raise HTTPException(status_code=401, detail='登录已失效，请重新登录。')
    return session['user']


def require_admin(
    authorization: str | None = Header(default=None),
    auth_cookie: str | None = Cookie(default=None, alias=AUTH_COOKIE_NAME),
) -> dict[str, Any]:
    user = require_auth(authorization, auth_cookie)
    if user.get('role') != 'admin':
        raise HTTPException(status_code=403, detail='只有管理员可以执行该操作。')
    return user


def resolve_auth_user(authorization: str | None = None, auth_cookie: str | None = None) -> dict[str, Any] | None:
    cleanup_auth_sessions()
    token = extract_auth_token(authorization, auth_cookie)
    if not token:
        return None
    session = AUTH_SESSIONS.get(token)
    if not session:
        return None
    return session['user']


class ConfigPayload(BaseModel):
    base_url: str = ''
    api_key: str = ''
    model: str = ''
    temperature: float = 0.1
    feishu_enabled: bool = False
    feishu_receive_mode: str = 'ws'
    feishu_app_id: str = ''
    feishu_app_secret: str = ''
    feishu_verification_token: str = ''
    feishu_encrypt_key: str = ''
    feishu_allowed_open_ids: str = ''
    feishu_allowed_chat_ids: str = ''


class FeishuTestPayload(BaseModel):
    receive_id_type: str = 'open_id'
    receive_id: str = ''
    text: str = '标书解析机器人测试消息'


class ChatPayload(BaseModel):
    session_id: str
    question: str = Field(min_length=1)


class ExportExtractionPayload(BaseModel):
    result: dict[str, Any] = Field(default_factory=dict)
    extraction_fields: dict[str, dict[str, str]] = Field(default_factory=dict)


class ExportRegisterPayload(BaseModel):
    rows: list[dict[str, Any]]
    sheet_name: str = ''


class ExportOverviewPayload(BaseModel):
    title: str = ''
    result: dict[str, Any]
    follow_up: dict[str, Any] = Field(default_factory=dict)
    our_quotes: list[dict[str, Any]] = Field(default_factory=list)
    competitor_quotes: list[dict[str, Any]] = Field(default_factory=list)
    timeline: list[dict[str, Any]] = Field(default_factory=list)


class ExportQuotePayload(BaseModel):
    title: str = ''
    follow_up: dict[str, Any] = Field(default_factory=dict)
    rows: list[dict[str, Any]] = Field(default_factory=list)


class ExportLedgerPayload(BaseModel):
    q: str = ''
    year: str = ''
    award_status: str = ''
    bid_status: str = ''


class MarketExtractPayload(BaseModel):
    kind: str
    text: str = Field(min_length=1)


class MarketRecordPayload(BaseModel):
    kind: str
    record: dict[str, Any] = Field(default_factory=dict)


class MarketExportPayload(BaseModel):
    kind: str
    q: str = ''
    segment: str = ''
    board_type: str = ''
    status: str = ''
    stage: str = ''


class MarketReportPayload(BaseModel):
    kind: str
    period: str = 'custom'
    start_date: str = ''
    end_date: str = ''
    custom_prompt: str = ''
    filters: dict[str, Any] = Field(default_factory=dict)


class MarketReportExportPayload(BaseModel):
    kind: str
    report: dict[str, Any] = Field(default_factory=dict)
    format: str = 'docx'


class ProjectPayload(BaseModel):
    title: str = ''
    source_file_name: str = ''
    source_channel: str = ''
    source_open_id: str = ''
    source_chat_id: str = ''
    source_message_id: str = ''
    source_text: str = ''
    confirmed_at: str = ''
    register_mode: str = 'packages'
    sheet_name: str = ''
    result: dict[str, Any] = Field(default_factory=dict)
    follow_up: dict[str, Any] = Field(default_factory=dict)
    our_quotes: list[dict[str, Any]] = Field(default_factory=list)
    competitor_quotes: list[dict[str, Any]] = Field(default_factory=list)
    timeline: list[dict[str, Any]] = Field(default_factory=list)


class LoginPayload(BaseModel):
    username: str = Field(min_length=1)
    password: str = Field(min_length=1)


class CreateUserPayload(BaseModel):
    username: str = Field(min_length=1)
    password: str = Field(min_length=1)
    display_name: str = ''
    role: str = 'user'


def cleanup_sessions() -> None:
    now = datetime.now()
    expired = [
        session_id
        for session_id, payload in SESSIONS.items()
        if now - payload['created_at'] > SESSION_TTL
    ]
    for session_id in expired:
        SESSIONS.pop(session_id, None)


def load_config() -> dict[str, Any]:
    if not CONFIG_PATH.exists():
        return DEFAULT_CONFIG.copy()
    try:
        raw = json.loads(CONFIG_PATH.read_text(encoding='utf-8'))
    except json.JSONDecodeError:
        return DEFAULT_CONFIG.copy()
    return {**DEFAULT_CONFIG, **raw}


def save_config(payload: dict[str, Any]) -> dict[str, Any]:
    current = load_config()
    api_key = payload.get('api_key', '').strip() or current.get('api_key', '')
    feishu_app_secret = payload.get('feishu_app_secret', '').strip() or current.get('feishu_app_secret', '')
    feishu_verification_token = payload.get('feishu_verification_token', '').strip() or current.get('feishu_verification_token', '')
    feishu_encrypt_key = payload.get('feishu_encrypt_key', '').strip() or current.get('feishu_encrypt_key', '')
    cleaned = {
        'base_url': payload.get('base_url', '').strip(),
        'api_key': api_key,
        'model': payload.get('model', '').strip(),
        'temperature': float(payload.get('temperature', 0.1) or 0.1),
        'feishu_enabled': bool(payload.get('feishu_enabled')),
        'feishu_receive_mode': payload.get('feishu_receive_mode', 'ws').strip() if payload.get('feishu_receive_mode') in {'ws', 'http'} else 'ws',
        'feishu_app_id': payload.get('feishu_app_id', '').strip(),
        'feishu_app_secret': feishu_app_secret,
        'feishu_verification_token': feishu_verification_token,
        'feishu_encrypt_key': feishu_encrypt_key,
        'feishu_allowed_open_ids': payload.get('feishu_allowed_open_ids', '').strip(),
        'feishu_allowed_chat_ids': payload.get('feishu_allowed_chat_ids', '').strip(),
    }
    CONFIG_PATH.write_text(
        json.dumps(cleaned, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )
    return cleaned


def mask_config(config: dict[str, Any]) -> dict[str, Any]:
    masked = config.copy()
    masked['has_api_key'] = bool(masked.get('api_key'))
    if masked.get('api_key'):
        masked['api_key'] = '*' * 8
    for key in ('feishu_app_secret', 'feishu_verification_token', 'feishu_encrypt_key'):
        masked[f'has_{key}'] = bool(masked.get(key))
        if masked.get(key):
            masked[key] = '*' * 8
    return masked


def require_config() -> dict[str, Any]:
    config = load_config()
    if not config['base_url'] or not config['api_key'] or not config['model']:
        raise HTTPException(
            status_code=400,
            detail='请先在左侧 AI 配置中保存 base URL、API Key 和模型名称。',
        )
    return config


def normalize_base_url(base_url: str) -> str:
    value = base_url.rstrip('/')
    if value.endswith('/chat/completions'):
        return value
    if value.endswith('/v1'):
        return f'{value}/chat/completions'
    return f'{value}/v1/chat/completions'


def parse_json_text(content: str) -> dict[str, Any]:
    def strip_code_fence(text: str) -> str:
        text = text.strip()
        if text.startswith('```'):
            text = re.sub(r'^```(?:json)?\s*', '', text)
            text = re.sub(r'\s*```$', '', text)
        return text.strip()

    def extract_balanced_object(text: str) -> str:
        start = text.find('{')
        if start < 0:
            return ''
        depth = 0
        in_string = False
        escape = False
        for index in range(start, len(text)):
            char = text[index]
            if in_string:
                if escape:
                    escape = False
                elif char == '\\':
                    escape = True
                elif char == '"':
                    in_string = False
                continue
            if char == '"':
                in_string = True
            elif char == '{':
                depth += 1
            elif char == '}':
                depth -= 1
                if depth == 0:
                    return text[start : index + 1]
        return text[start:]

    def sanitize_json_candidate(text: str) -> str:
        text = text.replace('\ufeff', '')
        text = (
            text.replace('“', '"')
            .replace('”', '"')
            .replace('‘', '"')
            .replace('’', '"')
        )
        text = re.sub(r',(\s*[}\]])', r'\1', text)
        buffer: list[str] = []
        in_string = False
        escape = False
        for char in text:
            if in_string:
                if escape:
                    buffer.append(char)
                    escape = False
                    continue
                if char == '\\':
                    buffer.append(char)
                    escape = True
                    continue
                if char == '"':
                    buffer.append(char)
                    in_string = False
                    continue
                if char == '\n':
                    buffer.append('\\n')
                    continue
                if char == '\r':
                    buffer.append('\\r')
                    continue
                if char == '\t':
                    buffer.append('\\t')
                    continue
            else:
                if char == '"':
                    in_string = True
            buffer.append(char)
        return ''.join(buffer).strip()

    text = strip_code_fence(content)
    candidates = [text]
    balanced = extract_balanced_object(text)
    if balanced:
        candidates.append(balanced)
    candidates.extend(sanitize_json_candidate(item) for item in list(candidates))

    seen: set[str] = set()
    last_error: Exception | None = None
    for candidate in candidates:
        candidate = candidate.strip()
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError as exc:
            last_error = exc
            continue

    detail = 'AI 返回的 JSON 无法解析。'
    if last_error:
        detail = f'AI 返回的 JSON 无法解析：{last_error.msg}'
    raise HTTPException(status_code=502, detail=detail)


def extract_ai_response_content(data: Any) -> str:
    def text_from_content(value: Any) -> str:
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, list):
            parts = [text_from_content(item) for item in value]
            return '\n'.join(part for part in parts if part).strip()
        if not isinstance(value, dict):
            return ''

        for key in ('text', 'output_text', 'content'):
            nested = text_from_content(value.get(key))
            if nested:
                return nested
        if isinstance(value.get('message'), dict):
            nested = text_from_content(value['message'].get('content'))
            if nested:
                return nested
        return ''

    if isinstance(data, str):
        return data.strip()
    if not isinstance(data, dict):
        return ''

    choices = data.get('choices')
    if isinstance(choices, list) and choices:
        for choice in choices:
            if not isinstance(choice, dict):
                continue
            message = choice.get('message') or choice.get('delta')
            if isinstance(message, dict):
                content = text_from_content(message.get('content'))
                if content:
                    return content
                function_call = message.get('function_call')
                if isinstance(function_call, dict):
                    content = text_from_content(function_call.get('arguments'))
                    if content:
                        return content
                tool_calls = message.get('tool_calls')
                if isinstance(tool_calls, list):
                    for tool_call in tool_calls:
                        if not isinstance(tool_call, dict):
                            continue
                        function = tool_call.get('function')
                        if isinstance(function, dict):
                            content = text_from_content(function.get('arguments'))
                            if content:
                                return content
            content = text_from_content(choice.get('content') or choice.get('text'))
            if content:
                return content

    output_text = text_from_content(data.get('output_text'))
    if output_text:
        return output_text

    output = data.get('output')
    if isinstance(output, list):
        output_parts: list[str] = []
        for item in output:
            if isinstance(item, dict):
                content = text_from_content(item.get('content') or item.get('text'))
            else:
                content = text_from_content(item)
            if content:
                output_parts.append(content)
        if output_parts:
            return '\n'.join(output_parts).strip()

    for key in ('result', 'response', 'answer', 'text', 'content'):
        content = text_from_content(data.get(key))
        if content:
            return content

    message = data.get('message')
    if isinstance(message, dict):
        content = text_from_content(message.get('content'))
        if content:
            return content
    elif isinstance(message, str) and message.strip().startswith('{'):
        try:
            nested = json.loads(message)
            content = extract_ai_response_content(nested)
            if content:
                return content
        except json.JSONDecodeError:
            pass

    data_node = data.get('data')
    if isinstance(data_node, dict):
        content = extract_ai_response_content(data_node)
        if content:
            return content

    return ''


def call_chat_completion(
    config: dict[str, Any],
    *,
    system_prompt: str,
    user_prompt: str,
    temperature: float | None = None,
) -> str:
    payload = {
        'model': config['model'],
        'temperature': config.get('temperature', 0.1) if temperature is None else temperature,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ],
    }
    headers = {
        'Authorization': f"Bearer {config['api_key']}",
        'Content-Type': 'application/json',
    }
    url = normalize_base_url(config['base_url'])
    timeout = httpx.Timeout(connect=30.0, read=300.0, write=60.0, pool=30.0)
    response = None
    try:
        response = httpx.post(url, json=payload, headers=headers, timeout=timeout)
    except httpx.ConnectTimeout as exc:
        # ponytail: 仅对连接阶段重试一次；读超时不重试，避免模型已成功执行却重复扣费。
        try:
            response = httpx.post(url, json=payload, headers=headers, timeout=timeout)
        except httpx.ConnectTimeout as retry_exc:
            raise HTTPException(
                status_code=504,
                detail=f'AI 接口连接超时：{type(retry_exc).__name__}: {retry_exc}',
            ) from retry_exc
    except (httpx.ReadTimeout, httpx.WriteTimeout) as exc:
        raise HTTPException(
            status_code=504,
            detail=(
                f'AI 接口响应超时：{type(exc).__name__}: {exc}。'
                '请求可能已到达模型侧并正在处理，请避免立即重复提交，可稍后重试或缩短解析内容。'
            ),
        ) from exc
    except httpx.TimeoutException as exc:
        raise HTTPException(
            status_code=504,
            detail=f'AI 接口超时：{type(exc).__name__}: {exc}',
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(
            status_code=502,
            detail=f'AI 接口调用失败：{type(exc).__name__}: {exc}',
        ) from exc
    if response is None:
        raise HTTPException(
            status_code=504,
            detail='AI 接口超时，请稍后重试。',
        )
    if response.status_code == 401:
        raise HTTPException(status_code=401, detail='AI 接口鉴权失败，请检查 API Key。')
    if response.status_code >= 400:
        raise HTTPException(
            status_code=502,
            detail=f"AI 接口返回错误：{response.status_code} {response.text[:200]}",
        )
    try:
        data = response.json()
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f'AI 接口响应不是有效 JSON：{response.text[:200]}',
        ) from exc
    if isinstance(data, dict) and data.get('error'):
        error = data['error']
        if isinstance(error, dict):
            message = error.get('message') or error.get('msg') or json.dumps(error, ensure_ascii=False)
        else:
            message = str(error)
        raise HTTPException(status_code=502, detail=f'AI 接口返回错误：{message[:200]}')
    content = extract_ai_response_content(data)
    if content:
        return content
    preview = json.dumps(data, ensure_ascii=False)[:300] if isinstance(data, (dict, list)) else str(data)[:300]
    raise HTTPException(
        status_code=502,
        detail=f'AI 接口响应中没有找到可解析的文本内容：{preview}',
    )


def clean_text(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    parts = [(page.extract_text() or '').strip() for page in reader.pages]
    return clean_text('\n\n'.join(part for part in parts if part))


def extract_docx_text(file_path: Path) -> str:
    try:
        with ZipFile(file_path) as archive:
            try:
                xml_bytes = archive.read('word/document.xml')
            except KeyError as exc:
                raise HTTPException(status_code=400, detail='DOCX 结构无效，无法读取正文。') from exc
    except BadZipFile as exc:
        raise HTTPException(
            status_code=400,
            detail='上传的 DOCX 文件格式无效，当前文件不是可解析的 Office 文档，请重新导出为标准 DOCX 后再上传。',
        ) from exc
    root = ET.fromstring(xml_bytes)
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts: list[str] = []
    for paragraph in root.findall('.//w:p', namespace):
        texts = [node.text for node in paragraph.findall('.//w:t', namespace) if node.text]
        line = ''.join(texts).strip()
        if line:
            parts.append(line)
    return clean_text('\n'.join(parts))


def extract_xlsx_text(file_path: Path) -> str:
    try:
        workbook = load_workbook(file_path, data_only=True)
    except (BadZipFile, InvalidFileException, OSError) as exc:
        raise HTTPException(
            status_code=400,
            detail='上传的 Excel 文件格式无效，当前文件不是可解析的标准 Excel 文档，请重新保存为 xlsx 后再上传。',
        ) from exc
    parts: list[str] = []
    for sheet in workbook.worksheets:
        sheet_lines: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [compact_text(cell) for cell in row if compact_text(cell)]
            if cells:
                sheet_lines.append(' | '.join(cells))
        if sheet_lines:
            parts.append(f'工作表：{sheet.title}')
            parts.extend(sheet_lines)
    workbook.close()
    return clean_text('\n'.join(parts))


def extract_with_office_com(file_path: Path, kind: str) -> str:
    escaped = str(file_path).replace("'", "''")
    if kind == 'word':
        script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{escaped}'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open($path, $false, $true)
  $text = $doc.Content.Text
  [Console]::OutputEncoding = [System.Text.Encoding]::UTF8
  Write-Output $text
}} finally {{
  if ($doc) {{ $doc.Close() }}
  if ($word) {{ $word.Quit() }}
}}
""".strip()
    else:
        script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{escaped}'
$excel = $null
$workbook = $null
try {{
  $excel = New-Object -ComObject Excel.Application
  $excel.Visible = $false
  $excel.DisplayAlerts = $false
  $workbook = $excel.Workbooks.Open($path, 0, $true)
  $lines = New-Object System.Collections.Generic.List[string]
  foreach ($sheet in $workbook.Worksheets) {{
    $lines.Add("工作表：$($sheet.Name)")
    $range = $sheet.UsedRange
    $rowCount = $range.Rows.Count
    $colCount = $range.Columns.Count
    for ($r = 1; $r -le $rowCount; $r++) {{
      $cells = New-Object System.Collections.Generic.List[string]
      for ($c = 1; $c -le $colCount; $c++) {{
        $value = $range.Item($r, $c).Text
        if ($value) {{ $cells.Add($value.Trim()) }}
      }}
      if ($cells.Count -gt 0) {{ $lines.Add(($cells -join ' | ')) }}
    }}
  }}
  [Console]::OutputEncoding = [System.Text.Encoding]::UTF8
  Write-Output ($lines -join [Environment]::NewLine)
}} finally {{
  if ($workbook) {{ $workbook.Close($false) }}
  if ($excel) {{ $excel.Quit() }}
}}
""".strip()
    try:
        completed = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-Command', script],
            capture_output=True,
            text=True,
            encoding='utf-8',
            timeout=120,
            check=False,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail='当前环境无法调用 Office 解析该文件，请优先上传 PDF、DOCX 或 XLSX。',
        ) from exc
    if completed.returncode != 0:
        raise HTTPException(
            status_code=400,
            detail='当前环境无法解析该 Office 文件。若本机未安装 Word/Excel，请先另存为 DOCX/XLSX/PDF 后再上传。',
        )
    return clean_text(completed.stdout)


def extract_doc_text(file_path: Path) -> str:
    return extract_with_office_com(file_path, 'word')


def extract_xls_text(file_path: Path) -> str:
    return extract_with_office_com(file_path, 'excel')


def extract_document_text(file_path: Path, suffix: str) -> str:
    if suffix == '.pdf':
        return extract_pdf_text(file_path)
    if suffix == '.docx':
        # ponytail: 只读 document.xml 正文，不处理页眉页脚和图片；要更完整再接专门的 docx 解析器。
        return extract_docx_text(file_path)
    if suffix == '.doc':
        # ponytail: 旧版 doc 直接走本机 Word COM；无 Office 时引导另存为 docx/pdf。
        return extract_doc_text(file_path)
    if suffix == '.xlsx':
        return extract_xlsx_text(file_path)
    if suffix == '.xls':
        # ponytail: 旧版 xls 直接走本机 Excel COM；无 Office 时引导另存为 xlsx/pdf。
        return extract_xls_text(file_path)
    raise HTTPException(status_code=400, detail='当前只支持 PDF、DOC、DOCX、XLS、XLSX。')


def compact_text(value: Any) -> str:
    text = str(value or '').replace('\u3000', ' ')
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', ' ', text)
    return text.strip(' ;；，,')


def compact_paragraph(value: Any) -> str:
    text = str(value or '').replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def normalize_awarded(value: Any) -> str:
    text = compact_text(value)
    lowered = text.lower()
    if lowered in {'√', '是', 'yes', 'y', 'true', '已中标'}:
        return '√'
    if lowered in {'×', '否', 'no', 'n', 'false', '未中标'}:
        return ''
    return text


def normalize_summary(raw: Any) -> dict[str, str]:
    source = raw if isinstance(raw, dict) else {}
    return {
        'project_name': compact_text(source.get('project_name')),
        'bid_no': compact_text(source.get('bid_no')),
        'tenderer': compact_text(source.get('tenderer')),
        'bid_deadline': compact_text(source.get('bid_deadline')),
        'open_time': compact_text(source.get('open_time')),
        'submission_method': compact_text(source.get('submission_method')),
        'deposit_amount': compact_text(source.get('deposit_amount')),
        'service_period': compact_text(source.get('service_period')),
        'qualification_requirements': compact_paragraph(source.get('qualification_requirements')),
        'vessel_requirements': compact_paragraph(source.get('vessel_requirements')),
        'technical_business': compact_paragraph(source.get('technical_business')),
        'quotation': compact_paragraph(source.get('quotation')),
        'evaluation_method': compact_paragraph(source.get('evaluation_method')),
    }


def normalize_extraction_fields(raw: Any) -> dict[str, dict[str, str]]:
    source = raw if isinstance(raw, dict) else {}
    normalized: dict[str, dict[str, str]] = {}
    for field in EXTRACTION_FIELDS:
        item = source.get(field['key'], {})
        if isinstance(item, dict):
            value = compact_paragraph(item.get('value', ''))
            excerpt = compact_paragraph(item.get('source_excerpt', ''))
        else:
            value = compact_paragraph(item)
            excerpt = ''
        normalized[field['key']] = {'value': value, 'source_excerpt': excerpt}
    return normalized


def refine_register_row(row: dict[str, Any], summary: dict[str, str]) -> dict[str, str]:
    return {
        'receive_date': compact_text(row.get('receive_date') or datetime.now().strftime('%Y-%m-%d')),
        'project_name': compact_text(row.get('project_name') or summary['project_name']),
        'is_awarded': normalize_awarded(row.get('is_awarded')),
        'bid_no': compact_text(row.get('bid_no') or summary['bid_no']),
        'tenderer': compact_text(row.get('tenderer') or summary['tenderer']),
        'bid_deadline': compact_text(row.get('bid_deadline') or summary['bid_deadline']),
        'deposit_amount': compact_text(row.get('deposit_amount') or summary['deposit_amount']),
        'contract_note': compact_text(row.get('contract_note') or summary['service_period']),
        'deposit_return_status': compact_text(row.get('deposit_return_status')),
        'payment_method': compact_text(row.get('payment_method')),
        'acquisition_method': compact_text(row.get('acquisition_method')),
        'submission_method': compact_text(row.get('submission_method') or summary['submission_method']),
        'remark': compact_text(row.get('remark')),
    }


def normalize_register_rows(raw: Any, summary: dict[str, str]) -> list[dict[str, str]]:
    rows = raw if isinstance(raw, list) and raw else [{}]
    normalized: list[dict[str, str]] = []
    for row in rows:
        normalized.append(refine_register_row(row if isinstance(row, dict) else {}, summary))
    return normalized


def normalize_analysis(raw: Any) -> dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    return {
        'summary': compact_paragraph(source.get('summary')),
        'qualification_files': [compact_text(item) for item in source.get('qualification_files', []) if compact_text(item)],
        'business_points': [compact_text(item) for item in source.get('business_points', []) if compact_text(item)],
        'scoring_points': [compact_text(item) for item in source.get('scoring_points', []) if compact_text(item)],
        'risks': [compact_text(item) for item in source.get('risks', []) if compact_text(item)],
    }


def normalize_match_review(raw: Any) -> list[dict[str, str]]:
    rows = raw if isinstance(raw, list) else []
    normalized: list[dict[str, str]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        status = compact_text(row.get('status') or '需人工确认')
        if status not in STATUS_LABELS:
            status = '需人工确认'
        normalized.append(
            {
                'area': compact_text(row.get('area')),
                'item': compact_text(row.get('item')),
                'status': status,
                'reason': compact_paragraph(row.get('reason')),
                'action': compact_paragraph(row.get('action')),
            }
        )
    return normalized


def normalize_follow_up(raw: Any, *, sheet_name: str = '') -> dict[str, str]:
    source = raw if isinstance(raw, dict) else {}
    bid_status = compact_text(source.get('bid_status') or '待跟进')
    award_status = compact_text(source.get('award_status') or '未知')
    if bid_status not in BID_STATUS_OPTIONS:
        bid_status = '待跟进'
    if award_status not in AWARD_STATUS_OPTIONS:
        award_status = '未知'
    return {
        'bid_status': bid_status,
        'award_status': award_status,
        'award_date': compact_text(source.get('award_date')),
        'award_company': compact_text(source.get('award_company')),
        'our_award_amount': compact_text(source.get('our_award_amount')),
        'competitor_award_amount': compact_text(source.get('competitor_award_amount')),
        'tracking_note': compact_paragraph(source.get('tracking_note')),
        'information_source': compact_text(source.get('information_source')),
        'register_year': compact_text(source.get('register_year') or sheet_name),
    }


def normalize_quote_row(raw: Any, *, default_company: str = '') -> dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    try:
        round_no = max(1, int(source.get('round_no') or 1))
    except (TypeError, ValueError):
        round_no = 1
    return {
        'id': compact_text(source.get('id') or uuid.uuid4().hex[:12]),
        'package_name': compact_text(source.get('package_name')),
        'round_no': round_no,
        'quote_date': compact_text(source.get('quote_date')),
        'quote_company': compact_text(source.get('quote_company') or default_company),
        'currency': compact_text(source.get('currency') or 'CNY'),
        'tax_mode': compact_text(source.get('tax_mode')),
        'unit_price': compact_text(source.get('unit_price')),
        'total_price': compact_text(source.get('total_price')),
        'ranking': compact_text(source.get('ranking')),
        'is_submitted': bool(source.get('is_submitted')),
        'is_awarded': bool(source.get('is_awarded')),
        'source': compact_text(source.get('source')),
        'remark': compact_paragraph(source.get('remark')),
    }


def normalize_quote_rows(raw: Any, *, default_company: str = '') -> list[dict[str, Any]]:
    rows = raw if isinstance(raw, list) else []
    return [normalize_quote_row(row, default_company=default_company) for row in rows]


def normalize_timeline_rows(raw: Any) -> list[dict[str, str]]:
    rows = raw if isinstance(raw, list) else []
    normalized: list[dict[str, str]] = []
    for row in rows:
        source = row if isinstance(row, dict) else {}
        item_type = compact_text(source.get('type') or 'note')
        if item_type not in TIMELINE_TYPE_OPTIONS:
            item_type = 'note'
        note = compact_paragraph(source.get('note'))
        if not note:
            continue
        normalized.append(
            {
                'id': compact_text(source.get('id') or uuid.uuid4().hex[:12]),
                'date': compact_text(source.get('date') or datetime.now().strftime('%Y-%m-%d')),
                'type': item_type,
                'note': note,
            }
        )
    return normalized


def merge_register_rows(rows: list[dict[str, Any]]) -> dict[str, Any]:
    if not rows:
        return {column['key']: '' for column in REGISTER_COLUMNS}
    first = rows[0].copy()
    if len(rows) == 1:
        return first
    bid_numbers = [row.get('bid_no', '') for row in rows if row.get('bid_no')]
    deposits = [row.get('deposit_amount', '') for row in rows if row.get('deposit_amount')]
    if bid_numbers:
        first['bid_no'] = ' / '.join(dict.fromkeys(bid_numbers))
    if deposits:
        first['deposit_amount'] = ' / '.join(dict.fromkeys(deposits))
    package_note = f'识别到 {len(rows)} 个标段/合同包'
    remark = first.get('remark', '')
    first['remark'] = package_note if not remark else f'{remark}；{package_note}'
    return first


def normalize_result_payload(data: dict[str, Any]) -> dict[str, Any]:
    summary = normalize_summary(data.get('document_summary'))
    register_rows = normalize_register_rows(data.get('register_rows'), summary)
    document_text_full = compact_paragraph(data.get('document_text_full') or data.get('document_text_excerpt') or '')
    package_count = len(register_rows)
    return {
        'document_summary': summary,
        'extraction_fields': normalize_extraction_fields(data.get('extraction_fields')),
        'register_rows': register_rows,
        'document_register_row': merge_register_rows(register_rows),
        'package_count': package_count,
        'has_multiple_packages': package_count > 1,
        'analysis': normalize_analysis(data.get('analysis')),
        'match_review': normalize_match_review(data.get('match_review')),
        'document_text_full': document_text_full,
        'document_text_excerpt': document_text_full[:12000],
    }


def build_auto_timeline(result: dict[str, Any], source_file_name: str = '') -> list[dict[str, str]]:
    note = '完成标书解析并生成摘取结果'
    if source_file_name:
        note = f'完成标书解析：{source_file_name}'
    return [
        {
            'id': uuid.uuid4().hex[:12],
            'date': datetime.now().strftime('%Y-%m-%d'),
            'type': 'parse',
            'note': note,
        }
    ]


def build_parser_prompt(document_text: str) -> str:
    return (
        '以下是招标文件正文，请按约定 JSON 结构解析。\n'
        '注意：本次抽取要求接近人工摘取表标准，source_excerpt 要尽可能详细完整，尤其不要遗漏多标段金额、分项条件、特别说明、备注、适用范围。\n\n'
        f'{document_text[:50000]}'
    )


def build_json_repair_prompt(raw_content: str, error_detail: str) -> str:
    return (
        '下面是上一次模型输出的内容，但它不是有效 JSON。\n'
        f'解析错误：{error_detail}\n\n'
        '请修复它，并只返回严格有效的 JSON 对象：\n\n'
        f'{raw_content[:60000]}'
    )


def build_chat_prompt(session: dict[str, Any], question: str) -> str:
    return json.dumps(
        {
            'document_summary': session['result']['document_summary'],
            'analysis': session['result']['analysis'],
            'match_review': session['result']['match_review'],
            'register_rows': session['result']['register_rows'],
            'document_text_excerpt': session['result']['document_text_excerpt'],
            'question': question,
        },
        ensure_ascii=False,
    )


def create_session_for_result(result: dict[str, Any]) -> str:
    cleanup_sessions()
    session_id = uuid.uuid4().hex
    SESSIONS[session_id] = {'created_at': datetime.now(), 'result': result}
    return session_id


def parse_ai_document(document_text: str, config: dict[str, Any]) -> dict[str, Any]:
    content = call_chat_completion(
        config,
        system_prompt=PARSER_SYSTEM_PROMPT,
        user_prompt=build_parser_prompt(document_text),
        temperature=min(float(config.get('temperature', 0.1)), 0.3),
    )
    raw = None
    current_content = content
    last_error_detail = 'AI 返回的 JSON 无法解析。'
    for repair_round in range(3):
        try:
            raw = parse_json_text(current_content)
            break
        except HTTPException as exc:
            last_error_detail = str(exc.detail)
            if repair_round == 2:
                break
            repair_prompt = build_json_repair_prompt(current_content, last_error_detail)
            try:
                current_content = call_chat_completion(
                    config,
                    system_prompt=JSON_REPAIR_SYSTEM_PROMPT,
                    user_prompt=repair_prompt,
                    temperature=0,
                )
            except HTTPException as repair_exc:
                raise HTTPException(
                    status_code=repair_exc.status_code,
                    detail=f'AI 首轮已返回，但第 {repair_round + 1} 轮 JSON 修复失败：{repair_exc.detail}',
                ) from repair_exc
    if raw is None:
        raise HTTPException(status_code=502, detail=f'AI 连续 3 轮 JSON 修复仍失败：{last_error_detail}')
    result = normalize_result_payload(raw)
    result['document_text_full'] = document_text
    result['document_text_excerpt'] = document_text[:12000]
    return result


def split_config_list(value: Any) -> set[str]:
    return {item.strip() for item in re.split(r'[\s,;，；]+', str(value or '')) if item.strip()}


def safe_feishu_id(value: str) -> str:
    return hashlib.sha256(value.encode('utf-8')).hexdigest()


def read_json_file(path: Path, default: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except (OSError, json.JSONDecodeError):
        return default


def write_json_file(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')


def cleanup_feishu_files() -> None:
    now = datetime.now()
    for folder, ttl in ((FEISHU_SESSIONS_DIR, FEISHU_SESSION_TTL), (FEISHU_EVENTS_DIR, FEISHU_EVENT_TTL)):
        for path in folder.glob('*.json'):
            try:
                mtime = datetime.fromtimestamp(path.stat().st_mtime)
            except OSError:
                continue
            if now - mtime > ttl:
                path.unlink(missing_ok=True)


def feishu_session_path(open_id: str, chat_id: str) -> Path:
    key = open_id or chat_id or 'anonymous'
    return FEISHU_SESSIONS_DIR / f'{safe_feishu_id(key)}.json'


def load_feishu_session(open_id: str, chat_id: str) -> dict[str, Any]:
    cleanup_feishu_files()
    return read_json_file(feishu_session_path(open_id, chat_id), {})


def save_feishu_session(open_id: str, chat_id: str, session: dict[str, Any]) -> None:
    session['updated_at'] = datetime.now().isoformat(timespec='seconds')
    write_json_file(feishu_session_path(open_id, chat_id), session)


def clear_feishu_session(open_id: str, chat_id: str) -> None:
    feishu_session_path(open_id, chat_id).unlink(missing_ok=True)


def feishu_event_seen(event_key: str) -> bool:
    cleanup_feishu_files()
    if not event_key:
        return False
    path = FEISHU_EVENTS_DIR / f'{safe_feishu_id(event_key)}.json'
    if path.exists():
        return True
    write_json_file(path, {'event_key': event_key, 'created_at': datetime.now().isoformat(timespec='seconds')})
    return False


def decrypt_feishu_payload(encrypted: str, encrypt_key: str) -> dict[str, Any]:
    if Cipher is None or algorithms is None or modes is None or padding is None:
        raise HTTPException(status_code=500, detail='缺少 cryptography 依赖，无法解密飞书事件。')
    key = hashlib.sha256(encrypt_key.encode('utf-8')).digest()
    try:
        cipher_text = base64.b64decode(encrypted)
        decryptor = Cipher(algorithms.AES(key), modes.CBC(key[:16])).decryptor()
        padded = decryptor.update(cipher_text) + decryptor.finalize()
        unpadder = padding.PKCS7(128).unpadder()
        plain = unpadder.update(padded) + unpadder.finalize()
        return json.loads(plain.decode('utf-8'))
    except Exception as exc:
        raise HTTPException(status_code=400, detail='飞书事件解密失败，请检查 Encrypt Key。') from exc


def load_feishu_event_payload(payload: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    if payload.get('encrypt'):
        encrypt_key = config.get('feishu_encrypt_key', '')
        if not encrypt_key:
            raise HTTPException(status_code=400, detail='飞书事件已加密，但系统未配置 Encrypt Key。')
        return decrypt_feishu_payload(str(payload.get('encrypt') or ''), encrypt_key)
    return payload


def verify_feishu_token(payload: dict[str, Any], config: dict[str, Any]) -> bool:
    expected = config.get('feishu_verification_token', '')
    if not expected:
        return True
    actual = payload.get('token') or (payload.get('header') or {}).get('token')
    return hmac.compare_digest(str(actual or ''), str(expected))


def extract_feishu_text(content: Any) -> str:
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError:
            return content.strip()
    if not isinstance(content, dict):
        return ''
    if isinstance(content.get('text'), str):
        return content['text'].strip()
    if isinstance(content.get('title'), str):
        return content['title'].strip()
    return ''


def collect_feishu_files(content: Any) -> list[dict[str, str]]:
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError:
            return []
    if not isinstance(content, dict):
        return []
    files: list[dict[str, str]] = []
    keys = ('file_key', 'fileKey', 'image_key', 'media_id')
    if any(content.get(key) for key in keys):
        files.append(
            {
                'file_key': str(next((content.get(key) for key in keys if content.get(key)), '')),
                'file_name': str(content.get('file_name') or content.get('name') or content.get('title') or 'feishu_file'),
                'file_type': str(content.get('file_type') or content.get('type') or ''),
            }
        )
    for value in content.values():
        if isinstance(value, list):
            for item in value:
                files.extend(collect_feishu_files(item))
        elif isinstance(value, dict):
            files.extend(collect_feishu_files(value))
    unique: dict[str, dict[str, str]] = {}
    for item in files:
        if item.get('file_key'):
            unique[item['file_key']] = item
    return list(unique.values())


def extract_feishu_message(payload: dict[str, Any]) -> dict[str, Any] | None:
    event = payload.get('event') if isinstance(payload.get('event'), dict) else payload
    message = event.get('message') if isinstance(event.get('message'), dict) else {}
    sender = event.get('sender') if isinstance(event.get('sender'), dict) else {}
    sender_id = sender.get('sender_id') if isinstance(sender.get('sender_id'), dict) else {}
    header = payload.get('header') if isinstance(payload.get('header'), dict) else {}
    message_id = str(message.get('message_id') or event.get('message_id') or '')
    content = message.get('content') or event.get('content') or {}
    return {
        'event_id': str(header.get('event_id') or payload.get('uuid') or event.get('event_id') or message_id),
        'event_type': str(header.get('event_type') or payload.get('type') or ''),
        'message_id': message_id,
        'message_type': str(message.get('message_type') or event.get('message_type') or ''),
        'chat_id': str(message.get('chat_id') or event.get('open_chat_id') or event.get('chat_id') or ''),
        'open_id': str(sender_id.get('open_id') or event.get('open_id') or ''),
        'text': extract_feishu_text(content),
        'files': collect_feishu_files(content),
    }


def feishu_token(config: dict[str, Any]) -> str:
    now = time.time()
    if FEISHU_TOKEN_CACHE.get('token') and float(FEISHU_TOKEN_CACHE.get('expires_at') or 0) > now + 60:
        return str(FEISHU_TOKEN_CACHE['token'])
    if not config.get('feishu_app_id') or not config.get('feishu_app_secret'):
        raise HTTPException(status_code=400, detail='飞书 App ID 或 App Secret 未配置。')
    response = httpx.post(
        'https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal',
        json={'app_id': config['feishu_app_id'], 'app_secret': config['feishu_app_secret']},
        timeout=20,
    )
    data = response.json()
    if response.status_code >= 400 or data.get('code') != 0:
        raise HTTPException(status_code=502, detail=f"飞书 tenant_access_token 获取失败：{data.get('msg') or response.text[:120]}")
    FEISHU_TOKEN_CACHE['token'] = data['tenant_access_token']
    FEISHU_TOKEN_CACHE['expires_at'] = now + int(data.get('expire') or 7200)
    return str(FEISHU_TOKEN_CACHE['token'])


def feishu_headers(config: dict[str, Any]) -> dict[str, str]:
    return {'Authorization': f'Bearer {feishu_token(config)}', 'Content-Type': 'application/json; charset=utf-8'}


def feishu_reply_message(config: dict[str, Any], message_id: str, text: str) -> None:
    if not message_id:
        return
    safe_text = sanitize_feishu_reply(text)
    response = httpx.post(
        f'https://open.feishu.cn/open-apis/im/v1/messages/{message_id}/reply',
        headers=feishu_headers(config),
        json={'msg_type': 'text', 'content': json.dumps({'text': safe_text}, ensure_ascii=False)},
        timeout=20,
    )
    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=f'飞书消息回复失败：{response.status_code} {response.text[:160]}')


def feishu_send_message(config: dict[str, Any], receive_id_type: str, receive_id: str, text: str) -> None:
    if receive_id_type not in {'open_id', 'chat_id'}:
        raise HTTPException(status_code=400, detail='receive_id_type 仅支持 open_id 或 chat_id。')
    if not receive_id:
        raise HTTPException(status_code=400, detail='请填写接收人的 open_id 或群 chat_id。')
    response = httpx.post(
        f'https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type={receive_id_type}',
        headers=feishu_headers(config),
        json={
            'receive_id': receive_id,
            'msg_type': 'text',
            'content': json.dumps({'text': sanitize_feishu_reply(text)}, ensure_ascii=False),
        },
        timeout=20,
    )
    data = {}
    try:
        data = response.json()
    except json.JSONDecodeError:
        pass
    if response.status_code >= 400 or (data and data.get('code') not in (0, None)):
        raise HTTPException(status_code=502, detail=f"飞书测试消息发送失败：{data.get('msg') if data else response.text[:160]}")


def feishu_download_file(config: dict[str, Any], message_id: str, file_key: str, file_name: str = '') -> Path:
    suffix = Path(file_name or '').suffix.lower()
    if suffix not in {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}:
        suffix = '.bin'
    temp_path = TMP_DIR / f'{uuid.uuid4().hex}{suffix}'
    urls = [
        f'https://open.feishu.cn/open-apis/im/v1/messages/{message_id}/resources/{file_key}?type=file',
        f'https://open.feishu.cn/open-apis/im/v1/messages/{message_id}/resources/{file_key}?type=docx',
    ]
    last_error = ''
    for url in urls:
        response = httpx.get(url, headers={'Authorization': f'Bearer {feishu_token(config)}'}, timeout=60)
        if response.status_code < 400 and response.content:
            temp_path.write_bytes(response.content)
            return temp_path
        last_error = f'{response.status_code} {response.text[:120]}'
    raise HTTPException(status_code=502, detail=f'飞书文件下载失败：{last_error}')


def sanitize_feishu_reply(text: Any) -> str:
    value = str(text or '')
    api_key = load_config().get('api_key', '')
    if api_key:
        value = value.replace(api_key, '[hidden]')
    for key in ('feishu_app_secret', 'feishu_verification_token', 'feishu_encrypt_key'):
        secret = load_config().get(key, '')
        if secret:
            value = value.replace(secret, '[hidden]')
    value = value.replace(str(BASE_DIR), '[project]')
    return value[:3800]


def sanitize_agent_log_value(value: Any) -> Any:
    config = load_config()
    secrets = [
        str(config.get('api_key') or ''),
        str(config.get('feishu_app_secret') or ''),
        str(config.get('feishu_verification_token') or ''),
        str(config.get('feishu_encrypt_key') or ''),
        str(BASE_DIR),
    ]

    def sanitize_text(text: str) -> str:
        cleaned = text
        for secret in secrets:
            if secret:
                cleaned = cleaned.replace(secret, '[hidden]')
        return cleaned

    if isinstance(value, str):
        return sanitize_text(value)
    if isinstance(value, list):
        return [sanitize_agent_log_value(item) for item in value]
    if isinstance(value, dict):
        sanitized: dict[str, Any] = {}
        for key, item in value.items():
            key_text = str(key)
            if any(secret_key in key_text.lower() for secret_key in ('api_key', 'secret', 'token', 'encrypt_key', 'authorization')):
                sanitized[key_text] = '[hidden]' if item else ''
            else:
                sanitized[key_text] = sanitize_agent_log_value(item)
        return sanitized
    return value


def write_feishu_agent_log(message: dict[str, Any], decision: dict[str, Any], tool_results: list[dict[str, Any]], reply: str, error: str = '') -> None:
    payload = {
        'id': uuid.uuid4().hex,
        'created_at': datetime.now().isoformat(timespec='seconds'),
        'open_id': message.get('open_id', ''),
        'chat_id': message.get('chat_id', ''),
        'message_id': message.get('message_id', ''),
        'user_text': message.get('text', ''),
        'files': message.get('files') or [],
        'decision': decision,
        'tool_results': tool_results,
        'reply': reply,
        'error': error,
    }
    safe_payload = sanitize_agent_log_value(payload)
    date_dir = FEISHU_AGENT_LOGS_DIR / datetime.now().strftime('%Y%m%d')
    date_dir.mkdir(parents=True, exist_ok=True)
    path = date_dir / f"{payload['created_at'].replace(':', '').replace('-', '').replace('T', '_')}_{payload['id']}.json"
    path.write_text(json.dumps(safe_payload, ensure_ascii=False, indent=2), encoding='utf-8')


def feishu_message_allowed(config: dict[str, Any], message: dict[str, Any]) -> bool:
    allowed_open_ids = split_config_list(config.get('feishu_allowed_open_ids'))
    allowed_chat_ids = split_config_list(config.get('feishu_allowed_chat_ids'))
    if not allowed_open_ids and not allowed_chat_ids:
        return True
    open_id = compact_text(message.get('open_id'))
    chat_id = compact_text(message.get('chat_id'))
    return bool((open_id and open_id in allowed_open_ids) or (chat_id and chat_id in allowed_chat_ids))


def load_feishu_agent_logs(date: str = '', limit: int = 50) -> list[dict[str, Any]]:
    limit = max(1, min(limit, 500))
    if date:
        if not re.fullmatch(r'\d{8}', date):
            raise HTTPException(status_code=400, detail='日志日期格式应为 YYYYMMDD。')
        paths = sorted((FEISHU_AGENT_LOGS_DIR / date).glob('*.json'), key=lambda path: path.stat().st_mtime, reverse=True)
    else:
        paths = sorted(FEISHU_AGENT_LOGS_DIR.glob('*/*.json'), key=lambda path: path.stat().st_mtime, reverse=True)
    items: list[dict[str, Any]] = []
    for path in paths[:limit]:
        try:
            item = json.loads(path.read_text(encoding='utf-8'))
        except (OSError, json.JSONDecodeError):
            continue
        if isinstance(item, dict):
            item['log_file'] = path.name
            item['log_date'] = path.parent.name
            items.append(item)
    return items


def feishu_help_text() -> str:
    return (
        '标书解析机器人可用指令：\n'
        '1. 帮助\n'
        '2. 查询 福海创\n'
        '3. 4000吨甲苯，张家港～东莞，月内装出\n'
        '4. 新造船资讯：船名/船厂/船东/DWT/状态...\n'
        '5. 生成本月商机市场分析\n'
        '6. 生成本月新造船市场分析\n'
        '7. 发送标书文件并输入“解析标书”\n'
        '下一步：直接发送你要处理的内容；识别出的市场情报不会自动保存，回复“确认保存”才入库。'
    )


def feishu_period_from_text(text: str) -> tuple[str, str, str]:
    now = datetime.now()
    if '本周' in text:
        start = now - timedelta(days=now.weekday())
        return 'weekly', start.strftime('%Y-%m-%d'), now.strftime('%Y-%m-%d')
    if '上月' in text:
        first = now.replace(day=1)
        end = first - timedelta(days=1)
        start = end.replace(day=1)
        return 'monthly', start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')
    if '本月' in text or '月' in text:
        return 'monthly', now.replace(day=1).strftime('%Y-%m-%d'), now.strftime('%Y-%m-%d')
    return 'custom', '', ''


def heuristic_feishu_route(text: str, files: list[dict[str, str]]) -> dict[str, Any]:
    stripped = text.strip()
    compact = stripped.replace(' ', '')
    if compact.lower() in {'help', '帮助', '菜单'}:
        return {'tool': 'help', 'arguments': {}}
    if compact in {'确认保存', '确认', '保存'}:
        return {'tool': 'confirm_save_market_record', 'arguments': {}}
    if compact in {'取消', '取消保存', '放弃'}:
        return {'tool': 'cancel_pending', 'arguments': {}}
    if files and any(word in stripped for word in ['解析', '标书', '招标']):
        return {'tool': 'parse_bid_file', 'arguments': {}}
    if any(word in stripped for word in ['新造船', '船厂', '造船', 'DWT', 'dwt', '完造', '出厂', '投运']):
        if any(word in stripped for word in ['分析', '报告', '周报', '月报']):
            period, start_date, end_date = feishu_period_from_text(stripped)
            return {'tool': 'market_report_newbuilding', 'arguments': {'period': period, 'start_date': start_date, 'end_date': end_date}}
        return {'tool': 'extract_newbuilding', 'arguments': {'text': stripped}}
    if '商机' in stripped and any(word in stripped for word in ['分析', '报告', '周报', '月报', '市场']):
        period, start_date, end_date = feishu_period_from_text(stripped)
        return {'tool': 'market_report_cargo', 'arguments': {'period': period, 'start_date': start_date, 'end_date': end_date}}
    if any(word in stripped for word in ['货盘', '吨', '装出', '装港', '卸港', '～', '~', '->']):
        return {'tool': 'extract_cargo', 'arguments': {'text': stripped}}
    if any(word in stripped for word in ['查询', '查找', '搜索']):
        keyword = re.sub(r'^(查询|查找|搜索)', '', stripped).replace('项目', '').strip()
        return {'tool': 'search_project', 'arguments': {'q': keyword or stripped}}
    return {'tool': 'search_project', 'arguments': {'q': stripped}}


def ai_feishu_route(text: str, files: list[dict[str, str]], config: dict[str, Any]) -> dict[str, Any]:
    if not (config.get('base_url') and config.get('api_key') and config.get('model')):
        return heuristic_feishu_route(text, files)
    prompt = json.dumps(
        {
            'text': text,
            'files': files,
            'tools': sorted(FEISHU_TOOL_NAMES),
            'required_json': {'tool': 'help', 'arguments': {}},
        },
        ensure_ascii=False,
    )
    try:
        content = call_chat_completion(
            config,
            system_prompt=(
                '你是飞书机器人指令路由器，只能输出严格 JSON。'
                'tool 必须是给定白名单之一；arguments 只能放工具需要的参数；'
                '保存、删除、覆盖类动作不要直接执行，必须路由到确认或取消。'
                '最终回复由后端生成，后端会提醒用户下一步操作。'
            ),
            user_prompt=prompt,
            temperature=0,
        )
        route = parse_json_text(content)
    except HTTPException:
        return heuristic_feishu_route(text, files)
    tool = str(route.get('tool') or '')
    if tool not in FEISHU_TOOL_NAMES:
        return heuristic_feishu_route(text, files)
    arguments = route.get('arguments') if isinstance(route.get('arguments'), dict) else {}
    return {'tool': tool, 'arguments': arguments}


def format_project_search(items: list[dict[str, Any]]) -> str:
    if not items:
        return '没有查到匹配的历史项目。\n下一步：请换一个项目关键词，或到网页端历史台账确认是否已保存。'
    lines = ['查询到以下历史项目：']
    for index, item in enumerate(items[:5], start=1):
        lines.append(
            f"{index}. {item.get('title') or item.get('project_name') or '未命名项目'}\n"
            f"   招标编号：{item.get('bid_no') or '-'}\n"
            f"   招标人：{item.get('tenderer') or '-'}\n"
            f"   投标状态：{item.get('bid_status') or '-'}，中标状态：{item.get('award_status') or '-'}"
        )
    lines.append('\n下一步：如果要看详情，请在网页端打开历史台账；如果要查其他项目，继续发送“查询 项目名”。')
    return '\n'.join(lines)


def format_market_record(kind: str, record: dict[str, Any]) -> str:
    if kind == 'cargo':
        lines = [
            '已识别商机货盘草稿：',
            f"货品：{record.get('cargo_name') or '-'}",
            f"吨数：{record.get('tonnage') or '-'}",
            f"航线：{record.get('load_port') or '-'} -> {record.get('discharge_port') or '-'}",
            f"装载期：{record.get('laycan') or '-'}",
            f"货主：{record.get('cargo_owner') or '-'}",
            f"板块：{record.get('segment') or '-'}，类型：{record.get('board_type') or '-'}",
        ]
    else:
        lines = [
            '已识别新造船资讯草稿：',
            f"船名：{record.get('ship_name') or '-'}",
            f"船厂：{record.get('shipyard') or '-'}",
            f"船东：{record.get('owner') or '-'}",
            f"DWT：{record.get('dwt') or '-'}",
            f"阶段：{record.get('stage') or '-'}",
            f"预计交付：{record.get('delivery_time') or '-'}",
        ]
    lines.append('下一步：请核对字段；回复“确认保存”写入市场情报台账，或回复“取消”放弃。')
    return '\n'.join(lines)


def format_market_report(report: dict[str, Any]) -> str:
    lines = [str(report.get('summary') or '已生成市场分析报告。')]
    for title, key in [('关键发现', 'key_findings'), ('风险提示', 'risks'), ('经营建议', 'recommendations')]:
        values = [str(item) for item in (report.get(key) or []) if str(item).strip()]
        if values:
            lines.append(f'\n{title}：')
            lines.extend(f'- {item}' for item in values[:5])
    lines.append('\n下一步：如需沉淀报告或导出 Word/Excel，请到网页端“市场情报”模块导出；也可以继续发送新的时间范围生成分析。')
    return '\n'.join(lines)


def format_bid_result(result: dict[str, Any]) -> str:
    summary = result.get('document_summary') or {}
    analysis = result.get('analysis') or {}
    risks = analysis.get('risks') or []
    return (
        '标书解析完成：\n'
        f"项目名称：{summary.get('project_name') or '-'}\n"
        f"招标编号：{summary.get('bid_no') or '-'}\n"
        f"招标人：{summary.get('tenderer') or '-'}\n"
        f"投标截止：{summary.get('bid_deadline') or '-'}\n"
        f"保证金：{summary.get('deposit_amount') or '-'}\n"
        f"资格要求：{compact_paragraph(summary.get('qualification_requirements'))[:600] or '-'}\n"
        f"风险提示：{'；'.join(map(str, risks[:5])) or '-'}\n"
        '下一步：文件未自动入历史，请到网页端核对摘取结果、导出 Excel，并确认是否保存到历史台账。'
    )


def execute_feishu_tool(route: dict[str, Any], message: dict[str, Any], config: dict[str, Any]) -> str:
    tool = route.get('tool')
    args = route.get('arguments') if isinstance(route.get('arguments'), dict) else {}
    text = args.get('text') or message.get('text') or ''
    open_id = message.get('open_id', '')
    chat_id = message.get('chat_id', '')
    if tool == 'help':
        return feishu_help_text()
    if tool == 'cancel_pending':
        clear_feishu_session(open_id, chat_id)
        return '已取消当前待确认内容。'
    if tool == 'confirm_save_market_record':
        session = load_feishu_session(open_id, chat_id)
        pending = session.get('pending') if isinstance(session.get('pending'), dict) else {}
        if pending.get('type') != 'market_record':
            return '当前没有待保存内容。\n下一步：请先发送商机货盘或新造船资讯，我识别成草稿后再回复“确认保存”。'
        record = save_market_record(str(pending.get('kind') or ''), pending.get('record') or {})
        clear_feishu_session(open_id, chat_id)
        return f"已保存到市场情报台账：{record.get('id')}\n下一步：可继续发送新的商机/新造船信息，或到网页端查看台账、统计和导出。"
    if tool == 'search_project':
        items, _ = collect_project_items(q=str(args.get('q') or text))
        return format_project_search(items)
    if tool in {'extract_cargo', 'extract_newbuilding'}:
        kind = 'cargo' if tool == 'extract_cargo' else 'newbuilding'
        record = extract_market_record(kind, str(text))
        save_feishu_session(
            open_id,
            chat_id,
            {'pending': {'type': 'market_record', 'kind': kind, 'record': record, 'created_at': datetime.now().isoformat(timespec='seconds')}},
        )
        return format_market_record(kind, record)
    if tool in {'market_report_cargo', 'market_report_newbuilding'}:
        kind = 'cargo' if tool == 'market_report_cargo' else 'newbuilding'
        report = build_market_report(
            MarketReportPayload(
                kind=kind,
                period=str(args.get('period') or 'custom'),
                start_date=str(args.get('start_date') or ''),
                end_date=str(args.get('end_date') or ''),
                filters=args.get('filters') if isinstance(args.get('filters'), dict) else {},
            )
        )
        return format_market_report(report)
    if tool == 'parse_bid_file':
        files = message.get('files') or []
        if not files:
            return '请在飞书里同时发送标书文件，并输入“解析标书”。\n下一步：上传 PDF/DOC/DOCX/XLS/XLSX 后再发“解析标书”。'
        file_info = files[0]
        temp_path = feishu_download_file(config, message.get('message_id', ''), file_info.get('file_key', ''), file_info.get('file_name', ''))
        try:
            suffix = Path(file_info.get('file_name') or temp_path.name).suffix.lower() or temp_path.suffix.lower()
            if suffix not in {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}:
                raise HTTPException(status_code=400, detail='当前仅支持 PDF、DOC、DOCX、XLS、XLSX 标书文件。')
            document_text = extract_document_text(temp_path, suffix)
            if len(document_text.strip()) < 80:
                raise HTTPException(status_code=400, detail='当前版本仅支持可提取文本的标书文件，扫描件暂不支持。')
            return format_bid_result(parse_ai_document(document_text, require_config()))
        finally:
            temp_path.unlink(missing_ok=True)
    return feishu_help_text()


def handle_feishu_message(message: dict[str, Any], config: dict[str, Any]) -> str:
    route = ai_feishu_route(message.get('text', ''), message.get('files') or [], config)
    return execute_feishu_tool(route, message, config)


def process_feishu_event(message: dict[str, Any], config: dict[str, Any]) -> None:
    try:
        reply = handle_feishu_message(message, config)
    except Exception as exc:
        detail = exc.detail if isinstance(exc, HTTPException) else str(exc)
        reply = f'处理失败：{detail}'
    try:
        feishu_reply_message(config, message.get('message_id', ''), reply)
    except Exception as exc:
        print(f'Feishu reply failed: {exc}')


def object_attr(source: Any, name: str, default: Any = '') -> Any:
    return getattr(source, name, default) if source is not None else default


def ws_message_to_payload(event: Any) -> dict[str, Any]:
    header = object_attr(event, 'header', None)
    event_data = object_attr(event, 'event', None)
    sender = object_attr(event_data, 'sender', None)
    sender_id = object_attr(sender, 'sender_id', None)
    message = object_attr(event_data, 'message', None)
    return {
        'schema': '2.0',
        'header': {
            'event_id': object_attr(header, 'event_id', '') or object_attr(message, 'message_id', ''),
            'event_type': object_attr(header, 'event_type', 'im.message.receive_v1'),
            'token': object_attr(header, 'token', ''),
        },
        'event': {
            'sender': {
                'sender_id': {
                    'open_id': object_attr(sender_id, 'open_id', ''),
                    'user_id': object_attr(sender_id, 'user_id', ''),
                    'union_id': object_attr(sender_id, 'union_id', ''),
                }
            },
            'message': {
                'message_id': object_attr(message, 'message_id', ''),
                'chat_id': object_attr(message, 'chat_id', ''),
                'message_type': object_attr(message, 'message_type', ''),
                'content': object_attr(message, 'content', ''),
            },
        },
    }


def handle_feishu_ws_message(event: Any) -> None:
    config = load_config()
    if not config.get('feishu_enabled') or config.get('feishu_receive_mode') != 'ws':
        return
    payload = ws_message_to_payload(event)
    message = extract_feishu_message(payload)
    if not message:
        return
    event_key = message.get('event_id') or message.get('message_id')
    if feishu_event_seen(event_key):
        return
    process_feishu_event(message, config)


def feishu_dialog_defaults() -> dict[str, Any]:
    return {
        'active_skill': '',
        'seed_text': '',
        'slots': {},
        'missing': [],
        'pending_confirm': {},
        'last_question': '',
        'retry_count': 0,
        'last_result': {},
        'agent_summary': '',
        'last_tool_result': {},
    }


def feishu_dialog_load(open_id: str, chat_id: str) -> dict[str, Any]:
    session = feishu_dialog_defaults()
    loaded = load_feishu_session(open_id, chat_id)
    if isinstance(loaded, dict):
        session.update(loaded)
    session['active_skill'] = compact_text(session.get('active_skill'))
    session['seed_text'] = str(session.get('seed_text') or '')
    session['slots'] = session.get('slots') if isinstance(session.get('slots'), dict) else {}
    session['missing'] = [compact_text(item) for item in (session.get('missing') or []) if compact_text(item)]
    pending = session.get('pending_confirm')
    if not isinstance(pending, dict):
        pending = session.get('pending') if isinstance(session.get('pending'), dict) else {}
    session['pending_confirm'] = pending if isinstance(pending, dict) else {}
    session['last_question'] = compact_text(session.get('last_question'))
    session['retry_count'] = int(session.get('retry_count') or 0)
    session['last_result'] = session.get('last_result') if isinstance(session.get('last_result'), dict) else {}
    session['agent_summary'] = compact_text(session.get('agent_summary'))
    session['last_tool_result'] = session.get('last_tool_result') if isinstance(session.get('last_tool_result'), dict) else {}
    return session


def feishu_dialog_save(open_id: str, chat_id: str, session: dict[str, Any]) -> None:
    payload = feishu_dialog_defaults()
    payload.update(session)
    payload['active_skill'] = compact_text(payload.get('active_skill'))
    payload['seed_text'] = str(payload.get('seed_text') or '')
    payload['slots'] = payload.get('slots') if isinstance(payload.get('slots'), dict) else {}
    payload['missing'] = [compact_text(item) for item in (payload.get('missing') or []) if compact_text(item)]
    payload['pending_confirm'] = payload.get('pending_confirm') if isinstance(payload.get('pending_confirm'), dict) else {}
    payload['last_question'] = compact_text(payload.get('last_question'))
    payload['retry_count'] = int(payload.get('retry_count') or 0)
    payload['last_result'] = payload.get('last_result') if isinstance(payload.get('last_result'), dict) else {}
    payload['agent_summary'] = compact_text(payload.get('agent_summary'))
    payload['last_tool_result'] = payload.get('last_tool_result') if isinstance(payload.get('last_tool_result'), dict) else {}
    save_feishu_session(open_id, chat_id, payload)


def feishu_dialog_clear(open_id: str, chat_id: str) -> None:
    clear_feishu_session(open_id, chat_id)


def feishu_skill_is_confirm(text: str) -> bool:
    compact = compact_text(text).replace(' ', '')
    return compact in {'确认', '确认保存', '保存', '好的', '可以', '行', '同意'}


def feishu_skill_is_cancel(text: str) -> bool:
    compact = compact_text(text).replace(' ', '')
    return compact in {'取消', '放弃', '不要了', '算了', '撤销'}


def feishu_skill_is_help(text: str) -> bool:
    compact = compact_text(text).replace(' ', '')
    return compact in {'帮助', '菜单', '功能', '你好', '在吗', '你能做什么', '怎么用', '说明'}


def feishu_skill_catalog_text() -> str:
    lines = ['我现在能做这些事：']
    for skill_name in FEISHU_SKILL_ORDER:
        meta = FEISHU_SKILLS[skill_name]
        lines.append(f"- {meta['title']}：{meta['description']}")
        example = compact_text(meta.get('example') or '')
        if example:
            lines.append(f"  例子：{example}")
    lines.extend(['', '你可以直接发自然语言，不用先选菜单。', '我会先告诉你系统能做什么，还缺什么，下一步怎么回。'])
    return '\n'.join(lines)


def feishu_guess_period(text: str) -> tuple[str, str, str]:
    now = datetime.now()
    compact = compact_text(text)
    if any(word in compact for word in ('本周', '这周', '最近7天', '近7天')):
        start = now - timedelta(days=now.weekday())
        return 'weekly', start.strftime('%Y-%m-%d'), now.strftime('%Y-%m-%d')
    if any(word in compact for word in ('上周', '上星期')):
        monday = now - timedelta(days=now.weekday())
        start = monday - timedelta(days=7)
        end = monday - timedelta(days=1)
        return 'weekly', start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')
    if any(word in compact for word in ('本月', '这个月', '当月')):
        return 'monthly', now.replace(day=1).strftime('%Y-%m-%d'), now.strftime('%Y-%m-%d')
    if any(word in compact for word in ('上月', '上个月')):
        first = now.replace(day=1)
        end = first - timedelta(days=1)
        start = end.replace(day=1)
        return 'monthly', start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')
    range_match = re.search(
        r'(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})[日]?\s*(?:到|至|~|-)\s*(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})[日]?',
        compact,
    )
    if range_match:
        start = datetime(int(range_match.group(1)), int(range_match.group(2)), int(range_match.group(3)))
        end = datetime(int(range_match.group(4)), int(range_match.group(5)), int(range_match.group(6)))
        period = 'weekly' if (end.date() - start.date()).days <= 7 else 'custom'
        return period, start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')
    one_date = re.search(r'(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})[日]?', compact)
    if one_date:
        date_text = datetime(int(one_date.group(1)), int(one_date.group(2)), int(one_date.group(3)))
        return 'custom', date_text.strftime('%Y-%m-%d'), date_text.strftime('%Y-%m-%d')
    return 'custom', '', ''


def feishu_skill_route_prompt(text: str, files: list[dict[str, str]], session: dict[str, Any]) -> str:
    return json.dumps(
        {
            'user_text': text,
            'files': files,
            'session': {
                'active_skill': session.get('active_skill', ''),
                'slots': session.get('slots', {}),
                'missing': session.get('missing', []),
                'pending_confirm': bool(session.get('pending_confirm')),
                'last_question': session.get('last_question', ''),
                'seed_text': session.get('seed_text', ''),
            },
            'skills': [
                {
                    'skill': skill_name,
                    'title': meta['title'],
                    'description': meta['description'],
                    'needs_file': bool(meta.get('needs_file')),
                    'needs_period': bool(meta.get('needs_period')),
                    'confirmation': bool(meta.get('confirmation')),
                    'example': meta.get('example') or '',
                }
                for skill_name, meta in FEISHU_SKILLS.items()
            ],
            'output_schema': {
                'skill': 'help',
                'action': 'clarify',
                'confidence': 0.0,
                'slots': {},
                'missing': [],
                'next_question': '',
                'reply_hint': '',
            },
        },
        ensure_ascii=False,
    )


def feishu_skill_route_from_ai(text: str, files: list[dict[str, str]], session: dict[str, Any], config: dict[str, Any]) -> dict[str, Any] | None:
    if not (config.get('base_url') and config.get('api_key') and config.get('model')):
        return None
    system_prompt = (
        '你是飞书里的系统总控 Skill。'
        '你的任务是先识别用户意图，再判断系统能做什么、还缺什么，最后输出严格 JSON。'
        '只能输出 JSON，不能输出 Markdown、解释、前后缀。'
        'skill 只能取 help、bid_parse、project_search、cargo_opportunity、newbuilding_info、cargo_report、newbuilding_report。'
        'action 只能取 help、clarify、preview、execute、confirm、cancel。'
        '如果用户是在延续上一轮对话，请结合 session.active_skill、session.slots、session.pending_confirm 判断。'
        '如果缺关键输入，action 必须是 clarify，并在 next_question 里给出一句可直接发送的补充问题。'
        '如果用户在问系统能做什么，直接返回 help。'
        '如果用户要确认保存，action 返回 confirm。'
        '如果用户要取消，action 返回 cancel。'
        '如果用户发了标书文件并说解析标书，skill 返回 bid_parse。'
    )
    try:
        current_content = call_chat_completion(
            config,
            system_prompt=system_prompt,
            user_prompt=feishu_skill_route_prompt(text, files, session),
            temperature=0,
        )
    except HTTPException:
        return None
    last_error = 'AI 返回内容无法解析。'
    for round_index in range(3):
        try:
            route = parse_json_text(current_content)
            if not isinstance(route, dict):
                raise HTTPException(status_code=502, detail='AI 返回的路由不是 JSON 对象。')
            skill = compact_text(route.get('skill'))
            action = compact_text(route.get('action'))
            if skill not in FEISHU_SKILLS or action not in {'help', 'clarify', 'preview', 'execute', 'confirm', 'cancel'}:
                raise HTTPException(status_code=502, detail='skill 或 action 不在允许列表内。')
            route['skill'] = skill
            route['action'] = action
            route['slots'] = route.get('slots') if isinstance(route.get('slots'), dict) else {}
            route['missing'] = [compact_text(item) for item in (route.get('missing') or []) if compact_text(item)]
            route['next_question'] = compact_text(route.get('next_question'))
            route['reply_hint'] = compact_text(route.get('reply_hint'))
            route['confidence'] = float(route.get('confidence') or 0)
            return route
        except HTTPException as exc:
            last_error = str(exc.detail)
            if round_index == 2:
                break
            try:
                current_content = call_chat_completion(
                    config,
                    system_prompt=JSON_REPAIR_SYSTEM_PROMPT,
                    user_prompt=build_json_repair_prompt(current_content, last_error),
                    temperature=0,
                )
            except HTTPException:
                break
    return None


def feishu_skill_from_text(text: str, files: list[dict[str, str]], session: dict[str, Any]) -> str:
    compact = compact_text(text).replace(' ', '')
    if feishu_skill_is_help(compact):
        return 'help'
    if files and any(name.lower().endswith(('.pdf', '.doc', '.docx', '.xls', '.xlsx')) for name in (item.get('file_name', '') for item in files)):
        if any(word in compact for word in ('解析', '标书', '招标', '投标')) or not compact:
            return 'bid_parse'
    if any(word in compact for word in ('商机报告', '商机分析', '商机市场分析', '运价', '成交价趋势')):
        return 'cargo_report'
    if any(word in compact for word in ('新造船报告', '新造船分析', '造船市场分析')):
        return 'newbuilding_report'
    if any(word in compact for word in ('新造船', '造船', '船厂', '船东', '交付', '出厂', 'DWT', 'dwt')):
        return 'newbuilding_info'
    if any(word in compact for word in ('货盘', '商机', '成交价', '运价', '装港', '卸港', '吨', '吨位', '甲苯', '溶剂', '成品油', '化学品')):
        return 'cargo_opportunity'
    if any(word in compact for word in ('查询', '查找', '搜索', '项目', '台账')):
        return 'project_search'
    if compact and session.get('active_skill') in FEISHU_SKILLS:
        return str(session.get('active_skill'))
    return 'help'


def feishu_skill_missing_fields(skill: str, text: str, slots: dict[str, Any]) -> list[str]:
    if skill == 'cargo_opportunity':
        name = compact_text(slots.get('cargo_name'))
        load_port = compact_text(slots.get('load_port'))
        discharge_port = compact_text(slots.get('discharge_port'))
        tonnage = compact_text(slots.get('tonnage'))
        if name and load_port and discharge_port and tonnage:
            return []
        return [field for field in ['cargo_name', 'tonnage', 'load_port', 'discharge_port'] if not compact_text(slots.get(field))]
    if skill == 'newbuilding_info':
        ship_name = compact_text(slots.get('ship_name'))
        shipyard = compact_text(slots.get('shipyard'))
        owner = compact_text(slots.get('owner'))
        dwt = compact_text(slots.get('dwt'))
        if sum(bool(item) for item in [ship_name, shipyard, owner, dwt]) >= 2:
            return []
        return [field for field in ['ship_name', 'shipyard', 'owner', 'dwt'] if not compact_text(slots.get(field))]
    if skill in {'cargo_report', 'newbuilding_report'}:
        period, start_date, end_date = feishu_guess_period(text)
        if period == 'custom' and not (start_date and end_date):
            return ['period']
        return []
    if skill == 'project_search':
        query = compact_text(slots.get('query') or text)
        return [] if query else ['query']
    return []


def feishu_skill_next_question(skill: str, missing: list[str]) -> str:
    if skill == 'cargo_opportunity':
        labels = {'cargo_name': '货名', 'tonnage': '吨数', 'load_port': '装港', 'discharge_port': '卸港'}
        missing_text = '、'.join(labels.get(item, item) for item in missing) or '关键信息'
        return f'我先按商机货盘处理。请补充：{missing_text}。'
    if skill == 'newbuilding_info':
        labels = {'ship_name': '船名', 'shipyard': '船厂', 'owner': '船东', 'dwt': 'DWT'}
        missing_text = '、'.join(labels.get(item, item) for item in missing) or '关键信息'
        return f'我先按新造船信息处理。请补充：{missing_text}。'
    if skill == 'project_search':
        return '请告诉我项目名、招标编号，或者直接发“查询 + 关键词”。'
    if skill in {'cargo_report', 'newbuilding_report'}:
        return '请先选时间范围，例如“本周”“本月”或者“2026-06-01 到 2026-06-30”。'
    if skill == 'bid_parse':
        return '请先上传标书文件，再说“解析标书”。当前版本支持 PDF、DOC、DOCX、XLS、XLSX。'
    return '你可以直接发自然语言，我来帮你判断要走哪个功能。'


def feishu_skill_reply_prefix(skill: str, action: str, missing: list[str]) -> str:
    title = FEISHU_SKILLS.get(skill, {}).get('title', skill)
    lines = [f'我识别到的意图：{title}']
    if action == 'clarify':
        lines.append(f"还缺：{', '.join(missing) if missing else '少量关键信息'}")
    elif action == 'confirm':
        lines.append('已按当前草稿执行保存。')
    elif action == 'cancel':
        lines.append('已取消当前待办。')
    return '\n'.join(lines)


FEISHU_AGENT_TOOLS = {
    'chat_general',
    'extract_cargo_opportunity',
    'extract_newbuilding_info',
    'update_pending_record',
    'save_pending_record',
    'check_last_saved_record',
    'search_project',
    'search_market_records',
    'parse_bid_file',
    'answer_bid_question',
    'save_bid_project',
    'generate_market_report',
    'cancel_pending',
}


def require_ai_chat_config(config: dict[str, Any]) -> None:
    if not (config.get('base_url') and config.get('api_key') and config.get('model')):
        raise HTTPException(status_code=400, detail='AI 暂不可用：请先在网页端 AI 配置中保存 Base URL、API Key 和模型名称。')


def feishu_agent_session_view(session: dict[str, Any]) -> dict[str, Any]:
    pending = session.get('pending_confirm') if isinstance(session.get('pending_confirm'), dict) else {}
    pending_record = pending.get('record') if isinstance(pending.get('record'), dict) else {}
    return {
        'active_skill': session.get('active_skill', ''),
        'agent_summary': session.get('agent_summary', ''),
        'last_question': session.get('last_question', ''),
        'has_pending_record': bool(pending_record),
        'pending_kind': pending.get('kind', ''),
        'pending_record': pending_record,
        'last_saved_record_id': (session.get('last_tool_result') or {}).get('saved_record_id', ''),
        'last_saved_record_kind': (session.get('last_tool_result') or {}).get('saved_record_kind', ''),
        'has_bid_result': bool(session.get('last_result') and session.get('active_skill') == 'bid_parse'),
        'last_tool_result': session.get('last_tool_result', {}),
    }


def feishu_agent_decision_prompt(text: str, files: list[dict[str, str]], session: dict[str, Any]) -> str:
    return json.dumps(
        {
            'user_text': text,
            'files': files,
            'session': feishu_agent_session_view(session),
            'available_tools': sorted(FEISHU_AGENT_TOOLS),
            'tool_policy': {
                'chat_general': '普通聊天、寒暄、解释系统能力、无法判断业务动作时使用。普通聊天必须用这个工具，不要返回固定菜单。',
                'extract_cargo_opportunity': '从货盘/商机文本识别商机草稿，但不保存。',
                'extract_newbuilding_info': '从新造船资讯识别草稿，但不保存。',
                'update_pending_record': '用户正在修改当前草稿，例如“装港改成宁波”。',
                'save_pending_record': '用户明确确认保存当前草稿时使用。',
                'check_last_saved_record': '用户询问刚才那条、上一条、是否保存成功、记录 ID 时使用。',
                'search_project': '查询历史项目。',
                'search_market_records': '查询商机货盘或新造船市场情报，可带 kind/query/status/stage/segment/board_type/start_date/end_date。',
                'parse_bid_file': '用户上传标书文件并要求解析时使用。',
                'answer_bid_question': '已有最近标书解析结果，用户继续追问标书内容时使用。',
                'save_bid_project': '用户明确确认保存最近一次标书解析结果为项目时使用。',
                'generate_market_report': '生成商机或新造船市场分析报告。',
                'cancel_pending': '用户取消当前草稿或待办。',
            },
            'required_json': {
                'reply': '',
                'tool_calls': [{'name': 'chat_general', 'arguments': {}}],
                'needs_confirmation': False,
                'pending_update': {},
            },
        },
        ensure_ascii=False,
    )


def feishu_agent_fallback_tool_call(text: str, files: list[dict[str, str]], session: dict[str, Any]) -> dict[str, Any] | None:
    compact = compact_text(text)
    pending = session.get('pending_confirm') if isinstance(session.get('pending_confirm'), dict) else {}
    if pending.get('type') == 'market_record':
        if any(word in compact for word in ('取消', '算了', '不要保存', '撤销')):
            return {'name': 'cancel_pending', 'arguments': {}}
        if compact in {'保存', '确认', '确认保存', '可以保存', '没问题'}:
            return {'name': 'save_pending_record', 'arguments': {}}
        if any(word in compact for word in ('改成', '改为', '调整为', '换成', '设为', '补充', '加上')):
            return {'name': 'update_pending_record', 'arguments': {'text': compact}}
    if session.get('active_skill') == 'bid_parse' and session.get('last_result'):
        if compact in {'保存', '确认', '确认保存', '保存项目', '保存成项目'}:
            return {'name': 'save_bid_project', 'arguments': {}}
        if compact:
            return {'name': 'answer_bid_question', 'arguments': {'question': compact}}
    skill = feishu_skill_from_text(compact, files, session)
    if skill == 'cargo_opportunity':
        return {'name': 'extract_cargo_opportunity', 'arguments': {'text': compact}}
    if skill == 'newbuilding_info':
        return {'name': 'extract_newbuilding_info', 'arguments': {'text': compact}}
    if skill == 'project_search':
        if any(word in compact for word in ('货盘', '商机', '新造船', '船厂', '船东', '航线')):
            return {'name': 'search_market_records', 'arguments': {'query': compact}}
        return {'name': 'search_project', 'arguments': {'query': compact}}
    if skill == 'bid_parse':
        return {'name': 'parse_bid_file', 'arguments': {}}
    if skill == 'cargo_report':
        return {'name': 'generate_market_report', 'arguments': {'kind': 'cargo'}}
    if skill == 'newbuilding_report':
        return {'name': 'generate_market_report', 'arguments': {'kind': 'newbuilding'}}
    return None


def feishu_agent_decide(text: str, files: list[dict[str, str]], session: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    require_ai_chat_config(config)
    system_prompt = (
        '你是飞书里的业务 AI Agent。用户是在和配置好的 AI 自然对话，不是在操作固定菜单。'
        '你必须先理解用户消息，再决定是否调用后端白名单工具。'
        '普通聊天、寒暄、泛泛提问必须调用 chat_general，由 AI 自然回复，不要返回固定菜单。'
        '保存、写入类动作只有用户明确确认时才能调用 save_pending_record。'
        '如果需要工具，只输出 JSON，不要 Markdown。tool_calls 最多 2 个。'
        '工具名必须来自 available_tools；如果不确定，用 chat_general。'
    )
    content = call_chat_completion(
        config,
        system_prompt=system_prompt,
        user_prompt=feishu_agent_decision_prompt(text, files, session),
        temperature=0,
    )
    try:
        decision = parse_json_text(content)
    except HTTPException:
        return {'reply': '', 'tool_calls': [{'name': 'chat_general', 'arguments': {}}], 'needs_confirmation': False, 'pending_update': {}}
    tool_calls = decision.get('tool_calls') if isinstance(decision.get('tool_calls'), list) else []
    cleaned_calls: list[dict[str, Any]] = []
    for call in tool_calls[:2]:
        if not isinstance(call, dict):
            continue
        name = compact_text(call.get('name'))
        if name not in FEISHU_AGENT_TOOLS:
            continue
        args = call.get('arguments') if isinstance(call.get('arguments'), dict) else {}
        cleaned_calls.append({'name': name, 'arguments': args})
    if not cleaned_calls:
        cleaned_calls = [{'name': 'chat_general', 'arguments': {}}]
    if len(cleaned_calls) == 1 and cleaned_calls[0]['name'] == 'chat_general':
        fallback_call = feishu_agent_fallback_tool_call(text, files, session)
        if fallback_call:
            cleaned_calls = [fallback_call]
    return {
        'reply': compact_text(decision.get('reply')),
        'tool_calls': cleaned_calls,
        'needs_confirmation': bool(decision.get('needs_confirmation')),
        'pending_update': decision.get('pending_update') if isinstance(decision.get('pending_update'), dict) else {},
    }


def feishu_agent_chat_reply(text: str, session: dict[str, Any], config: dict[str, Any], tool_results: list[dict[str, Any]] | None = None) -> str:
    require_ai_chat_config(config)
    prompt = json.dumps(
        {
            'user_text': text,
            'session': feishu_agent_session_view(session),
            'tool_results': tool_results or [],
            'reply_rules': [
                '用自然中文回复，像一个可协作的业务 AI 助手。',
                '不要说自己是固定菜单机器人。',
                '如果工具已生成草稿，简洁列出关键字段并提示确认保存或继续修改。',
                '如果没有工具结果，就正常回答用户问题。',
            ],
        },
        ensure_ascii=False,
    )
    return call_chat_completion(
        config,
        system_prompt='你是飞书里的业务 AI Agent，负责自然回复用户，并在需要时说明系统已完成的操作。',
        user_prompt=prompt,
        temperature=0.3,
    ).strip()


def feishu_extract_direct_market_updates(kind: str, text: str) -> dict[str, str]:
    updates: dict[str, str] = {}
    field_aliases = {
        'cargo': {
            'load_port': ['装港', '装货港', '起运港', '起始港'],
            'discharge_port': ['卸港', '卸货港', '目的港', '到港'],
            'cargo_name': ['货品', '货名', '品名'],
            'tonnage': ['吨数', '货量', '吨位'],
            'laycan': ['装期', '装载期', '受载期'],
            'cargo_owner': ['货主', '租家', '客户'],
            'board_type': ['货盘类型', '类型'],
            'segment': ['业务板块', '板块'],
        },
        'newbuilding': {
            'ship_name': ['船名'],
            'shipyard': ['船厂', '造船厂'],
            'owner': ['船东'],
            'dwt': ['载重吨', 'DWT', 'dwt'],
            'delivery_time': ['交付', '交付时间', '预计交付'],
            'ship_type': ['船型'],
            'stage': ['阶段', '建造阶段'],
        },
    }.get(kind, {})
    for key, aliases in field_aliases.items():
        for alias in aliases:
            pattern = rf'{re.escape(alias)}\s*(?:改成|改为|调整为|换成|设为|是|为|：|:)\s*([^，,。；;\n]+)'
            match = re.search(pattern, text, flags=re.I)
            if match:
                updates[key] = compact_text(match.group(1))
                break
    return updates


def feishu_agent_merge_record(kind: str, current: dict[str, Any], text: str) -> dict[str, Any]:
    update = extract_market_record(kind, text)
    update.update(feishu_extract_direct_market_updates(kind, text))
    merged = dict(current)
    for key, value in update.items():
        if key in {'id', 'kind', 'created_at', 'updated_at'}:
            continue
        if compact_text(value):
            merged[key] = value
    raw_parts = [compact_text(current.get('raw_text')), text]
    merged['raw_text'] = '\n'.join(part for part in raw_parts if part)
    return normalize_market_record(kind, merged, record_id=compact_text(current.get('id')), created_at=compact_text(current.get('created_at')))


def feishu_record_source_fields(message: dict[str, Any], *, confirmed_at: str = '') -> dict[str, str]:
    return {
        'source': '飞书',
        'source_channel': 'feishu',
        'source_open_id': compact_text(message.get('open_id')),
        'source_chat_id': compact_text(message.get('chat_id')),
        'source_message_id': compact_text(message.get('message_id')),
        'source_text': compact_text(message.get('text')),
        'confirmed_at': confirmed_at,
    }


def feishu_agent_execute_tool(call: dict[str, Any], message: dict[str, Any], session: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    name = compact_text(call.get('name'))
    args = call.get('arguments') if isinstance(call.get('arguments'), dict) else {}
    text = compact_text(args.get('text') or message.get('text') or '')
    files = message.get('files') or []

    if name == 'chat_general':
        return {'tool': name, 'ok': True, 'reply': feishu_agent_chat_reply(text, session, config)}

    if name == 'cancel_pending':
        session.update({'active_skill': '', 'seed_text': '', 'slots': {}, 'missing': [], 'pending_confirm': {}, 'last_question': text})
        return {'tool': name, 'ok': True, 'reply': '已取消当前待办。你可以继续发新内容，我会重新理解。'}

    if name == 'save_pending_record':
        pending = session.get('pending_confirm') if isinstance(session.get('pending_confirm'), dict) else {}
        if pending.get('type') != 'market_record':
            return {'tool': name, 'ok': False, 'error': '当前没有可保存的草稿。'}
        raw_record = dict(pending.get('record') or {})
        source_fields = feishu_record_source_fields(message, confirmed_at=datetime.now().isoformat(timespec='seconds'))
        for key, value in source_fields.items():
            if key in {'source', 'confirmed_at'} or not compact_text(raw_record.get(key)):
                raw_record[key] = value
        if compact_text(raw_record.get('raw_text')):
            raw_record['source_text'] = compact_text(raw_record.get('raw_text'))
        record_kind = str(pending.get('kind') or '')
        record = save_market_record(record_kind, raw_record)
        session.update({'active_skill': '', 'seed_text': '', 'slots': {}, 'missing': [], 'pending_confirm': {}, 'last_tool_result': {'tool': name, 'saved_record_kind': record_kind, 'saved_record_id': record.get('id'), 'saved_at': datetime.now().isoformat(timespec='seconds')}})
        return {'tool': name, 'ok': True, 'saved_record_kind': record_kind, 'saved_record_id': record.get('id'), 'record': record, 'reply': f"已保存到市场情报台账，记录 ID：{record.get('id')}"}

    if name == 'check_last_saved_record':
        last_result = session.get('last_tool_result') if isinstance(session.get('last_tool_result'), dict) else {}
        record_id = compact_text(args.get('record_id') or last_result.get('saved_record_id'))
        kind = compact_text(args.get('kind') or last_result.get('saved_record_kind'))
        if kind not in {'cargo', 'newbuilding'} or not record_id:
            return {'tool': name, 'ok': False, 'error': '我这里没有最近保存记录的上下文。'}
        path = market_record_path(kind, record_id)
        if not path.exists():
            return {'tool': name, 'ok': False, 'kind': kind, 'record_id': record_id, 'error': f'最近记录 {record_id} 未在台账中找到，可能已被删除或数据目录已更换。'}
        record = load_market_record(kind, record_id)
        label = '商机货盘' if kind == 'cargo' else '新造船信息'
        title = record.get('cargo_name') or record.get('ship_name') or record.get('route') or record.get('owner') or '未命名'
        return {'tool': name, 'ok': True, 'kind': kind, 'record_id': record_id, 'record': record, 'reply': f'刚才那条{label}已经保存，记录 ID：{record_id}，摘要：{title}。'}

    if name in {'extract_cargo_opportunity', 'extract_newbuilding_info'}:
        kind = 'cargo' if name == 'extract_cargo_opportunity' else 'newbuilding'
        skill = 'cargo_opportunity' if kind == 'cargo' else 'newbuilding_info'
        record = extract_market_record(kind, text)
        missing = feishu_skill_missing_fields(skill, text, record)
        session.update(
            {
                'active_skill': skill,
                'seed_text': text,
                'slots': record,
                'missing': missing,
                'pending_confirm': {} if missing else {'type': 'market_record', 'kind': kind, 'record': record, 'created_at': datetime.now().isoformat(timespec='seconds')},
                'last_question': text,
            }
        )
        return {'tool': name, 'ok': True, 'kind': kind, 'record': record, 'missing': missing, 'reply': feishu_skill_preview_record(skill, record) if not missing else feishu_skill_next_question(skill, missing)}

    if name == 'update_pending_record':
        pending = session.get('pending_confirm') if isinstance(session.get('pending_confirm'), dict) else {}
        current = pending.get('record') if isinstance(pending.get('record'), dict) else session.get('slots', {})
        kind = str(pending.get('kind') or current.get('kind') or ('newbuilding' if session.get('active_skill') == 'newbuilding_info' else 'cargo'))
        if not isinstance(current, dict) or not current:
            return {'tool': name, 'ok': False, 'error': '当前没有可修改的草稿。'}
        record = feishu_agent_merge_record(kind, current, text)
        skill = 'cargo_opportunity' if kind == 'cargo' else 'newbuilding_info'
        missing = feishu_skill_missing_fields(skill, text, record)
        session.update(
            {
                'active_skill': skill,
                'seed_text': f"{session.get('seed_text')}\n{text}".strip(),
                'slots': record,
                'missing': missing,
                'pending_confirm': {} if missing else {'type': 'market_record', 'kind': kind, 'record': record, 'created_at': datetime.now().isoformat(timespec='seconds')},
                'last_question': text,
            }
        )
        return {'tool': name, 'ok': True, 'kind': kind, 'record': record, 'missing': missing, 'reply': feishu_skill_preview_record(skill, record) if not missing else feishu_skill_next_question(skill, missing)}

    if name == 'search_project':
        query = compact_text(args.get('query') or args.get('q') or text)
        items, _ = collect_project_items(q=query)
        return {'tool': name, 'ok': True, 'query': query, 'items': items[:5], 'reply': format_project_search(items)}

    if name == 'search_market_records':
        query = compact_text(args.get('query') or args.get('q') or text)
        kind = compact_text(args.get('kind'))
        if kind not in {'cargo', 'newbuilding'}:
            kind = 'newbuilding' if any(word in query for word in ('新造船', '船厂', '船东', 'DWT', '交付')) else 'cargo'
        items, stats = list_market_records(
            kind,
            q=query,
            segment=compact_text(args.get('segment')),
            board_type=compact_text(args.get('board_type')),
            status=compact_text(args.get('status')),
            stage=compact_text(args.get('stage')),
        )
        start_date = compact_text(args.get('start_date'))
        end_date = compact_text(args.get('end_date'))
        if start_date or end_date:
            date_key = 'cargo_date' if kind == 'cargo' else 'update_date'
            items = [
                item
                for item in items
                if (not start_date or compact_text(item.get(date_key)) >= start_date)
                and (not end_date or compact_text(item.get(date_key)) <= end_date)
            ]
        if not items and kind == 'cargo':
            query_parts = []
            cargo_record = guess_market_cargo(query)
            for key in ('cargo_name', 'load_port', 'discharge_port', 'cargo_owner'):
                value = compact_text(cargo_record.get(key))
                if value:
                    query_parts.append(value)
            for part in query_parts:
                items, stats = list_market_records(kind, q=part)
                if items:
                    break
        label = '商机货盘' if kind == 'cargo' else '新造船'
        if not items:
            return {'tool': name, 'ok': True, 'kind': kind, 'query': query, 'items': [], 'stats': stats, 'reply': f'没有查到匹配的{label}记录，可以换个货品、航线、船厂或时间关键词。'}
        filter_parts = [part for part in [query, compact_text(args.get('status') or args.get('stage')), start_date and f'{start_date}起', end_date and f'{end_date}止'] if part]
        lines = [f"我查了系统里的{label}台账（条件：{' / '.join(filter_parts) or '全部'}），共找到 {len(items)} 条，先汇总前 5 条："]
        for index, item in enumerate(items[:5], 1):
            if kind == 'cargo':
                lines.append(f"{index}. {item.get('cargo_name') or '-'} | {item.get('tonnage') or '-'} | {item.get('load_port') or '-'} -> {item.get('discharge_port') or '-'} | {item.get('status') or '-'} | ID：{item.get('id')}")
            else:
                lines.append(f"{index}. {item.get('ship_name') or item.get('owner') or '-'} | {item.get('shipyard') or '-'} | {item.get('dwt') or '-'} | {item.get('stage') or '-'} | ID：{item.get('id')}")
        lines.append('这些结果来自当前系统台账；没有出现在台账里的信息我不会编造。')
        return {'tool': name, 'ok': True, 'kind': kind, 'query': query, 'items': items[:5], 'stats': stats, 'reply': '\n'.join(lines)}

    if name == 'generate_market_report':
        kind = compact_text(args.get('kind') or ('newbuilding' if '新造船' in text else 'cargo'))
        if kind not in {'cargo', 'newbuilding'}:
            kind = 'cargo'
        period = compact_text(args.get('period'))
        start_date = compact_text(args.get('start_date'))
        end_date = compact_text(args.get('end_date'))
        if not (period and start_date and end_date):
            period, start_date, end_date = feishu_guess_period(text)
        if not (period and start_date and end_date):
            return {'tool': name, 'ok': False, 'error': '请补充报告时间范围，例如本月、上月，或 2026-06-01 到 2026-06-30。'}
        report = build_market_report(MarketReportPayload(kind=kind, period=period, start_date=start_date, end_date=end_date, filters={}))
        session.update({'active_skill': 'cargo_report' if kind == 'cargo' else 'newbuilding_report', 'last_result': report, 'last_question': text})
        return {'tool': name, 'ok': True, 'kind': kind, 'report': report, 'reply': feishu_skill_preview_report('cargo_report' if kind == 'cargo' else 'newbuilding_report', report)}

    if name == 'answer_bid_question':
        if not session.get('last_result'):
            return {'tool': name, 'ok': False, 'error': '当前没有最近一次标书解析结果，请先上传并解析标书。'}
        answer = feishu_answer_bid_question(session['last_result'], text, config)
        session['last_question'] = text
        return {'tool': name, 'ok': True, 'reply': answer}

    if name == 'save_bid_project':
        result = session.get('last_result') if isinstance(session.get('last_result'), dict) else {}
        if not result:
            return {'tool': name, 'ok': False, 'error': '当前没有可保存的标书解析结果，请先上传并解析标书。'}
        summary = result.get('document_summary') if isinstance(result.get('document_summary'), dict) else {}
        project_id = uuid.uuid4().hex
        source_file_name = compact_text(args.get('source_file_name') or session.get('last_file_name') or message.get('file_name'))
        source_note = '飞书确认保存'
        if message.get('message_id'):
            source_note += f"；message_id={message.get('message_id')}"
        project = save_project(
            project_id,
            ProjectPayload(
                title=compact_text(summary.get('project_name')) or '飞书标书解析项目',
                source_file_name=source_file_name,
                source_channel='feishu',
                source_open_id=compact_text(message.get('open_id')),
                source_chat_id=compact_text(message.get('chat_id')),
                source_message_id=compact_text(message.get('message_id')),
                source_text=compact_text(session.get('last_parse_text') or session.get('last_question') or message.get('text')),
                confirmed_at=datetime.now().isoformat(timespec='seconds'),
                sheet_name=str(datetime.now().year),
                result=result,
                follow_up={'bid_status': '待跟进', 'award_status': '未知', 'information_source': '飞书', 'register_year': str(datetime.now().year)},
                timeline=[{'date': today_text(), 'type': 'parse', 'note': source_note}],
            ),
        )
        session.update({'last_tool_result': {'tool': name, 'saved_project_id': project_id, 'saved_at': datetime.now().isoformat(timespec='seconds')}})
        return {'tool': name, 'ok': True, 'saved_project_id': project_id, 'project': project, 'reply': f"已保存到历史项目台账，项目 ID：{project_id}"}

    if name == 'parse_bid_file':
        if not files:
            return {'tool': name, 'ok': False, 'error': '请在飞书里同时上传 PDF/DOC/DOCX/XLS/XLSX 标书文件。'}
        file_info = files[0]
        temp_path: Path | None = None
        try:
            temp_path = feishu_download_file(config, message.get('message_id', ''), file_info.get('file_key', ''), file_info.get('file_name', ''))
            suffix = Path(file_info.get('file_name') or temp_path.name).suffix.lower() or temp_path.suffix.lower()
            if suffix not in {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}:
                raise HTTPException(status_code=400, detail='当前只支持 PDF、DOC、DOCX、XLS、XLSX 标书文件。')
            document_text = extract_document_text(temp_path, suffix)
            if len(document_text.strip()) < 80:
                raise HTTPException(status_code=400, detail='当前版本只支持可提取文本的文件，扫描件暂不支持。')
            result = parse_ai_document(document_text, require_config())
            session.update({'active_skill': 'bid_parse', 'last_result': result, 'last_question': text, 'last_parse_text': text, 'last_file_name': file_info.get('file_name', ''), 'pending_confirm': {}})
            return {'tool': name, 'ok': True, 'result': result, 'reply': feishu_skill_preview_bid(result)}
        except HTTPException as exc:
            return {'tool': name, 'ok': False, 'error': compact_text(exc.detail)}
        finally:
            if temp_path:
                temp_path.unlink(missing_ok=True)

    return {'tool': name, 'ok': False, 'error': '工具不在白名单内。'}


def feishu_skill_preview_bid(result: dict[str, Any]) -> str:
    summary = result.get('document_summary') or {}
    analysis = result.get('analysis') or {}
    review = result.get('match_review') or []
    return '\n'.join(
        [
            '标书解析完成：',
            f"项目名称：{summary.get('project_name') or '-'}",
            f"招标编号：{summary.get('bid_no') or '-'}",
            f"招标人：{summary.get('tenderer') or '-'}",
            f"投标截止时间：{summary.get('bid_deadline') or '-'}",
            f"开标时间：{summary.get('open_time') or '-'}",
            f"保证金：{summary.get('deposit_amount') or '-'}",
            f"资格要求：{compact_paragraph(summary.get('qualification_requirements'))[:500] or '-'}",
            f"风险提示：{', '.join(str(item.get('reason') or item) for item in review[:3]) or '-'}",
            f"解析结论：{analysis.get('summary') or '-'}",
            '下一步：如果要继续追问标书内容，我可以基于这次解析结果继续回答。',
        ]
    )


def feishu_skill_preview_record(skill: str, record: dict[str, Any]) -> str:
    if skill == 'cargo_opportunity':
        lines = [
            '我已识别到一条商机货盘草稿：',
            f"货名：{record.get('cargo_name') or '-'}",
            f"吨数：{record.get('tonnage') or '-'}",
            f"装港：{record.get('load_port') or '-'}",
            f"卸港：{record.get('discharge_port') or '-'}",
            f"装载期：{record.get('laycan') or '-'}",
            f"货主：{record.get('cargo_owner') or '-'}",
            f"航线：{record.get('route') or '-'}",
            f"成交状态：{record.get('status') or '-'}",
            f"最终成交价：{record.get('deal_price') or record.get('final_price') or '-'}",
        ]
    else:
        lines = [
            '我已识别到一条新造船草稿：',
            f"船名：{record.get('ship_name') or '-'}",
            f"船厂：{record.get('shipyard') or '-'}",
            f"船东：{record.get('owner') or '-'}",
            f"DWT：{record.get('dwt') or '-'}",
            f"阶段：{record.get('stage') or '-'}",
            f"预计交付：{record.get('delivery_time') or '-'}",
        ]
    lines.append('下一步：回复“确认保存”写入台账，或者回复“取消”放弃。')
    return '\n'.join(lines)


def feishu_skill_preview_report(skill: str, report: dict[str, Any]) -> str:
    title = '商机市场分析报告' if skill == 'cargo_report' else '新造船市场分析报告'
    lines = [f'已生成{title}：', f"摘要：{report.get('summary') or '-'}"]
    for label, key in [('关键发现', 'key_findings'), ('风险提示', 'risks'), ('经营建议', 'recommendations')]:
        values = [str(item) for item in (report.get(key) or []) if str(item).strip()]
        if values:
            lines.append(f'{label}：')
            lines.extend(f'- {item}' for item in values[:5])
    lines.append('下一步：如果要导出 Word / Excel，请到网页端的“市场情报”里操作。')
    return '\n'.join(lines)


def feishu_answer_bid_question(result: dict[str, Any], question: str, config: dict[str, Any]) -> str:
    summary = result.get('document_summary') or {}
    q = compact_text(question)
    if any(word in q for word in ('截止', '投标截止')):
        return f"本次标书里能直接确认的投标截止时间是：{summary.get('bid_deadline') or '-'}。"
    if any(word in q for word in ('开标', '开标时间')):
        return f"本次标书里能直接确认的开标时间是：{summary.get('open_time') or '-'}。"
    if '保证金' in q:
        return f"本次标书里能直接确认的保证金是：{summary.get('deposit_amount') or '-'}。"
    if any(word in q for word in ('资格', '资质', '条件')):
        return f"本次标书里能直接确认的资格要求是：{compact_paragraph(summary.get('qualification_requirements'))[:500] or '-'}。"
    if any(word in q for word in ('评标', '评分')):
        return f"本次标书里能直接确认的评标方法是：{summary.get('evaluation_method') or '-'}。"
    try:
        answer = call_chat_completion(
            config,
            system_prompt=CHAT_SYSTEM_PROMPT,
            user_prompt=build_chat_prompt({'result': result}, question),
            temperature=0.2,
        )
        return f"{answer.strip()}\n下一步：如果还要继续问这份标书，直接接着发问题就行。"
    except HTTPException:
        return '我已保留这次标书解析结果。你可以继续问我“投标截止时间 / 保证金 / 资格要求 / 评标标准”。'


def feishu_skill_route(text: str, files: list[dict[str, str]], session: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    compact = compact_text(text)
    if feishu_skill_is_cancel(compact):
        return {'skill': 'help', 'action': 'cancel', 'slots': {}, 'missing': [], 'next_question': ''}
    if feishu_skill_is_confirm(compact):
        return {'skill': compact_text(session.get('active_skill') or 'help'), 'action': 'confirm', 'slots': {}, 'missing': [], 'next_question': ''}
    if feishu_skill_is_help(compact) and not session.get('pending_confirm'):
        return {'skill': 'help', 'action': 'help', 'slots': {}, 'missing': [], 'next_question': ''}

    ai_route = feishu_skill_route_from_ai(text, files, session, config)
    if ai_route:
        return ai_route

    skill = feishu_skill_from_text(text, files, session)
    if skill == 'project_search':
        query = compact_text(re.sub(r'^(查询|查找|搜索|搜索一下|查询项目)', '', compact).strip())
        if not query and session.get('active_skill') == 'project_search':
            query = compact_text(session.get('slots', {}).get('query'))
        if not query:
            return {'skill': skill, 'action': 'clarify', 'slots': {}, 'missing': ['query'], 'next_question': feishu_skill_next_question(skill, ['query'])}
        return {'skill': skill, 'action': 'execute', 'slots': {'query': query}, 'missing': [], 'next_question': ''}
    if skill in {'cargo_report', 'newbuilding_report'}:
        period, start_date, end_date = feishu_guess_period(text)
        if not (period and start_date and end_date) or (period == 'custom' and not (start_date and end_date)):
            return {'skill': skill, 'action': 'clarify', 'slots': {}, 'missing': ['period'], 'next_question': feishu_skill_next_question(skill, ['period'])}
        return {'skill': skill, 'action': 'execute', 'slots': {'period': period, 'start_date': start_date, 'end_date': end_date}, 'missing': [], 'next_question': ''}
    if skill == 'bid_parse':
        if not files and not session.get('last_result'):
            return {'skill': skill, 'action': 'clarify', 'slots': {}, 'missing': ['file'], 'next_question': feishu_skill_next_question(skill, ['file'])}
        return {'skill': skill, 'action': 'execute', 'slots': {}, 'missing': [], 'next_question': ''}
    if skill in {'cargo_opportunity', 'newbuilding_info'}:
        return {'skill': skill, 'action': 'execute', 'slots': {}, 'missing': [], 'next_question': ''}
    return {'skill': 'help', 'action': 'help', 'slots': {}, 'missing': [], 'next_question': ''}


def feishu_skill_execute(route: dict[str, Any], message: dict[str, Any], config: dict[str, Any]) -> str:
    open_id = message.get('open_id', '')
    chat_id = message.get('chat_id', '')
    text = compact_text(message.get('text') or '')
    files = message.get('files') or []
    session = feishu_dialog_load(open_id, chat_id)
    skill = compact_text(route.get('skill'))
    action = compact_text(route.get('action'))
    slots = route.get('slots') if isinstance(route.get('slots'), dict) else {}
    missing = [compact_text(item) for item in (route.get('missing') or []) if compact_text(item)]
    next_question = compact_text(route.get('next_question'))

    if action == 'cancel':
        feishu_dialog_clear(open_id, chat_id)
        return '已取消当前待办。你可以继续发新内容，我会重新识别。'

    if action == 'help':
        feishu_dialog_clear(open_id, chat_id)
        return feishu_skill_catalog_text()

    if action == 'confirm':
        pending = session.get('pending_confirm') if isinstance(session.get('pending_confirm'), dict) else {}
        if pending.get('type') != 'market_record':
            return '当前没有待保存内容。你可以直接发新的商机货盘、新造船信息或标书文件。'
        record = save_market_record(str(pending.get('kind') or ''), pending.get('record') or {})
        feishu_dialog_clear(open_id, chat_id)
        return f"已保存到市场情报台账，记录 ID：{record.get('id')}\n下一步：可以继续发新的商机或新造船信息。"

    if skill == 'project_search':
        query = compact_text(slots.get('query') or re.sub(r'^(查询|查找|搜索|搜索一下|查询项目)', '', text))
        items, _ = collect_project_items(q=query)
        if not items:
            session.update({'active_skill': 'project_search', 'slots': {'query': query}, 'last_question': next_question or feishu_skill_next_question('project_search', ['query'])})
            feishu_dialog_save(open_id, chat_id, session)
            return f'没有找到匹配的历史项目。请换一个关键词，比如项目名、招标编号或更短的关键词：{query or "（空）"}'
        lines = ['查询到以下历史项目：']
        for index, item in enumerate(items[:5], start=1):
            lines.append(
                f"{index}. {item.get('title') or item.get('project_name') or '未命名项目'}\n"
                f"   招标编号：{item.get('bid_no') or '-'}\n"
                f"   招标人：{item.get('tenderer') or '-'}\n"
                f"   投标状态：{item.get('bid_status') or '-'}，中标状态：{item.get('award_status') or '-'}"
            )
        lines.append('下一步：如果要看详情，请在网页端打开历史台账；如果要查其他项目，继续发送“查询 + 关键词”。')
        feishu_dialog_clear(open_id, chat_id)
        return '\n'.join(lines)

    if skill == 'bid_parse':
        if not files and not session.get('last_result'):
            session.update({'active_skill': 'bid_parse', 'last_question': next_question or feishu_skill_next_question('bid_parse', ['file'])})
            feishu_dialog_save(open_id, chat_id, session)
            return next_question or feishu_skill_next_question('bid_parse', ['file'])
        if not files and session.get('last_result') and text:
            answer = feishu_answer_bid_question(session['last_result'], text, config)
            session['last_question'] = text
            feishu_dialog_save(open_id, chat_id, session)
            return answer
        file_info = files[0]
        temp_path = feishu_download_file(config, message.get('message_id', ''), file_info.get('file_key', ''), file_info.get('file_name', ''))
        try:
            suffix = Path(file_info.get('file_name') or temp_path.name).suffix.lower() or temp_path.suffix.lower()
            if suffix not in {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}:
                raise HTTPException(status_code=400, detail='当前只支持 PDF、DOC、DOCX、XLS、XLSX 标书文件。')
            document_text = extract_document_text(temp_path, suffix)
            if len(document_text.strip()) < 80:
                raise HTTPException(status_code=400, detail='当前版本只支持可提取文本的文件，扫描件暂不支持。')
            result = parse_ai_document(document_text, require_config())
            session.update(
                {
                    'active_skill': 'bid_parse',
                    'seed_text': text,
                    'slots': {},
                    'missing': [],
                    'pending_confirm': {},
                    'last_question': '',
                    'retry_count': 0,
                    'last_result': result,
                }
            )
            feishu_dialog_save(open_id, chat_id, session)
            return feishu_skill_preview_bid(result)
        finally:
            temp_path.unlink(missing_ok=True)

    if skill in {'cargo_opportunity', 'newbuilding_info'}:
        record_kind = 'cargo' if skill == 'cargo_opportunity' else 'newbuilding'
        merged_text = text
        if session.get('active_skill') == skill and compact_text(session.get('seed_text')):
            merged_text = f"{session.get('seed_text')}\n{text}".strip()
        record = extract_market_record(record_kind, merged_text)
        missing = feishu_skill_missing_fields(skill, merged_text, record)
        if missing:
            session.update(
                {
                    'active_skill': skill,
                    'seed_text': merged_text,
                    'slots': record,
                    'missing': missing,
                    'pending_confirm': {},
                    'last_question': next_question or feishu_skill_next_question(skill, missing),
                    'retry_count': 0,
                }
            )
            feishu_dialog_save(open_id, chat_id, session)
            return '\n'.join([feishu_skill_reply_prefix(skill, 'clarify', missing), next_question or feishu_skill_next_question(skill, missing)])
        session.update(
            {
                'active_skill': skill,
                'seed_text': merged_text,
                'slots': record,
                'missing': [],
                'pending_confirm': {'type': 'market_record', 'kind': record_kind, 'record': record, 'created_at': datetime.now().isoformat(timespec='seconds')},
                'last_question': '回复“确认保存”即可写入台账，或回复“取消”放弃。',
                'retry_count': 0,
            }
        )
        feishu_dialog_save(open_id, chat_id, session)
        return feishu_skill_preview_record(skill, record)

    if skill in {'cargo_report', 'newbuilding_report'}:
        period = compact_text(slots.get('period'))
        start_date = compact_text(slots.get('start_date'))
        end_date = compact_text(slots.get('end_date'))
        if not (period and start_date and end_date):
            period, start_date, end_date = feishu_guess_period(text)
        if not (period and start_date and end_date):
            session.update({'active_skill': skill, 'last_question': next_question or feishu_skill_next_question(skill, ['period'])})
            feishu_dialog_save(open_id, chat_id, session)
            return '\n'.join([feishu_skill_reply_prefix(skill, 'clarify', ['period']), next_question or feishu_skill_next_question(skill, ['period'])])
        report = build_market_report(
            MarketReportPayload(
                kind='cargo' if skill == 'cargo_report' else 'newbuilding',
                period=period,
                start_date=start_date,
                end_date=end_date,
                filters={},
            )
        )
        session.update(
            {
                'active_skill': skill,
                'seed_text': text,
                'slots': {'period': period, 'start_date': start_date, 'end_date': end_date},
                'missing': [],
                'pending_confirm': {},
                'last_question': '',
                'retry_count': 0,
                'last_result': report,
            }
        )
        feishu_dialog_save(open_id, chat_id, session)
        return feishu_skill_preview_report(skill, report)

    if session.get('active_skill') == 'bid_parse' and session.get('last_result') and text and not files:
        answer = feishu_answer_bid_question(session['last_result'], text, config)
        session['last_question'] = text
        feishu_dialog_save(open_id, chat_id, session)
        return answer

    feishu_dialog_clear(open_id, chat_id)
    return feishu_skill_catalog_text()


def handle_feishu_message(message: dict[str, Any], config: dict[str, Any]) -> str:
    open_id = message.get('open_id', '')
    chat_id = message.get('chat_id', '')
    text = compact_text(message.get('text') or '')
    files = message.get('files') or []
    if not feishu_message_allowed(config, message):
        reply = '你当前没有权限使用这个飞书 Agent，请联系管理员把你的 open_id 或群 chat_id 加入允许列表。'
        write_feishu_agent_log(message, {}, [], reply, error='unauthorized_feishu_user')
        return reply
    session = feishu_dialog_load(open_id, chat_id)
    try:
        decision = feishu_agent_decide(text, files, session, config)
    except HTTPException as exc:
        if exc.status_code == 400 and 'AI 暂不可用' in str(exc.detail):
            reply = compact_text(exc.detail)
            write_feishu_agent_log(message, {}, [], reply, error=reply)
            return reply
        raise

    tool_results: list[dict[str, Any]] = []
    for call in decision.get('tool_calls') or [{'name': 'chat_general', 'arguments': {}}]:
        result = feishu_agent_execute_tool(call, message, session, config)
        tool_results.append(result)
        session['last_tool_result'] = result

    pending_update = decision.get('pending_update') if isinstance(decision.get('pending_update'), dict) else {}
    if pending_update:
        for key in ('agent_summary', 'last_question'):
            if key in pending_update:
                session[key] = compact_text(pending_update.get(key))
    if text:
        session['last_question'] = text
    feishu_dialog_save(open_id, chat_id, session)

    explicit_reply = compact_text(decision.get('reply'))
    if explicit_reply and not any(result.get('tool') != 'chat_general' for result in tool_results):
        write_feishu_agent_log(message, decision, tool_results, explicit_reply)
        return explicit_reply
    try:
        reply = feishu_agent_chat_reply(text, session, config, tool_results)
        if reply:
            write_feishu_agent_log(message, decision, tool_results, reply)
            return reply
    except HTTPException:
        pass
    for result in tool_results:
        if result.get('reply'):
            reply = sanitize_feishu_reply(result.get('reply'))
            write_feishu_agent_log(message, decision, tool_results, reply)
            return reply
        if result.get('error'):
            reply = f"处理时遇到问题：{compact_text(result.get('error'))}"
            write_feishu_agent_log(message, decision, tool_results, reply, error=compact_text(result.get('error')))
            return reply
    reply = 'AI 已处理这条消息，但没有生成可发送的回复。'
    write_feishu_agent_log(message, decision, tool_results, reply, error='empty_reply')
    return reply


def process_feishu_event(message: dict[str, Any], config: dict[str, Any]) -> None:
    try:
        reply = handle_feishu_message(message, config)
    except HTTPException as exc:
        reply = f'处理时遇到问题：{compact_text(exc.detail)}\n下一步：请补充缺失信息，或者回复“取消”重新开始。'
    except Exception as exc:
        print(f'Feishu skill error: {exc}')
        reply = '处理时遇到问题：系统暂时没法完成这次请求。\n下一步：请稍后再试，或者换一种说法。'
    try:
        feishu_reply_message(config, message.get('message_id', ''), sanitize_feishu_reply(reply))
    except Exception as exc:
        print(f'Feishu reply failed: {exc}')

def start_feishu_ws_client() -> None:
    global FEISHU_WS_THREAD, FEISHU_WS_STARTED_FOR
    config = load_config()
    if not config.get('feishu_enabled') or config.get('feishu_receive_mode') != 'ws':
        return
    app_id = config.get('feishu_app_id', '')
    app_secret = config.get('feishu_app_secret', '')
    if not app_id or not app_secret:
        print('Feishu WS disabled: App ID/App Secret is missing.')
        return
    if lark is None or P2ImMessageReceiveV1 is None:
        print('Feishu WS disabled: lark-oapi is not installed.')
        return
    started_key = hashlib.sha256(f'{app_id}:{app_secret}'.encode('utf-8')).hexdigest()
    if FEISHU_WS_THREAD and FEISHU_WS_THREAD.is_alive() and FEISHU_WS_STARTED_FOR == started_key:
        return

    def run() -> None:
        try:
            if lark_ws_client is not None:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                lark_ws_client.loop = loop
            handler = (
                lark.EventDispatcherHandler.builder(
                    config.get('feishu_encrypt_key', ''),
                    config.get('feishu_verification_token', ''),
                )
                .register_p2_im_message_receive_v1(handle_feishu_ws_message)
                .build()
            )
            client = lark.ws.Client(app_id, app_secret, event_handler=handler)
            print('Feishu WS client starting.')
            client.start()
        except Exception as exc:
            print(f'Feishu WS client stopped: {exc}')

    FEISHU_WS_STARTED_FOR = started_key
    FEISHU_WS_THREAD = threading.Thread(target=run, name='feishu-ws-client', daemon=True)
    FEISHU_WS_THREAD.start()


def parse_excel_date(value: str) -> Any:
    text = compact_text(value)
    if not text:
        return ''
    for fmt in (
        '%Y-%m-%d',
        '%Y/%m/%d',
        '%Y.%m.%d',
        '%Y年%m月%d日',
        '%Y-%m-%d %H:%M',
        '%Y/%m/%d %H:%M',
        '%Y年%m月%d日 %H:%M',
    ):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return text


def parse_excel_amount(value: str) -> Any:
    text = compact_text(value)
    if not text:
        return ''
    plain = text.replace(',', '').replace('，', '')
    if re.fullmatch(r'-?\d+(?:\.\d+)?', plain):
        return float(plain) if '.' in plain else int(plain)
    match = re.fullmatch(r'(-?\d+(?:\.\d+)?)\s*万(?:元)?', plain)
    if match:
        number = float(match.group(1)) * 10000
        return int(number) if number.is_integer() else number
    match = re.fullmatch(r'(-?\d+(?:\.\d+)?)\s*元', plain)
    if match:
        number = float(match.group(1))
        return int(number) if number.is_integer() else number
    return text


def find_register_start_row(sheet: Any) -> int:
    row = 3
    while sheet[f'A{row}'].value not in (None, ''):
        row += 1
    return row


def get_register_template_path() -> Path:
    if not PACKAGE_REGISTER_TEMPLATE_PATH.exists():
        raise HTTPException(
            status_code=500,
            detail=f'缺少登记模板文件：{PACKAGE_REGISTER_TEMPLATE_PATH.name}，请放到 模板文件 目录。',
        )
    return PACKAGE_REGISTER_TEMPLATE_PATH


def get_extraction_template_path() -> Path:
    if not PACKAGE_EXTRACTION_TEMPLATE_PATH.exists():
        raise HTTPException(
            status_code=500,
            detail=f'缺少摘取模板文件：{PACKAGE_EXTRACTION_TEMPLATE_PATH.name}，请放到 模板文件 目录。',
        )
    return PACKAGE_EXTRACTION_TEMPLATE_PATH


def ensure_register_sheet(workbook: Any, sheet_name: str) -> Any:
    if sheet_name in workbook.sheetnames:
        return workbook[sheet_name]
    source_sheet = workbook[workbook.sheetnames[0]]
    cloned = workbook.copy_worksheet(source_sheet)
    cloned.title = sheet_name
    return cloned


def apply_register_row_style(sheet: Any, row_index: int) -> None:
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    for column in range(1, 15):
        cell = sheet.cell(row_index, column)
        cell.border = border
        cell.alignment = center
    sheet[f'B{row_index}'].number_format = 'yyyy/m/d'
    sheet[f'G{row_index}'].number_format = 'yyyy/m/d h:mm'
    if isinstance(sheet[f'H{row_index}'].value, (int, float)):
        sheet[f'H{row_index}'].number_format = '#,##0.00'


def project_path(project_id: str) -> Path:
    if not re.fullmatch(r'[a-f0-9]{32}', project_id):
        raise HTTPException(status_code=400, detail='项目 ID 无效。')
    return PROJECTS_DIR / f'{project_id}.json'


def project_meta(project: dict[str, Any]) -> dict[str, Any]:
    summary = project.get('result', {}).get('document_summary', {})
    follow_up = project.get('follow_up', {})
    return {
        'project_id': project['project_id'],
        'title': project.get('title', ''),
        'source_file_name': project.get('source_file_name', ''),
        'source_channel': project.get('source_channel', ''),
        'source_message_id': project.get('source_message_id', ''),
        'register_mode': project.get('register_mode', 'packages'),
        'sheet_name': project.get('sheet_name', ''),
        'created_at': project.get('created_at', ''),
        'updated_at': project.get('updated_at', ''),
        'project_name': summary.get('project_name', ''),
        'bid_no': summary.get('bid_no', ''),
        'tenderer': summary.get('tenderer', ''),
        'register_year': follow_up.get('register_year', '') or project.get('sheet_name', ''),
        'bid_status': follow_up.get('bid_status', '待跟进'),
        'award_status': follow_up.get('award_status', '未知'),
        'our_quote_count': len(project.get('our_quotes', [])),
        'competitor_quote_count': len(project.get('competitor_quotes', [])),
    }


def read_project_file(path: Path) -> dict[str, Any] | None:
    try:
        raw = json.loads(path.read_text(encoding='utf-8'))
    except (OSError, json.JSONDecodeError):
        return None
    project_id = str(raw.get('project_id') or path.stem)
    sheet_name = compact_text(raw.get('sheet_name'))
    result = normalize_result_payload(raw.get('result') or {})
    follow_up = normalize_follow_up(raw.get('follow_up'), sheet_name=sheet_name)
    our_quotes = normalize_quote_rows(raw.get('our_quotes'), default_company='我司')
    competitor_quotes = normalize_quote_rows(raw.get('competitor_quotes'))
    timeline = normalize_timeline_rows(raw.get('timeline'))
    if not timeline and result.get('document_summary', {}).get('project_name'):
        timeline = build_auto_timeline(result, compact_text(raw.get('source_file_name')))
    return {
        'project_id': project_id,
        'title': compact_text(raw.get('title')),
        'source_file_name': compact_text(raw.get('source_file_name')),
        'source_channel': compact_text(raw.get('source_channel')),
        'source_open_id': compact_text(raw.get('source_open_id')),
        'source_chat_id': compact_text(raw.get('source_chat_id')),
        'source_message_id': compact_text(raw.get('source_message_id')),
        'source_text': compact_text(raw.get('source_text')),
        'confirmed_at': compact_text(raw.get('confirmed_at')),
        'register_mode': compact_text(raw.get('register_mode') or 'packages') or 'packages',
        'sheet_name': sheet_name,
        'created_at': compact_text(raw.get('created_at')),
        'updated_at': compact_text(raw.get('updated_at')),
        'result': result,
        'follow_up': follow_up,
        'our_quotes': our_quotes,
        'competitor_quotes': competitor_quotes,
        'timeline': timeline,
    }


def load_project(project_id: str) -> dict[str, Any]:
    path = project_path(project_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail='项目不存在。')
    project = read_project_file(path)
    if not project:
        raise HTTPException(status_code=500, detail='项目文件损坏，无法读取。')
    return project


def save_project(project_id: str, payload: ProjectPayload, *, created_at: str | None = None) -> dict[str, Any]:
    now = datetime.now().isoformat(timespec='seconds')
    result = normalize_result_payload(payload.result)
    title = compact_text(payload.title) or result['document_summary']['project_name'] or '未命名项目'
    sheet_name = compact_text(payload.sheet_name)
    follow_up = normalize_follow_up(payload.follow_up, sheet_name=sheet_name)
    our_quotes = normalize_quote_rows(payload.our_quotes, default_company='我司')
    competitor_quotes = normalize_quote_rows(payload.competitor_quotes)
    timeline = normalize_timeline_rows(payload.timeline)
    if not timeline and result.get('document_summary', {}).get('project_name'):
        timeline = build_auto_timeline(result, compact_text(payload.source_file_name))
    project = {
        'project_id': project_id,
        'title': title,
        'source_file_name': compact_text(payload.source_file_name),
        'source_channel': compact_text(payload.source_channel),
        'source_open_id': compact_text(payload.source_open_id),
        'source_chat_id': compact_text(payload.source_chat_id),
        'source_message_id': compact_text(payload.source_message_id),
        'source_text': compact_text(payload.source_text),
        'confirmed_at': compact_text(payload.confirmed_at),
        'register_mode': payload.register_mode if payload.register_mode in {'packages', 'document'} else 'packages',
        'sheet_name': sheet_name,
        'created_at': created_at or now,
        'updated_at': now,
        'result': result,
        'follow_up': follow_up,
        'our_quotes': our_quotes,
        'competitor_quotes': competitor_quotes,
        'timeline': timeline,
    }
    project_path(project_id).write_text(
        json.dumps(project, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )
    return project


def build_overview_text(
    title: str,
    result: dict[str, Any],
    *,
    follow_up: dict[str, Any] | None = None,
    our_quotes: list[dict[str, Any]] | None = None,
    competitor_quotes: list[dict[str, Any]] | None = None,
    timeline: list[dict[str, Any]] | None = None,
) -> str:
    summary = result['document_summary']
    analysis = result['analysis']
    rows = result['register_rows']
    review = result['match_review']
    follow_up = normalize_follow_up(follow_up or {})
    our_quotes = normalize_quote_rows(our_quotes or [], default_company='我司')
    competitor_quotes = normalize_quote_rows(competitor_quotes or [])
    timeline = normalize_timeline_rows(timeline or [])

    def block(title: str, items: list[str]) -> str:
        lines = [item for item in items if item]
        return f"{title}\n" + ('\n'.join(f'- {item}' for item in lines) if lines else '- 无')

    sections = [
        '# 标书解析总览',
        '',
        f"归档标题：{title or summary['project_name'] or '未命名项目'}",
        '',
        '## 核心信息',
        f"项目名称：{summary['project_name'] or '未识别'}",
        f"招标编号：{summary['bid_no'] or '未识别'}",
        f"招标人：{summary['tenderer'] or '未识别'}",
        f"投标截止时间：{summary['bid_deadline'] or '未识别'}",
        f"开标时间：{summary['open_time'] or '未识别'}",
        f"投标文件递交方式：{summary['submission_method'] or '未识别'}",
        f"保证金金额：{summary['deposit_amount'] or '未识别'}",
        f"服务期限：{summary['service_period'] or '未识别'}",
        '',
        '## AI 解析',
        analysis['summary'] or '无',
        '',
        '## 跟踪结果',
        f"投标状态：{follow_up['bid_status'] or '待跟进'}",
        f"中标状态：{follow_up['award_status'] or '未知'}",
        f"中标日期：{follow_up['award_date'] or '-'}",
        f"中标单位：{follow_up['award_company'] or '-'}",
        f"我司中标价：{follow_up['our_award_amount'] or '-'}",
        f"竞对中标价：{follow_up['competitor_award_amount'] or '-'}",
        f"信息来源：{follow_up['information_source'] or '-'}",
        f"跟进备注：{follow_up['tracking_note'] or '-'}",
        '',
        block('## 资格文件清单', analysis['qualification_files']),
        '',
        block('## 报价与商务关注点', analysis['business_points']),
        '',
        block('## 评标得分关注点', analysis['scoring_points']),
        '',
        block('## 风险提示', analysis['risks']),
        '',
        '## 登记预览',
    ]

    if rows:
        for index, row in enumerate(rows, start=1):
            sections.extend(
                [
                    f"{index}. 项目名称：{row.get('project_name') or '未识别'}",
                    f"   招标编号：{row.get('bid_no') or '未识别'}",
                    f"   招标人：{row.get('tenderer') or '未识别'}",
                    f"   投标截止日期：{row.get('bid_deadline') or '未识别'}",
                    f"   保证金金额：{row.get('deposit_amount') or '未识别'}",
                    f"   备注：{row.get('remark') or '-'}",
                ]
            )
    else:
        sections.append('无')

    sections.extend(['', '## 我司报价一览'])
    if our_quotes:
        for index, row in enumerate(our_quotes, start=1):
            sections.extend(
                [
                    f"{index}. 标段：{row.get('package_name') or '-'}",
                    f"   轮次：第 {row.get('round_no') or 1} 轮",
                    f"   报价日期：{row.get('quote_date') or '-'}",
                    f"   币种：{row.get('currency') or 'CNY'}",
                    f"   单价：{row.get('unit_price') or '-'}",
                    f"   总价：{row.get('total_price') or '-'}",
                    f"   是否已投：{'是' if row.get('is_submitted') else '否'}",
                    f"   是否中标：{'是' if row.get('is_awarded') else '否'}",
                    f"   备注：{row.get('remark') or '-'}",
                ]
            )
    else:
        sections.append('- 无')

    sections.extend(['', '## 竞对报价一览'])
    if competitor_quotes:
        for index, row in enumerate(competitor_quotes, start=1):
            sections.extend(
                [
                    f"{index}. 公司：{row.get('quote_company') or '-'}",
                    f"   标段：{row.get('package_name') or '-'}",
                    f"   报价日期：{row.get('quote_date') or '-'}",
                    f"   币种：{row.get('currency') or 'CNY'}",
                    f"   单价：{row.get('unit_price') or '-'}",
                    f"   总价：{row.get('total_price') or '-'}",
                    f"   排名：{row.get('ranking') or '-'}",
                    f"   是否中标：{'是' if row.get('is_awarded') else '否'}",
                    f"   来源：{row.get('source') or '-'}",
                    f"   备注：{row.get('remark') or '-'}",
                ]
            )
    else:
        sections.append('- 无')

    sections.extend(['', '## 跟进时间线'])
    if timeline:
        for item in timeline:
            sections.append(f"- [{item.get('date') or '-'}] {item.get('type') or 'note'}：{item.get('note') or '-'}")
    else:
        sections.append('- 无')

    sections.extend(['', '## 匹配复核'])
    if review:
        for item in review:
            sections.extend(
                [
                    f"- 领域：{item.get('area') or '未分类'}",
                    f"  事项：{item.get('item') or '-'}",
                    f"  状态：{item.get('status') or '需人工确认'}",
                    f"  原因：{item.get('reason') or '-'}",
                    f"  建议动作：{item.get('action') or '-'}",
                ]
            )
    else:
        sections.append('- 无')
    return '\n'.join(sections).strip() + '\n'


def clear_sheet_rows(sheet: Any, start_row: int) -> None:
    if sheet.max_row >= start_row:
        sheet.delete_rows(start_row, sheet.max_row - start_row + 1)


def autosize_sheet(sheet: Any, widths: dict[int, int]) -> None:
    for index, width in widths.items():
        sheet.column_dimensions[chr(64 + index)].width = width


def apply_plain_table_style(sheet: Any, row_count: int, col_count: int) -> None:
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_font = Font(bold=True)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    for row_index in range(1, row_count + 1):
        for col_index in range(1, col_count + 1):
            cell = sheet.cell(row_index, col_index)
            cell.border = border
            cell.alignment = center if row_index <= 2 or col_index <= 2 else Alignment(vertical='top', wrap_text=True)
            if row_index in {1, 2, 4}:
                cell.font = header_font


def build_quote_workbook(
    *,
    title: str,
    follow_up: dict[str, Any],
    rows: list[dict[str, Any]],
    is_competitor: bool,
) -> BytesIO:
    follow_up = normalize_follow_up(follow_up or {})
    normalized_rows = normalize_quote_rows(rows or [], default_company='' if is_competitor else '我司')
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '报价一览'
    sheet['A1'] = '竞对报价对比' if is_competitor else '我司报价一览表'
    sheet['A2'] = '项目名称'
    sheet['B2'] = title or '未命名项目'
    sheet['C2'] = '投标状态'
    sheet['D2'] = follow_up.get('bid_status', '待跟进')
    sheet['E2'] = '中标状态'
    sheet['F2'] = follow_up.get('award_status', '未知')
    sheet['G2'] = '我司中标价' if not is_competitor else '竞对中标价'
    sheet['H2'] = follow_up.get('our_award_amount' if not is_competitor else 'competitor_award_amount', '')

    headers = (
        ['序号', '标段/合同包', '轮次', '报价日期', '报价主体', '币种', '税率/口径', '单价', '总价', '已投递', '是否中标', '备注']
        if not is_competitor
        else ['序号', '竞对公司', '标段/合同包', '报价日期', '币种', '单价', '总价', '排名', '是否中标', '来源', '备注']
    )
    start_row = 4
    for index, header in enumerate(headers, start=1):
        sheet.cell(start_row, index, header)

    for offset, row in enumerate(normalized_rows, start=1):
        row_index = start_row + offset
        if not is_competitor:
            values = [
                offset,
                row.get('package_name', ''),
                row.get('round_no', 1),
                row.get('quote_date', ''),
                row.get('quote_company', '我司'),
                row.get('currency', 'CNY'),
                row.get('tax_mode', ''),
                row.get('unit_price', ''),
                row.get('total_price', ''),
                '是' if row.get('is_submitted') else '否',
                '是' if row.get('is_awarded') else '否',
                row.get('remark', ''),
            ]
        else:
            values = [
                offset,
                row.get('quote_company', ''),
                row.get('package_name', ''),
                row.get('quote_date', ''),
                row.get('currency', 'CNY'),
                row.get('unit_price', ''),
                row.get('total_price', ''),
                row.get('ranking', ''),
                '是' if row.get('is_awarded') else '否',
                row.get('source', ''),
                row.get('remark', ''),
            ]
        for col_index, value in enumerate(values, start=1):
            sheet.cell(row_index, col_index, value)

    final_row = max(start_row + len(normalized_rows), start_row)
    apply_plain_table_style(sheet, final_row, len(headers))
    autosize_sheet(
        sheet,
        {
            1: 8,
            2: 24,
            3: 18,
            4: 16,
            5: 18,
            6: 12,
            7: 14,
            8: 14,
            9: 14,
            10: 14,
            11: 26,
            12: 24,
        },
    )
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def collect_project_items(
    *,
    q: str = '',
    year: str = '',
    award_status: str = '',
    bid_status: str = '',
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    keyword = q.strip().lower()
    year = compact_text(year)
    award_status = compact_text(award_status)
    bid_status = compact_text(bid_status)
    items: list[dict[str, Any]] = []
    for path in PROJECTS_DIR.glob('*.json'):
        project = read_project_file(path)
        if not project:
            continue
        meta = project_meta(project)
        text = ' '.join(
            [
                meta.get('title', ''),
                meta.get('project_name', ''),
                meta.get('bid_no', ''),
                meta.get('tenderer', ''),
                meta.get('source_file_name', ''),
            ]
        ).lower()
        if keyword and keyword not in text:
            continue
        if year and meta.get('register_year', '') != year:
            continue
        if award_status and meta.get('award_status', '') != award_status:
            continue
        if bid_status and meta.get('bid_status', '') != bid_status:
            continue
        items.append(meta)
    items.sort(key=lambda item: item.get('updated_at', ''), reverse=True)
    stats = {
        'total_projects': len(items),
        'awarded_projects': sum(1 for item in items if item.get('award_status') == '已中标'),
        'submitted_projects': sum(1 for item in items if item.get('bid_status') == '已投标'),
        'pending_projects': sum(1 for item in items if item.get('bid_status') == '待跟进'),
        'our_quote_rows': sum(int(item.get('our_quote_count') or 0) for item in items),
        'competitor_quote_rows': sum(int(item.get('competitor_quote_count') or 0) for item in items),
    }
    return items, stats


def build_ledger_workbook(
    *,
    q: str = '',
    year: str = '',
    award_status: str = '',
    bid_status: str = '',
) -> BytesIO:
    items, stats = collect_project_items(q=q, year=year, award_status=award_status, bid_status=bid_status)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '历史台账'
    sheet['A1'] = '历史台账汇总'
    sheet['A2'] = '筛选关键词'
    sheet['B2'] = compact_text(q) or '全部'
    sheet['C2'] = '登记年份'
    sheet['D2'] = compact_text(year) or '全部'
    sheet['E2'] = '投标状态'
    sheet['F2'] = compact_text(bid_status) or '全部'
    sheet['G2'] = '中标状态'
    sheet['H2'] = compact_text(award_status) or '全部'

    stat_labels = [
        ('项目总数', stats['total_projects']),
        ('已中标', stats['awarded_projects']),
        ('已投标', stats['submitted_projects']),
        ('待跟进', stats['pending_projects']),
        ('我司报价条数', stats['our_quote_rows']),
        ('竞对报价条数', stats['competitor_quote_rows']),
    ]
    for index, (label, value) in enumerate(stat_labels, start=1):
        column = (index - 1) * 2 + 1
        sheet.cell(3, column, label)
        sheet.cell(3, column + 1, value)

    headers = ['序号', '项目标题', '项目名称', '招标编号', '招标人', '登记年份', '投标状态', '中标状态', '我司报价条数', '竞对报价条数', '最后更新时间']
    for col_index, header in enumerate(headers, start=1):
        sheet.cell(5, col_index, header)

    for row_index, item in enumerate(items, start=6):
        values = [
            row_index - 5,
            item.get('title', ''),
            item.get('project_name', ''),
            item.get('bid_no', ''),
            item.get('tenderer', ''),
            item.get('register_year', ''),
            item.get('bid_status', ''),
            item.get('award_status', ''),
            item.get('our_quote_count', 0),
            item.get('competitor_quote_count', 0),
            item.get('updated_at', ''),
        ]
        for col_index, value in enumerate(values, start=1):
            sheet.cell(row_index, col_index, value)

    final_row = max(5, 5 + len(items))
    apply_plain_table_style(sheet, final_row, len(headers))
    autosize_sheet(sheet, {1: 8, 2: 26, 3: 30, 4: 22, 5: 24, 6: 12, 7: 12, 8: 12, 9: 12, 10: 14, 11: 20})
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def write_overview_block(sheet: Any, row_index: int, title: str, lines: list[str]) -> int:
    sheet.cell(row_index, 1, title)
    sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=4)
    row_index += 1
    if not lines:
        lines = ['无']
    for line in lines:
        sheet.cell(row_index, 1, line)
        sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=4)
        row_index += 1
    return row_index + 1


def build_overview_workbook(
    title: str,
    result: dict[str, Any],
    *,
    follow_up: dict[str, Any] | None = None,
    our_quotes: list[dict[str, Any]] | None = None,
    competitor_quotes: list[dict[str, Any]] | None = None,
    timeline: list[dict[str, Any]] | None = None,
) -> BytesIO:
    summary = result['document_summary']
    analysis = result['analysis']
    rows = result['register_rows']
    review = result['match_review']
    follow_up = normalize_follow_up(follow_up or {})
    our_quotes = normalize_quote_rows(our_quotes or [], default_company='我司')
    competitor_quotes = normalize_quote_rows(competitor_quotes or [])
    timeline = normalize_timeline_rows(timeline or [])

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '解析总览'
    sheet['A1'] = '标书解析总览'
    sheet.merge_cells('A1:D1')
    sheet['A2'] = '归档标题'
    sheet['B2'] = title or summary['project_name'] or '未命名项目'
    sheet['C2'] = '生成时间'
    sheet['D2'] = datetime.now().strftime('%Y-%m-%d %H:%M')

    core_pairs = [
        ('项目名称', summary['project_name'] or '未识别'),
        ('招标编号', summary['bid_no'] or '未识别'),
        ('招标人', summary['tenderer'] or '未识别'),
        ('投标截止时间', summary['bid_deadline'] or '未识别'),
        ('开标时间', summary['open_time'] or '未识别'),
        ('投标文件递交方式', summary['submission_method'] or '未识别'),
        ('保证金金额', summary['deposit_amount'] or '未识别'),
        ('服务期限', summary['service_period'] or '未识别'),
        ('投标状态', follow_up['bid_status'] or '待跟进'),
        ('中标状态', follow_up['award_status'] or '未知'),
        ('中标日期', follow_up['award_date'] or '-'),
        ('中标单位', follow_up['award_company'] or '-'),
        ('我司中标价', follow_up['our_award_amount'] or '-'),
        ('竞对中标价', follow_up['competitor_award_amount'] or '-'),
        ('信息来源', follow_up['information_source'] or '-'),
        ('跟进备注', follow_up['tracking_note'] or '-'),
    ]
    sheet['A4'] = '核心信息'
    sheet.merge_cells('A4:D4')
    row_index = 5
    for label, value in core_pairs:
        sheet.cell(row_index, 1, label)
        sheet.cell(row_index, 2, value)
        sheet.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=4)
        row_index += 1

    row_index += 1
    row_index = write_overview_block(sheet, row_index, 'AI 解析结论', [analysis['summary'] or '无'])
    row_index = write_overview_block(sheet, row_index, '资格文件清单', analysis['qualification_files'])
    row_index = write_overview_block(sheet, row_index, '报价与商务关注点', analysis['business_points'])
    row_index = write_overview_block(sheet, row_index, '评标得分关注点', analysis['scoring_points'])
    row_index = write_overview_block(sheet, row_index, '风险提示', analysis['risks'])

    register_lines = []
    for index, row in enumerate(rows, start=1):
        register_lines.append(
            f"{index}. 项目名称：{row.get('project_name') or '未识别'} | 招标编号：{row.get('bid_no') or '未识别'} | 招标人：{row.get('tenderer') or '未识别'} | 投标截止日期：{row.get('bid_deadline') or '未识别'} | 保证金金额：{row.get('deposit_amount') or '未识别'} | 备注：{row.get('remark') or '-'}"
        )
    row_index = write_overview_block(sheet, row_index, '登记预览', register_lines)

    our_lines = []
    for index, row in enumerate(our_quotes, start=1):
        our_lines.append(
            f"{index}. 标段：{row.get('package_name') or '-'} | 轮次：第 {row.get('round_no') or 1} 轮 | 报价日期：{row.get('quote_date') or '-'} | 币种：{row.get('currency') or 'CNY'} | 单价：{row.get('unit_price') or '-'} | 总价：{row.get('total_price') or '-'} | 已投：{'是' if row.get('is_submitted') else '否'} | 中标：{'是' if row.get('is_awarded') else '否'} | 备注：{row.get('remark') or '-'}"
        )
    row_index = write_overview_block(sheet, row_index, '我司报价一览', our_lines)

    competitor_lines = []
    for index, row in enumerate(competitor_quotes, start=1):
        competitor_lines.append(
            f"{index}. 公司：{row.get('quote_company') or '-'} | 标段：{row.get('package_name') or '-'} | 报价日期：{row.get('quote_date') or '-'} | 币种：{row.get('currency') or 'CNY'} | 单价：{row.get('unit_price') or '-'} | 总价：{row.get('total_price') or '-'} | 排名：{row.get('ranking') or '-'} | 中标：{'是' if row.get('is_awarded') else '否'} | 来源：{row.get('source') or '-'} | 备注：{row.get('remark') or '-'}"
        )
    row_index = write_overview_block(sheet, row_index, '竞对报价一览', competitor_lines)

    timeline_lines = [f"[{item.get('date') or '-'}] {item.get('type') or 'note'}：{item.get('note') or '-'}" for item in timeline]
    row_index = write_overview_block(sheet, row_index, '跟进时间线', timeline_lines)

    review_lines = []
    for item in review:
        review_lines.append(
            f"领域：{item.get('area') or '未分类'} | 事项：{item.get('item') or '-'} | 状态：{item.get('status') or '需人工确认'} | 原因：{item.get('reason') or '-'} | 建议动作：{item.get('action') or '-'}"
        )
    write_overview_block(sheet, row_index, '匹配复核', review_lines)

    apply_plain_table_style(sheet, sheet.max_row, 4)
    sheet.row_dimensions[1].height = 28
    sheet.column_dimensions['A'].width = 20
    sheet.column_dimensions['B'].width = 34
    sheet.column_dimensions['C'].width = 18
    sheet.column_dimensions['D'].width = 34
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def resolve_source_excerpt(value: str, excerpt: str, document_text_full: str) -> str:
    clean_excerpt = compact_paragraph(excerpt)
    if clean_excerpt and not any(marker in clean_excerpt for marker in ELLIPSIS_MARKERS):
        return clean_excerpt
    if not value or not document_text_full:
        return clean_excerpt
    index = document_text_full.find(value)
    if index < 0:
        match = SequenceMatcher(None, value, document_text_full).find_longest_match(
            0,
            len(value),
            0,
            len(document_text_full),
        )
        # ponytail: 只做一次最长公共片段回退；还不够稳时再上专门的模糊匹配。
        if match.size < 4:
            return clean_excerpt
        index = match.b
    start = document_text_full.rfind('\n', 0, index)
    end = document_text_full.find('\n', index + len(value))
    start = 0 if start < 0 else start + 1
    end = len(document_text_full) if end < 0 else end
    line = compact_paragraph(document_text_full[start:end])
    return line or clean_excerpt


def project_register_rows(rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    return [refine_register_row(row if isinstance(row, dict) else {}, normalize_summary({})) for row in rows]


MARKET_KINDS = {'cargo', 'newbuilding'}
MARKET_CARGO_FIELDS = [
    ('board_type', '货盘类型'),
    ('segment', '业务板块'),
    ('cargo_name', '货品名称'),
    ('tonnage', '货物吨数'),
    ('load_port', '装港'),
    ('discharge_port', '卸港'),
    ('laycan', '装载期'),
    ('cargo_owner', '货主/租家'),
    ('cargo_date', '货盘日期'),
    ('source', '货盘来源'),
    ('status', '是否达成合作'),
    ('final_price', '成交价'),
    ('deal_date', '成交日期'),
    ('deal_price', '最终成交价格'),
    ('price_unit', '价格单位'),
    ('currency', '币种'),
    ('route', '航线'),
    ('cargo_standard_name', '货品标准名'),
    ('market_info', '市场了解信息'),
    ('loss_reason', '未达成/放弃原因'),
    ('competitor_name', '竞争对手'),
    ('competitor_price', '竞争对手价格'),
    ('remark', '备注'),
]
MARKET_NEWBUILDING_FIELDS = [
    ('stage', '建造阶段'),
    ('ship_name', '船名'),
    ('update_date', '更新日期'),
    ('shipyard', '造船厂'),
    ('owner', '船东'),
    ('dwt', '载重吨DWT'),
    ('build_status', '当前建造状态'),
    ('delivery_time', '预计交付时间'),
    ('actual_delivery_date', '实际出厂时间'),
    ('status_update_date', '状态更新时间'),
    ('status_note', '状态变化备注'),
    ('contract_date', '合同签订日期'),
    ('contract_price', '新造船合同价格'),
    ('ship_type', '船型'),
    ('source', '信息来源'),
    ('remark', '备注'),
]
MARKET_SEARCH_FIELDS = {
    'cargo': (
        'cargo_name',
        'tonnage',
        'load_port',
        'discharge_port',
        'laycan',
        'cargo_owner',
        'cargo_date',
        'source',
        'status',
        'final_price',
        'deal_date',
        'deal_price',
        'price_unit',
        'currency',
        'route',
        'cargo_standard_name',
        'market_info',
        'loss_reason',
        'competitor_name',
        'competitor_price',
        'remark',
    ),
    'newbuilding': (
        'stage',
        'ship_name',
        'update_date',
        'shipyard',
        'owner',
        'dwt',
        'build_status',
        'delivery_time',
        'actual_delivery_date',
        'status_update_date',
        'status_note',
        'contract_date',
        'contract_price',
        'ship_type',
        'source',
        'remark',
    ),
}


def market_kind_dir(kind: str) -> Path:
    if kind not in MARKET_KINDS:
        raise HTTPException(status_code=400, detail='市场情报类型无效。')
    folder = MARKET_SKILL_DIR / kind
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def market_record_path(kind: str, record_id: str) -> Path:
    if not re.fullmatch(r'[a-f0-9]{32}', record_id):
        raise HTTPException(status_code=400, detail='市场情报记录 ID 无效。')
    return market_kind_dir(kind) / f'{record_id}.json'


def market_text(value: Any) -> str:
    return compact_paragraph(value)


def today_text() -> str:
    return datetime.now().strftime('%Y-%m-%d')


def normalize_market_cargo(raw: Any, *, record_id: str = '', created_at: str = '') -> dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    now = datetime.now().isoformat(timespec='seconds')
    board_type = market_text(source.get('board_type') or '即时货盘')
    if board_type not in {'长期货盘', '即时货盘'}:
        board_type = '即时货盘'
    segment = market_text(source.get('segment') or '内贸化')
    if segment not in {'内贸化', '内贸油', '外贸'}:
        segment = '内贸化'
    status = market_text(source.get('status') or '跟进中')
    if status not in {'跟进中', '已成交', '未成交', '放弃'}:
        status = '跟进中'
    load_port = market_text(source.get('load_port'))
    discharge_port = market_text(source.get('discharge_port'))
    final_price = market_text(source.get('final_price'))
    deal_price = market_text(source.get('deal_price') or final_price)
    route = market_text(source.get('route')) or ' - '.join(part for part in [load_port, discharge_port] if part)
    cargo_name = market_text(source.get('cargo_name'))
    return {
        'id': record_id or market_text(source.get('id')) or uuid.uuid4().hex,
        'kind': 'cargo',
        'board_type': board_type,
        'segment': segment,
        'cargo_name': cargo_name,
        'tonnage': market_text(source.get('tonnage')),
        'load_port': load_port,
        'discharge_port': discharge_port,
        'laycan': market_text(source.get('laycan')),
        'cargo_owner': market_text(source.get('cargo_owner')),
        'cargo_date': market_text(source.get('cargo_date')) or today_text(),
        'source': market_text(source.get('source') or '手动录入'),
        'status': status,
        'final_price': final_price or deal_price,
        'deal_date': market_text(source.get('deal_date')),
        'deal_price': deal_price,
        'price_unit': market_text(source.get('price_unit') or '元/吨'),
        'currency': market_text(source.get('currency') or 'CNY'),
        'route': route,
        'cargo_standard_name': market_text(source.get('cargo_standard_name')) or cargo_name,
        'market_info': market_text(source.get('market_info')),
        'loss_reason': market_text(source.get('loss_reason')),
        'competitor_name': market_text(source.get('competitor_name')),
        'competitor_price': market_text(source.get('competitor_price')),
        'remark': market_text(source.get('remark')),
        'raw_text': market_text(source.get('raw_text')),
        'source_channel': market_text(source.get('source_channel')),
        'source_open_id': market_text(source.get('source_open_id')),
        'source_chat_id': market_text(source.get('source_chat_id')),
        'source_message_id': market_text(source.get('source_message_id')),
        'source_text': market_text(source.get('source_text')),
        'confirmed_at': market_text(source.get('confirmed_at')),
        'created_at': created_at or market_text(source.get('created_at')) or now,
        'updated_at': now,
    }


def normalize_market_newbuilding(raw: Any, *, record_id: str = '', created_at: str = '') -> dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    now = datetime.now().isoformat(timespec='seconds')
    stage = market_text(source.get('stage') or '信息待获取')
    allowed_stages = {'已完造并出厂投运', '合同签订未开造', '预计2027年完造出厂', '预计2028年完造出厂', '信息待获取'}
    if stage not in allowed_stages:
        stage = '信息待获取'
    return {
        'id': record_id or market_text(source.get('id')) or uuid.uuid4().hex,
        'kind': 'newbuilding',
        'stage': stage,
        'ship_name': market_text(source.get('ship_name')),
        'update_date': market_text(source.get('update_date')) or today_text(),
        'shipyard': market_text(source.get('shipyard')),
        'owner': market_text(source.get('owner')),
        'dwt': market_text(source.get('dwt')),
        'build_status': market_text(source.get('build_status')),
        'delivery_time': market_text(source.get('delivery_time')),
        'actual_delivery_date': market_text(source.get('actual_delivery_date')),
        'status_update_date': market_text(source.get('status_update_date') or source.get('update_date')) or today_text(),
        'status_note': market_text(source.get('status_note')),
        'contract_date': market_text(source.get('contract_date')),
        'contract_price': market_text(source.get('contract_price')),
        'ship_type': market_text(source.get('ship_type')),
        'source': market_text(source.get('source') or '手动录入'),
        'remark': market_text(source.get('remark')),
        'raw_text': market_text(source.get('raw_text')),
        'source_channel': market_text(source.get('source_channel')),
        'source_open_id': market_text(source.get('source_open_id')),
        'source_chat_id': market_text(source.get('source_chat_id')),
        'source_message_id': market_text(source.get('source_message_id')),
        'source_text': market_text(source.get('source_text')),
        'confirmed_at': market_text(source.get('confirmed_at')),
        'created_at': created_at or market_text(source.get('created_at')) or now,
        'updated_at': now,
    }


def normalize_market_record(kind: str, raw: Any, *, record_id: str = '', created_at: str = '') -> dict[str, Any]:
    if kind == 'cargo':
        return normalize_market_cargo(raw, record_id=record_id, created_at=created_at)
    if kind == 'newbuilding':
        return normalize_market_newbuilding(raw, record_id=record_id, created_at=created_at)
    raise HTTPException(status_code=400, detail='市场情报类型无效。')


def read_market_record(kind: str, path: Path) -> dict[str, Any] | None:
    try:
        raw = json.loads(path.read_text(encoding='utf-8'))
    except (OSError, json.JSONDecodeError):
        return None
    return normalize_market_record(kind, raw, record_id=path.stem, created_at=market_text(raw.get('created_at')))


def load_market_record(kind: str, record_id: str) -> dict[str, Any]:
    path = market_record_path(kind, record_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail='市场情报记录不存在。')
    record = read_market_record(kind, path)
    if not record:
        raise HTTPException(status_code=500, detail='市场情报记录文件损坏。')
    return record


def save_market_record(kind: str, raw: dict[str, Any], *, record_id: str | None = None, created_at: str = '') -> dict[str, Any]:
    record = normalize_market_record(kind, raw, record_id=record_id or '', created_at=created_at)
    path = market_record_path(kind, record['id'])
    path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding='utf-8')
    return record


def list_market_records(
    kind: str,
    *,
    q: str = '',
    segment: str = '',
    board_type: str = '',
    status: str = '',
    stage: str = '',
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    keyword = market_text(q).lower()
    items: list[dict[str, Any]] = []
    for path in market_kind_dir(kind).glob('*.json'):
        record = read_market_record(kind, path)
        if not record:
            continue
        haystack = ' '.join(str(record.get(key, '')) for key in MARKET_SEARCH_FIELDS[kind]).lower()
        if keyword and keyword not in haystack:
            continue
        if kind == 'cargo':
            if segment and record.get('segment') != segment:
                continue
            if board_type and record.get('board_type') != board_type:
                continue
            if status and record.get('status') != status:
                continue
        if kind == 'newbuilding' and stage and record.get('stage') != stage:
            continue
        items.append(record)
    items.sort(key=lambda item: item.get('updated_at', ''), reverse=True)
    if kind == 'cargo':
        stats = {
            'total': len(items),
            'won': sum(1 for item in items if item.get('status') == '已成交'),
            'lost': sum(1 for item in items if item.get('status') in {'未成交', '放弃'}),
            'tracking': sum(1 for item in items if item.get('status') == '跟进中'),
            'with_final_price': sum(1 for item in items if item.get('final_price')),
        }
    else:
        stats = {'total': len(items)}
        for stage_name in ['已完造并出厂投运', '合同签订未开造', '预计2027年完造出厂', '预计2028年完造出厂', '信息待获取']:
            stats[stage_name] = sum(1 for item in items if item.get('stage') == stage_name)
    return items, stats


def parse_market_number(value: Any) -> float | None:
    text = market_text(value).replace(',', '')
    match = re.search(r'-?\d+(?:\.\d+)?', text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def parse_market_date(value: Any) -> Any:
    text = market_text(value)
    match = re.search(r'(20\d{2})[年/\-.]?\s*(\d{1,2})?[月/\-.]?\s*(\d{1,2})?', text)
    if not match:
        return None
    year = int(match.group(1))
    month = int(match.group(2) or 1)
    day = int(match.group(3) or 1)
    try:
        return datetime(year, month, day).date()
    except ValueError:
        return None


def market_date_text(value: Any) -> str:
    parsed = parse_market_date(value)
    return parsed.isoformat() if parsed else market_text(value)


def market_month_text(value: Any) -> str:
    parsed = parse_market_date(value)
    return parsed.strftime('%Y-%m') if parsed else ''


def market_in_date_range(value: Any, start_date: Any, end_date: Any) -> bool:
    parsed = parse_market_date(value)
    if start_date and (not parsed or parsed < start_date):
        return False
    if end_date and (not parsed or parsed > end_date):
        return False
    return True


def market_filter_value(filters: dict[str, Any], key: str) -> str:
    return market_text(filters.get(key, '')) if isinstance(filters, dict) else ''


def filtered_market_report_items(kind: str, payload: MarketReportPayload) -> list[dict[str, Any]]:
    filters = payload.filters if isinstance(payload.filters, dict) else {}
    items, _ = list_market_records(
        kind,
        q=market_filter_value(filters, 'q'),
        segment=market_filter_value(filters, 'segment'),
        board_type=market_filter_value(filters, 'board_type'),
        status=market_filter_value(filters, 'status'),
        stage=market_filter_value(filters, 'stage'),
    )
    start_date = parse_market_date(payload.start_date)
    end_date = parse_market_date(payload.end_date)
    cargo_filter = market_filter_value(filters, 'cargo')
    route_filter = market_filter_value(filters, 'route')
    filtered: list[dict[str, Any]] = []
    for item in items:
        if kind == 'cargo':
            if cargo_filter and cargo_filter not in (item.get('cargo_standard_name') or item.get('cargo_name') or ''):
                continue
            if route_filter and route_filter not in (item.get('route') or ''):
                continue
            record_date = item.get('deal_date') or item.get('cargo_date') or item.get('updated_at')
        else:
            record_date = item.get('status_update_date') or item.get('update_date') or item.get('updated_at')
        if not market_in_date_range(record_date, start_date, end_date):
            continue
        filtered.append(item)
    return filtered


def top_counts(values: list[str], limit: int = 6) -> list[dict[str, Any]]:
    counts: dict[str, int] = {}
    for value in values:
        value = market_text(value)
        if value:
            counts[value] = counts.get(value, 0) + 1
    return [{'name': name, 'count': count} for name, count in sorted(counts.items(), key=lambda item: item[1], reverse=True)[:limit]]


def fallback_cargo_report(payload: MarketReportPayload) -> dict[str, Any]:
    items = filtered_market_report_items('cargo', payload)
    deal_rows: list[dict[str, Any]] = []
    groups: dict[tuple[str, str, str], dict[str, Any]] = {}
    for item in items:
        if item.get('status') != '已成交':
            continue
        price = parse_market_number(item.get('deal_price') or item.get('final_price'))
        if price is None:
            continue
        deal_date = item.get('deal_date') or item.get('cargo_date') or item.get('updated_at')
        month = market_month_text(deal_date) or '未填月份'
        cargo = item.get('cargo_standard_name') or item.get('cargo_name') or '未填货品'
        route = item.get('route') or ' - '.join(part for part in [item.get('load_port'), item.get('discharge_port')] if part) or '未填航线'
        row = {
            'deal_date': market_date_text(deal_date),
            'month': month,
            'cargo': cargo,
            'route': route,
            'price': price,
            'price_unit': item.get('price_unit') or '元/吨',
            'currency': item.get('currency') or 'CNY',
            'cargo_owner': item.get('cargo_owner') or '',
            'competitor_name': item.get('competitor_name') or '',
            'competitor_price': item.get('competitor_price') or '',
            'remark': item.get('remark') or '',
        }
        deal_rows.append(row)
        group = groups.setdefault((month, cargo, route), {'month': month, 'cargo': cargo, 'route': route, 'sum': 0.0, 'count': 0, 'price_unit': row['price_unit'], 'currency': row['currency']})
        group['sum'] += price
        group['count'] += 1
    trend_points = [
        {
            'month': group['month'],
            'cargo': group['cargo'],
            'route': group['route'],
            'avg_price': round(group['sum'] / group['count'], 2),
            'count': group['count'],
            'price_unit': group['price_unit'],
            'currency': group['currency'],
        }
        for group in groups.values()
    ]
    trend_points.sort(key=lambda item: (item['month'], item['cargo'], item['route']))
    deal_rows.sort(key=lambda item: item.get('deal_date') or '', reverse=True)
    won = sum(1 for item in items if item.get('status') == '已成交')
    lost = sum(1 for item in items if item.get('status') in {'未成交', '放弃'})
    tracking = sum(1 for item in items if item.get('status') == '跟进中')
    avg_price = round(sum(row['price'] for row in deal_rows) / len(deal_rows), 2) if deal_rows else 0
    source_stats = {
        'total': len(items),
        'won': won,
        'lost': lost,
        'tracking': tracking,
        'deal_count': len(deal_rows),
        'avg_price': avg_price,
        'win_rate': round(won * 100 / len(items), 1) if items else 0,
        'loss_reasons': top_counts([item.get('loss_reason', '') for item in items if item.get('status') in {'未成交', '放弃'}]),
        'competitor_price_count': sum(1 for item in items if parse_market_number(item.get('competitor_price')) is not None),
    }
    direction = '成交样本不足，暂不判断运价趋势'
    if len(trend_points) >= 2:
        direction = '近期均价上行' if trend_points[-1]['avg_price'] > trend_points[0]['avg_price'] else '近期均价持平或下行'
    return {
        'summary': f"本期共纳入 {len(items)} 条商机，已成交 {won} 条，成交价样本 {len(deal_rows)} 条，平均成交价 {avg_price or '-'}。",
        'key_findings': [direction, f"成交率约 {source_stats['win_rate']}%。", f"竞对价格样本 {source_stats['competitor_price_count']} 条。"],
        'trend_points': trend_points,
        'detail_rows': deal_rows,
        'risks': ['成交样本少时不要直接外推市场价格。'] if len(deal_rows) < 3 else [],
        'recommendations': ['继续补全成交日期、最终成交价、竞对价格和丢单原因，优先沉淀同航线同货品样本。'],
        'source_stats': source_stats,
        'generated_at': datetime.now().isoformat(timespec='seconds'),
        'period': payload.period,
        'kind': 'cargo',
    }


def fallback_newbuilding_report(payload: MarketReportPayload) -> dict[str, Any]:
    items = filtered_market_report_items('newbuilding', payload)
    stage_counts = top_counts([item.get('stage', '') for item in items], limit=10)
    delivery_years: list[str] = []
    for item in items:
        delivery_text = ' '.join([item.get('actual_delivery_date', ''), item.get('delivery_time', ''), item.get('stage', '')])
        match = re.search(r'(20\d{2})', delivery_text)
        if match:
            delivery_years.append(match.group(1))
    recent_changes = sorted(
        [
            {
                'status_update_date': market_date_text(item.get('status_update_date') or item.get('update_date')),
                'ship_name': item.get('ship_name') or '',
                'shipyard': item.get('shipyard') or '',
                'owner': item.get('owner') or '',
                'stage': item.get('stage') or '',
                'delivery_time': item.get('delivery_time') or '',
                'actual_delivery_date': item.get('actual_delivery_date') or '',
                'status_note': item.get('status_note') or item.get('build_status') or item.get('remark') or '',
            }
            for item in items
        ],
        key=lambda item: item.get('status_update_date') or '',
        reverse=True,
    )
    delivered = sum(1 for item in items if item.get('stage') == '已完造并出厂投运' or item.get('actual_delivery_date'))
    source_stats = {
        'total': len(items),
        'delivered': delivered,
        'stage_counts': stage_counts,
        'delivery_year_counts': top_counts(delivery_years, limit=8),
        'recent_change_count': len(recent_changes),
    }
    stage_text = ', '.join(f"{item['name']} {item['count']} 条" for item in stage_counts) or '暂无'
    return {
        'summary': f"本期共纳入 {len(items)} 条新造船信息，已出厂/投运 {delivered} 条。",
        'key_findings': [f"主要阶段分布：{stage_text}。"],
        'trend_points': recent_changes[:30],
        'detail_rows': recent_changes,
        'risks': ['若未来交付集中，需关注供给释放对运价和船舶投资节奏的影响。'] if items else [],
        'recommendations': ['持续维护状态更新时间、预计交付时间和实际出厂时间，形成周/月状态变化口径。'],
        'source_stats': source_stats,
        'generated_at': datetime.now().isoformat(timespec='seconds'),
        'period': payload.period,
        'kind': 'newbuilding',
    }


def ai_enhance_market_report(report: dict[str, Any], custom_prompt: str = '') -> dict[str, Any]:
    config = load_config()
    if not config.get('base_url') or not config.get('api_key') or not config.get('model'):
        return report
    prompt = json.dumps(
        {
            'kind': report.get('kind'),
            'period': report.get('period'),
            'source_stats': report.get('source_stats'),
            'trend_points': report.get('trend_points', [])[:80],
            'detail_rows': report.get('detail_rows', [])[:80],
            'user_prompt': market_text(custom_prompt)[:2000],
            'required_json': {
                'summary': '',
                'key_findings': [],
                'risks': [],
                'recommendations': [],
            },
        },
        ensure_ascii=False,
    )
    try:
        content = call_chat_completion(
            config,
            system_prompt='你是航运市场分析助手。优先满足用户提示词，但只能基于输入数据输出严格 JSON，不要 Markdown，不要编造未提供的数据。',
            user_prompt=prompt,
            temperature=0.1,
        )
        ai_report = parse_json_text(content)
    except HTTPException as exc:
        report['risks'] = [*report.get('risks', []), f"AI 结论生成失败，已保留本地统计报告：{exc.detail}"]
        return report
    for key in ['summary', 'key_findings', 'risks', 'recommendations']:
        if ai_report.get(key):
            report[key] = ai_report[key]
    return report


def build_market_report(payload: MarketReportPayload) -> dict[str, Any]:
    kind = payload.kind.strip()
    if kind not in MARKET_KINDS:
        raise HTTPException(status_code=400, detail='市场情报类型无效。')
    if payload.period not in {'weekly', 'monthly', 'custom'}:
        payload.period = 'custom'
    report = fallback_cargo_report(payload) if kind == 'cargo' else fallback_newbuilding_report(payload)
    return ai_enhance_market_report(report, payload.custom_prompt)


def build_market_report_workbook(kind: str, report: dict[str, Any]) -> BytesIO:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '分析报告'
    title = '商机市场分析报告' if kind == 'cargo' else '新造船市场分析报告'
    sheet['A1'] = title
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)
    rows = [
        ('生成时间', report.get('generated_at', datetime.now().isoformat(timespec='seconds'))),
        ('报告周期', '周报' if report.get('period') == 'weekly' else '月报'),
        ('报告摘要', report.get('summary', '')),
        ('关键发现', '\n'.join(map(str, report.get('key_findings') or []))),
        ('风险提示', '\n'.join(map(str, report.get('risks') or []))),
        ('经营建议', '\n'.join(map(str, report.get('recommendations') or []))),
    ]
    for row_index, (label, value) in enumerate(rows, start=2):
        sheet.cell(row_index, 1, label)
        sheet.cell(row_index, 2, value)
        sheet.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=4)
    apply_plain_table_style(sheet, 1 + len(rows), 4)
    autosize_sheet(sheet, {1: 16, 2: 36, 3: 36, 4: 36})

    stats_sheet = workbook.create_sheet('统计数据')
    stats_sheet.append(['指标', '值'])
    for key, value in (report.get('source_stats') or {}).items():
        stats_sheet.append([key, json.dumps(value, ensure_ascii=False) if isinstance(value, (list, dict)) else value])
    apply_plain_table_style(stats_sheet, max(1, stats_sheet.max_row), 2)
    autosize_sheet(stats_sheet, {1: 22, 2: 72})

    detail_sheet = workbook.create_sheet('数据明细')
    detail_rows = report.get('detail_rows') or report.get('trend_points') or []
    headers = list(detail_rows[0].keys()) if detail_rows else ['说明']
    detail_sheet.append(headers)
    for row in detail_rows:
        detail_sheet.append([row.get(header, '') for header in headers])
    if not detail_rows:
        detail_sheet.append(['暂无明细数据'])
    apply_plain_table_style(detail_sheet, max(1, detail_sheet.max_row), len(headers))
    autosize_sheet(detail_sheet, {index: 18 for index in range(1, min(len(headers), 12) + 1)})

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def build_market_report_docx(kind: str, report: dict[str, Any]) -> BytesIO:
    if Document is None:
        raise HTTPException(status_code=500, detail='缺少 python-docx 依赖，请运行：python -m pip install python-docx，或重新使用一键启动安装依赖。')

    def set_cell_shading(cell: Any, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), fill)
        tc_pr.append(shade)

    def set_cell_text(cell: Any, text: Any, *, bold: bool = False, color: str = '222222') -> None:
        cell.text = ''
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(text if text is not None else ''))
        run.bold = bold
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(33, 91, 116)

    def add_bullets(items: list[Any]) -> None:
        for item in items or ['暂无']:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            run = paragraph.add_run(f"• {item}")
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10.5)

    def add_simple_table(headers: list[str], rows: list[list[Any]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for index, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[index], header, bold=True, color='FFFFFF')
            set_cell_shading(table.rows[0].cells[index], '2F5D73')
        for row in rows or [['暂无数据'] + [''] * (len(headers) - 1)]:
            cells = table.add_row().cells
            for index, value in enumerate(row[: len(headers)]):
                set_cell_text(cells[index], value)

    title = '商机市场 AI 分析报告' if kind == 'cargo' else '新造船市场 AI 分析报告'
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(33, 91, 116)

    info_rows = [
        ['生成时间', report.get('generated_at', datetime.now().isoformat(timespec='seconds'))],
        ['分析口径', '用户自选时间范围'],
        ['报告类型', '商机成交价趋势' if kind == 'cargo' else '新造船状态追踪'],
    ]
    add_simple_table(['项目', '内容'], info_rows)

    add_heading('一、AI 分析结论')
    paragraph = document.add_paragraph(str(report.get('summary') or '暂无结论'))
    paragraph.paragraph_format.first_line_indent = Pt(21)

    add_heading('二、关键发现')
    add_bullets(report.get('key_findings') or [])

    add_heading('三、风险提示')
    add_bullets(report.get('risks') or [])

    add_heading('四、经营建议')
    add_bullets(report.get('recommendations') or [])

    stats_rows = []
    for key, value in (report.get('source_stats') or {}).items():
        stats_rows.append([key, json.dumps(value, ensure_ascii=False) if isinstance(value, (list, dict)) else value])
    add_heading('五、来源统计')
    add_simple_table(['指标', '值'], stats_rows)

    detail_rows = report.get('detail_rows') or report.get('trend_points') or []
    if detail_rows:
        headers = ['日期/月', '货品/船名', '航线/船厂', '价格/阶段', '备注']
        rows = []
        for row in detail_rows[:30]:
            if kind == 'cargo':
                rows.append([row.get('deal_date') or row.get('month'), row.get('cargo'), row.get('route'), row.get('price'), row.get('competitor_price') or row.get('remark')])
            else:
                rows.append([row.get('status_update_date'), row.get('ship_name') or row.get('owner'), row.get('shipyard'), row.get('stage'), row.get('status_note')])
        add_heading('六、趋势/状态明细')
        add_simple_table(headers, rows)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def market_extract_json(kind: str, text: str) -> dict[str, Any] | None:
    config = load_config()
    if not config.get('base_url') or not config.get('api_key') or not config.get('model'):
        return None
    if kind == 'cargo':
        schema = {
            'board_type': '长期货盘或即时货盘',
            'segment': '内贸化、内贸油或外贸',
            'cargo_name': '',
            'tonnage': '',
            'load_port': '',
            'discharge_port': '',
            'laycan': '',
            'cargo_owner': '',
            'cargo_date': '',
            'source': '',
            'remark': '',
        }
    else:
        schema = {
            'stage': '已完造并出厂投运、合同签订未开造、预计2027年完造出厂、预计2028年完造出厂、信息待获取',
            'ship_name': '',
            'update_date': '',
            'shipyard': '',
            'owner': '',
            'dwt': '',
            'build_status': '',
            'delivery_time': '',
            'contract_date': '',
            'contract_price': '',
            'ship_type': '',
            'source': '',
            'remark': '',
        }
    prompt = json.dumps({'text': text[:12000], 'schema': schema}, ensure_ascii=False)
    try:
        content = call_chat_completion(
            config,
            system_prompt='你是市场情报结构化抽取助手，只返回严格 JSON 对象，不要 Markdown。',
            user_prompt=prompt,
            temperature=0,
        )
        return parse_json_text(content)
    except HTTPException:
        return None


def guess_market_cargo(text: str) -> dict[str, Any]:
    record: dict[str, Any] = {'raw_text': text, 'cargo_date': today_text(), 'source': '手动录入'}
    tonnage_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:吨|t|T)', text)
    if tonnage_match:
        record['tonnage'] = f"{tonnage_match.group(1)}吨"
    route_match = re.search(r'([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,20})\s*[~～\-—→至到]\s*([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,20})', text)
    if route_match:
        record['load_port'] = route_match.group(1).strip()
        record['discharge_port'] = route_match.group(2).strip()
    load_match = re.search(r'(?:装港|装货港|起运港|从|由)?\s*([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,20})\s*(?:装|装出|起运)', text)
    if load_match and not record.get('load_port'):
        record['load_port'] = load_match.group(1).strip()
    discharge_match = re.search(r'(?:卸港|卸货港|目的港|到|至)?\s*([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,20})\s*(?:卸|卸货|到港)', text)
    if discharge_match and not record.get('discharge_port'):
        record['discharge_port'] = discharge_match.group(1).strip()
    cargo_match = re.search(r'(?:\d+(?:\.\d+)?\s*(?:吨|t|T)\s*)?([\u4e00-\u9fa5A-Za-z0-9]{2,18})(?:，|,|、|\s)', text)
    if cargo_match:
        record['cargo_name'] = cargo_match.group(1).strip()
    owner_match = re.search(r'(?:货主|租家|客户)[：:\s]*([\u4e00-\u9fa5A-Za-z0-9]{2,18})', text)
    if not owner_match:
        owner_match = re.search(r'([\u4e00-\u9fa5A-Za-z0-9]{2,18})(?:的计划|计划|货盘)', text)
    if owner_match:
        record['cargo_owner'] = owner_match.group(1).strip()
    if '月内' in text:
        record['laycan'] = '月内装出'
    elif '月底' in text:
        record['laycan'] = '月底前'
    if any(word in text for word in ['外贸', '出口', '进口', '境外']):
        record['segment'] = '外贸'
    elif any(word in text for word in ['油', '柴油', '汽油', '燃料油', '成品油']):
        record['segment'] = '内贸油'
    else:
        record['segment'] = '内贸化'
    record['board_type'] = '长期货盘' if any(word in text for word in ['长期', '年度', '月度', '包运']) else '即时货盘'
    return record


def guess_market_newbuilding(text: str) -> dict[str, Any]:
    record: dict[str, Any] = {'raw_text': text, 'update_date': today_text(), 'source': '手动录入'}
    dwt_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:万)?\s*(?:载重吨|DWT|dwt|吨级)', text)
    if dwt_match:
        record['dwt'] = dwt_match.group(0)
    year_match = re.search(r'(2027|2028)\s*年', text)
    if year_match:
        record['delivery_time'] = f'{year_match.group(1)}年'
        record['stage'] = f'预计{year_match.group(1)}年完造出厂'
    elif any(word in text for word in ['交付', '投运', '出厂']):
        record['stage'] = '已完造并出厂投运'
    elif any(word in text for word in ['签订', '合同', '订单']):
        record['stage'] = '合同签订未开造'
    else:
        record['stage'] = '信息待获取'
    yard_match = re.search(r'([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,24}(?:船厂|造船|重工|船舶))', text)
    if yard_match:
        record['shipyard'] = yard_match.group(1)
    owner_match = re.search(r'(?:船东|订造方|owner)[：:\s]*([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,24})', text, re.I)
    if not owner_match:
        owner_match = re.search(r'为\s*([\u4e00-\u9fa5A-Za-z0-9·（）()]{2,24}船东)\s*(?:建造|订造)', text)
    if owner_match:
        record['owner'] = owner_match.group(1)
    ship_match = re.search(r'(?:船名|命名为|名为)[：:\s]*([\u4e00-\u9fa5A-Za-z0-9·（）()#-]{2,24})', text)
    if not ship_match:
        ship_match = re.search(r'^([\u4e00-\u9fa5A-Za-z0-9·（）()#-]{2,24}船[\u4e00-\u9fa5A-Za-z0-9·（）()#-]*)\s*(?:，|,|。|\s)', text)
    if ship_match:
        record['ship_name'] = ship_match.group(1)
    if '化学品' in text:
        record['ship_type'] = '化学品船'
    elif '油' in text:
        record['ship_type'] = '油船'
    return record


def extract_market_record(kind: str, text: str) -> dict[str, Any]:
    ai_data = market_extract_json(kind, text)
    guessed = ai_data if isinstance(ai_data, dict) else (guess_market_cargo(text) if kind == 'cargo' else guess_market_newbuilding(text))
    guessed['raw_text'] = text
    return normalize_market_record(kind, guessed)


def build_market_workbook(kind: str, records: list[dict[str, Any]], stats: dict[str, Any]) -> BytesIO:
    fields = MARKET_CARGO_FIELDS if kind == 'cargo' else MARKET_NEWBUILDING_FIELDS
    title = '商机收集台账' if kind == 'cargo' else '新造船收集台账'
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = title
    sheet['A1'] = title
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(fields) + 1)
    sheet['A2'] = '生成时间'
    sheet['B2'] = datetime.now().strftime('%Y-%m-%d %H:%M')
    sheet['C2'] = '记录数'
    sheet['D2'] = stats.get('total', len(records))
    for col_index, header in enumerate(['序号'] + [label for _, label in fields], start=1):
        sheet.cell(4, col_index, header)
    for row_index, record in enumerate(records, start=5):
        sheet.cell(row_index, 1, row_index - 4)
        for col_index, (key, _) in enumerate(fields, start=2):
            sheet.cell(row_index, col_index, record.get(key, ''))
    apply_plain_table_style(sheet, max(4, 4 + len(records)), len(fields) + 1)
    autosize_sheet(sheet, {1: 8, 2: 16, 3: 16, 4: 18, 5: 14, 6: 16, 7: 16, 8: 18, 9: 18, 10: 16, 11: 16, 12: 18, 13: 18, 14: 18, 15: 18, 16: 18, 17: 28, 18: 28})
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


@app.get('/')
def root() -> RedirectResponse:
    return RedirectResponse(url='/login', status_code=302)


@app.get('/app')
def index(
    authorization: str | None = Header(default=None),
    auth_cookie: str | None = Cookie(default=None, alias=AUTH_COOKIE_NAME),
):
    if not resolve_auth_user(authorization, auth_cookie):
        return RedirectResponse(url='/login', status_code=302)
    return FileResponse(STATIC_DIR / 'index.html')


@app.get('/login')
def login_page() -> FileResponse:
    return FileResponse(STATIC_DIR / 'login.html')


@app.get('/api/template-meta')
def template_meta(_: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    workbook = load_workbook(get_register_template_path(), read_only=True)
    sheet_names = workbook.sheetnames
    workbook.close()
    current_year = str(datetime.now().year)
    return {
        'extraction_fields': EXTRACTION_FIELDS,
        'register_columns': REGISTER_COLUMNS,
        'register_sheets': sheet_names,
        'default_sheet': current_year,
    }


@app.post('/api/auth/login')
def login(payload: LoginPayload, response: Response) -> dict[str, Any]:
    user = find_user(payload.username)
    if not user:
        raise HTTPException(status_code=401, detail='账号或密码错误。')
    password_hash = hash_password(payload.password)
    if not hmac.compare_digest(user.get('password_hash', ''), password_hash):
        raise HTTPException(status_code=401, detail='账号或密码错误。')
    session = create_auth_session(user)
    response.set_cookie(
        key=AUTH_COOKIE_NAME,
        value=session['token'],
        httponly=True,
        samesite='lax',
        max_age=int(AUTH_TTL.total_seconds()),
        path='/',
    )
    return session


@app.post('/api/auth/logout')
def logout(
    response: Response,
    authorization: str | None = Header(default=None),
    auth_cookie: str | None = Cookie(default=None, alias=AUTH_COOKIE_NAME),
) -> dict[str, Any]:
    token = extract_auth_token(authorization, auth_cookie)
    if token:
        AUTH_SESSIONS.pop(token, None)
    response.delete_cookie(AUTH_COOKIE_NAME, path='/')
    return {'ok': True}


@app.get('/api/auth/me')
def me(user: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    return {'user': user}


@app.get('/api/users')
def list_users(_: dict[str, Any] = Depends(require_admin)) -> dict[str, Any]:
    users = [sanitize_user(user) for user in load_users().get('users', []) if isinstance(user, dict)]
    return {'items': users}


@app.post('/api/users')
def create_user(payload: CreateUserPayload, _: dict[str, Any] = Depends(require_admin)) -> dict[str, Any]:
    username = compact_text(payload.username)
    if not username:
        raise HTTPException(status_code=400, detail='用户名不能为空。')
    if find_user(username):
        raise HTTPException(status_code=400, detail='该用户名已存在。')
    role = compact_text(payload.role).lower() or 'user'
    if role not in {'admin', 'user'}:
        raise HTTPException(status_code=400, detail='角色只能是 admin 或 user。')
    users_payload = load_users()
    users = users_payload.get('users', [])
    users.append(
        {
            'username': username,
            'password_hash': hash_password(payload.password),
            'display_name': compact_text(payload.display_name) or username,
            'role': role,
            'created_at': datetime.now().isoformat(timespec='seconds'),
        }
    )
    save_users(users_payload)
    return {'user': sanitize_user(users[-1])}


@app.get('/api/config')
def get_config(_: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    return mask_config(load_config())


@app.post('/api/config')
def update_config(payload: ConfigPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    config = save_config(payload.model_dump())
    start_feishu_ws_client()
    return mask_config(config)


@app.post('/api/config/test')
def test_config(payload: ConfigPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    config = payload.model_dump()
    if not config['base_url'] or not config['api_key'] or not config['model']:
        raise HTTPException(status_code=400, detail='请先填写 URL、API Key 和模型名称。')
    content = call_chat_completion(
        config,
        system_prompt='你是连通性测试助手，只返回 OK。',
        user_prompt='回复 OK',
        temperature=0,
    )
    return {'ok': 'OK' in content.upper(), 'message': content.strip()}


@app.post('/api/feishu/test-message')
def test_feishu_message(payload: FeishuTestPayload, _: dict[str, Any] = Depends(require_admin)) -> dict[str, Any]:
    config = load_config()
    if not config.get('feishu_enabled'):
        raise HTTPException(status_code=400, detail='请先启用飞书机器人。')
    feishu_send_message(config, payload.receive_id_type, payload.receive_id.strip(), payload.text)
    return {'ok': True}


@app.get('/api/feishu/agent-logs')
def feishu_agent_logs(date: str = '', limit: int = 50, _: dict[str, Any] = Depends(require_admin)) -> dict[str, Any]:
    items = load_feishu_agent_logs(date=date, limit=limit)
    return {'items': items, 'count': len(items)}


@app.get('/api/export/feishu-agent-logs')
def export_feishu_agent_logs(date: str = '', limit: int = 500, _: dict[str, Any] = Depends(require_admin)) -> StreamingResponse:
    items = load_feishu_agent_logs(date=date, limit=limit)
    output = io.BytesIO()
    for item in items:
        output.write(json.dumps(item, ensure_ascii=False).encode('utf-8') + b'\n')
    output.seek(0)
    filename = f"feishu-agent-logs-{date or datetime.now().strftime('%Y%m%d')}.jsonl"
    return StreamingResponse(
        output,
        media_type='application/x-ndjson',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/feishu/events')
async def feishu_events(request: Request, background_tasks: BackgroundTasks) -> dict[str, Any]:
    config = load_config()
    try:
        payload = await request.json()
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail='飞书事件 JSON 无效。')
    payload = load_feishu_event_payload(payload, config)
    if not verify_feishu_token(payload, config):
        raise HTTPException(status_code=403, detail='飞书 verification token 校验失败。')
    if payload.get('type') == 'url_verification':
        return {'challenge': payload.get('challenge', '')}
    if not config.get('feishu_enabled'):
        return {'ok': True, 'ignored': 'disabled'}
    if not config.get('feishu_verification_token'):
        return {'ok': True, 'ignored': 'missing_verification_token'}
    message = extract_feishu_message(payload)
    if not message or (message.get('event_type') and message.get('event_type') != 'im.message.receive_v1'):
        return {'ok': True, 'ignored': 'unsupported_event'}
    event_key = message.get('event_id') or message.get('message_id')
    if feishu_event_seen(event_key):
        return {'ok': True, 'duplicate': True}
    background_tasks.add_task(process_feishu_event, message, config)
    return {'ok': True}


@app.get('/api/projects')
def list_projects(
    q: str = '',
    year: str = '',
    award_status: str = '',
    bid_status: str = '',
    _: dict[str, Any] = Depends(require_auth),
) -> dict[str, Any]:
    items, stats = collect_project_items(q=q, year=year, award_status=award_status, bid_status=bid_status)
    return {'items': items, 'stats': stats}


@app.post('/api/projects')
def create_project(payload: ProjectPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    project = save_project(uuid.uuid4().hex, payload)
    return {**project_meta(project), 'project': project}


@app.get('/api/projects/{project_id}')
def get_project(project_id: str, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    project = load_project(project_id)
    session_id = create_session_for_result(project['result'])
    return {**project, 'session_id': session_id}


@app.put('/api/projects/{project_id}')
def update_project(project_id: str, payload: ProjectPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    existing = load_project(project_id)
    project = save_project(project_id, payload, created_at=existing.get('created_at'))
    return {**project_meta(project), 'project': project}


@app.delete('/api/projects/{project_id}')
def delete_project(project_id: str, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    path = project_path(project_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail='项目不存在。')
    path.unlink()
    return {'ok': True}


@app.post('/api/market-skill/extract')
def market_skill_extract(payload: MarketExtractPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    kind = payload.kind.strip()
    if kind not in MARKET_KINDS:
        raise HTTPException(status_code=400, detail='市场情报类型无效。')
    return {'record': extract_market_record(kind, payload.text)}


@app.post('/api/market-skill/report')
def market_skill_report(payload: MarketReportPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    return build_market_report(payload)


@app.get('/api/market-skill/{kind}')
def market_skill_list(
    kind: str,
    q: str = '',
    segment: str = '',
    board_type: str = '',
    status: str = '',
    stage: str = '',
    _: dict[str, Any] = Depends(require_auth),
) -> dict[str, Any]:
    items, stats = list_market_records(kind, q=q, segment=segment, board_type=board_type, status=status, stage=stage)
    return {'items': items, 'stats': stats}


@app.post('/api/market-skill/{kind}')
def market_skill_create(kind: str, payload: MarketRecordPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    record = save_market_record(kind, payload.record)
    return {'record': record}


@app.get('/api/market-skill/{kind}/{record_id}')
def market_skill_get(kind: str, record_id: str, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    return {'record': load_market_record(kind, record_id)}


@app.put('/api/market-skill/{kind}/{record_id}')
def market_skill_update(kind: str, record_id: str, payload: MarketRecordPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    existing = load_market_record(kind, record_id)
    record = save_market_record(kind, payload.record, record_id=record_id, created_at=existing.get('created_at', ''))
    return {'record': record}


@app.delete('/api/market-skill/{kind}/{record_id}')
def market_skill_delete(kind: str, record_id: str, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    path = market_record_path(kind, record_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail='市场情报记录不存在。')
    path.unlink()
    return {'ok': True}


@app.post('/api/export/market-skill')
def export_market_skill(payload: MarketExportPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    kind = payload.kind.strip()
    items, stats = list_market_records(
        kind,
        q=payload.q,
        segment=payload.segment,
        board_type=payload.board_type,
        status=payload.status,
        stage=payload.stage,
    )
    output = build_market_workbook(kind, items, stats)
    prefix = 'cargo_opportunities' if kind == 'cargo' else 'newbuildings'
    filename = f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/market-skill-report')
def export_market_skill_report(payload: MarketReportExportPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    kind = payload.kind.strip()
    if kind not in MARKET_KINDS:
        raise HTTPException(status_code=400, detail='市场情报类型无效。')
    report = payload.report or {}
    prefix = 'cargo_market_report' if kind == 'cargo' else 'newbuilding_market_report'
    if payload.format == 'xlsx':
        output = build_market_report_workbook(kind, report)
        filename = f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        media_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    else:
        output = build_market_report_docx(kind, report)
        filename = f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return StreamingResponse(
        output,
        media_type=media_type,
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/parse')
async def parse_bid(file: UploadFile = File(...), _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    cleanup_sessions()
    suffix = Path(file.filename or '').suffix.lower()
    if suffix not in {'.pdf', '.doc', '.docx', '.xls', '.xlsx'}:
        raise HTTPException(status_code=400, detail='当前只支持 PDF、DOC、DOCX、XLS、XLSX。')
    temp_path = TMP_DIR / f'{uuid.uuid4().hex}{suffix}'
    temp_path.write_bytes(await file.read())
    try:
        document_text = extract_document_text(temp_path, suffix)
    finally:
        temp_path.unlink(missing_ok=True)
    if len(document_text.strip()) < 80:
        raise HTTPException(
            status_code=400,
            detail='当前版本仅支持可提取文本的 PDF、Word、Excel 文件，扫描件、空白文件或只有图片的内容暂不支持。',
        )
    result = parse_ai_document(document_text, require_config())
    session_id = create_session_for_result(result)
    return {'session_id': session_id, 'source_file_name': file.filename or '', **result}


@app.post('/api/chat')
def chat(payload: ChatPayload, _: dict[str, Any] = Depends(require_auth)) -> dict[str, Any]:
    cleanup_sessions()
    session = SESSIONS.get(payload.session_id)
    if not session:
        raise HTTPException(status_code=404, detail='当前会话已失效，请重新解析或重新打开记录。')
    answer = call_chat_completion(
        require_config(),
        system_prompt=CHAT_SYSTEM_PROMPT,
        user_prompt=build_chat_prompt(session, payload.question),
        temperature=0.2,
    )
    return {'answer': answer.strip()}


@app.post('/api/export/extraction')
def export_extraction(payload: ExportExtractionPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    raw_result = payload.result or {'extraction_fields': payload.extraction_fields}
    result = normalize_result_payload(raw_result)
    workbook = load_workbook(get_extraction_template_path())
    value_sheet = workbook[workbook.sheetnames[0]]
    value_sheet.title = '摘取结果'
    for field in EXTRACTION_FIELDS:
        item = result['extraction_fields'].get(field['key'], {})
        value_sheet[field['cell']] = compact_paragraph(item.get('value', ''))

    evidence_sheet = workbook.copy_worksheet(value_sheet)
    evidence_sheet.title = '原文依据'
    for field in EXTRACTION_FIELDS:
        item = result['extraction_fields'].get(field['key'], {})
        evidence_sheet[field['cell']] = resolve_source_excerpt(
            compact_paragraph(item.get('value', '')),
            compact_paragraph(item.get('source_excerpt', '')),
            result.get('document_text_full', ''),
        )

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    filename = f"bid_extraction_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/register')
def export_register(payload: ExportRegisterPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    workbook = load_workbook(get_register_template_path())
    sheet_name = compact_text(payload.sheet_name) or str(datetime.now().year)
    sheet = ensure_register_sheet(workbook, sheet_name)
    row_index = find_register_start_row(sheet)
    previous_no = sheet[f'A{row_index - 1}'].value if row_index > 3 else 0
    try:
        serial = int(previous_no or 0)
    except (TypeError, ValueError):
        serial = row_index - 3

    for offset, row in enumerate(payload.rows, start=1):
        current_row = row_index + offset - 1
        serial += 1
        clean_row = refine_register_row(row if isinstance(row, dict) else {}, normalize_summary({}))
        sheet[f'A{current_row}'] = serial
        sheet[f'B{current_row}'] = parse_excel_date(clean_row.get('receive_date', ''))
        sheet[f'C{current_row}'] = clean_row.get('project_name', '')
        sheet[f'D{current_row}'] = clean_row.get('is_awarded', '')
        sheet[f'E{current_row}'] = clean_row.get('bid_no', '')
        sheet[f'F{current_row}'] = clean_row.get('tenderer', '')
        sheet[f'G{current_row}'] = parse_excel_date(clean_row.get('bid_deadline', ''))
        sheet[f'H{current_row}'] = parse_excel_amount(clean_row.get('deposit_amount', ''))
        sheet[f'I{current_row}'] = clean_row.get('contract_note', '')
        sheet[f'J{current_row}'] = clean_row.get('deposit_return_status', '')
        sheet[f'K{current_row}'] = clean_row.get('payment_method', '')
        sheet[f'L{current_row}'] = clean_row.get('acquisition_method', '')
        sheet[f'M{current_row}'] = clean_row.get('submission_method', '')
        sheet[f'N{current_row}'] = clean_row.get('remark', '')
        apply_register_row_style(sheet, current_row)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    filename = f"bid_register_{sheet_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/overview')
def export_overview(payload: ExportOverviewPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    result = normalize_result_payload(payload.result)
    output = build_overview_workbook(
        payload.title,
        result,
        follow_up=payload.follow_up,
        our_quotes=payload.our_quotes,
        competitor_quotes=payload.competitor_quotes,
        timeline=payload.timeline,
    )
    filename = f"bid_overview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/our-quotes')
def export_our_quotes(payload: ExportQuotePayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    output = build_quote_workbook(
        title=compact_text(payload.title),
        follow_up=payload.follow_up,
        rows=payload.rows,
        is_competitor=False,
    )
    filename = f"our_quotes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/competitor-quotes')
def export_competitor_quotes(payload: ExportQuotePayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    output = build_quote_workbook(
        title=compact_text(payload.title),
        follow_up=payload.follow_up,
        rows=payload.rows,
        is_competitor=True,
    )
    filename = f"competitor_quotes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.post('/api/export/ledger')
def export_ledger(payload: ExportLedgerPayload, _: dict[str, Any] = Depends(require_auth)) -> StreamingResponse:
    output = build_ledger_workbook(
        q=payload.q,
        year=payload.year,
        award_status=payload.award_status,
        bid_status=payload.bid_status,
    )
    year = compact_text(payload.year) or 'all'
    filename = f"history_ledger_{year}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.exception_handler(HTTPException)
async def http_error_handler(_: Any, exc: HTTPException) -> JSONResponse:
    return JSONResponse(status_code=exc.status_code, content={'detail': exc.detail})


def _self_check() -> None:
    users_payload = load_users()
    assert any(user.get('username') == DEFAULT_ADMIN_USERNAME for user in users_payload.get('users', []))
    admin = find_user(DEFAULT_ADMIN_USERNAME)
    assert admin is not None
    assert hmac.compare_digest(admin.get('password_hash', ''), hash_password(DEFAULT_ADMIN_PASSWORD))
    sample = parse_json_text('```json {"ok": true} ```')
    assert sample['ok'] is True
    merged = merge_register_rows(
        [
            {'project_name': '项目A', 'bid_no': '01', 'deposit_amount': '1000', 'remark': ''},
            {'project_name': '项目A', 'bid_no': '02', 'deposit_amount': '2000', 'remark': ''},
        ]
    )
    assert '02' in merged['bid_no']
    assert '2 个标段' in merged['remark']
    assert resolve_source_excerpt('电子化股份有限公司', '...', '福建省东南电化股份有限公司') == '福建省东南电化股份有限公司'
    project = save_project(
        '0' * 32,
        ProjectPayload(
            title='测试项目',
            result={
                'document_summary': {'project_name': '测试项目'},
                'register_rows': [{'project_name': '测试项目'}],
                'document_text_full': '测试项目原文',
            },
        ),
    )
    assert project['title'] == '测试项目'
    assert project_path('0' * 32).exists()
    project_path('0' * 32).unlink(missing_ok=True)


def handle_feishu_message(message: dict[str, Any], config: dict[str, Any]) -> str:
    open_id = message.get('open_id', '')
    chat_id = message.get('chat_id', '')
    text = compact_text(message.get('text') or '')
    files = message.get('files') or []
    if not feishu_message_allowed(config, message):
        reply = '你当前没有权限使用这个飞书 Agent，请联系管理员把你的 open_id 或群 chat_id 加入允许列表。'
        write_feishu_agent_log(message, {}, [], reply, error='unauthorized_feishu_user')
        return reply
    session = feishu_dialog_load(open_id, chat_id)
    try:
        decision = feishu_agent_decide(text, files, session, config)
    except HTTPException as exc:
        if exc.status_code == 400 and 'AI 暂不可用' in str(exc.detail):
            reply = compact_text(exc.detail)
            write_feishu_agent_log(message, {}, [], reply, error=reply)
            return reply
        raise

    tool_results: list[dict[str, Any]] = []
    for call in decision.get('tool_calls') or [{'name': 'chat_general', 'arguments': {}}]:
        result = feishu_agent_execute_tool(call, message, session, config)
        tool_results.append(result)
        session['last_tool_result'] = result

    pending_update = decision.get('pending_update') if isinstance(decision.get('pending_update'), dict) else {}
    if pending_update:
        for key in ('agent_summary', 'last_question'):
            if key in pending_update:
                session[key] = compact_text(pending_update.get(key))
    if text:
        session['last_question'] = text
    feishu_dialog_save(open_id, chat_id, session)

    explicit_reply = compact_text(decision.get('reply'))
    if explicit_reply and not any(result.get('tool') != 'chat_general' for result in tool_results):
        write_feishu_agent_log(message, decision, tool_results, explicit_reply)
        return explicit_reply
    try:
        reply = feishu_agent_chat_reply(text, session, config, tool_results)
        if reply:
            write_feishu_agent_log(message, decision, tool_results, reply)
            return reply
    except HTTPException:
        pass
    for result in tool_results:
        if result.get('reply'):
            reply = sanitize_feishu_reply(result.get('reply'))
            write_feishu_agent_log(message, decision, tool_results, reply)
            return reply
        if result.get('error'):
            reply = f"处理时遇到问题：{compact_text(result.get('error'))}"
            write_feishu_agent_log(message, decision, tool_results, reply, error=compact_text(result.get('error')))
            return reply
    reply = 'AI 已处理这条消息，但没有生成可发送的回复。'
    write_feishu_agent_log(message, decision, tool_results, reply, error='empty_reply')
    return reply


def process_feishu_event(message: dict[str, Any], config: dict[str, Any]) -> None:
    try:
        reply = handle_feishu_message(message, config)
    except HTTPException as exc:
        reply = f'处理时遇到问题：{compact_text(exc.detail)}'
        write_feishu_agent_log(message, {}, [], reply, error=compact_text(exc.detail))
    except Exception as exc:
        print(f'Feishu agent error: {exc}')
        reply = '处理时遇到问题：AI 暂时没法完成这次请求，请稍后再试。'
        write_feishu_agent_log(message, {}, [], reply, error=str(exc))
    try:
        feishu_reply_message(config, message.get('message_id', ''), sanitize_feishu_reply(reply))
    except Exception as exc:
        print(f'Feishu reply failed: {exc}')


if __name__ == '__main__':
    _self_check()
    uvicorn.run('app:app', host='127.0.0.1', port=8008, reload=False, app_dir=str(BASE_DIR))
